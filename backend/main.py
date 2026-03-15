from __future__ import annotations

import argparse
import json
import re
import shutil
import threading
import time
import zipfile
from datetime import date, datetime
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

from backend.agent_executor import (
    build_command,
    executor_exists,
    get_settings,
    launch_job,
    launch_local_terminal,
    read_job_log,
    save_settings,
)
from backend.families import get_family, get_prompt_guidance, list_families
from backend.intake import answers_to_instructions, get_intake_questions
from backend.providers import fetch_provider_data, list_providers, set_settings_path
from backend.storage import DATA_ROOT, JOB_DIR, PROJECT_DIR, TEMPLATE_DIR, copy_any, ensure_storage, read_state, write_json, write_state

SKILLS_DIR = Path(__file__).resolve().parents[1] / "skills"

# Map each family to the skill directories most useful to it
_FAMILY_SKILL_DIRS: dict[str, list[str]] = {
    "equity-research": [
        "equity-research/skills/initiating-coverage",
        "equity-research/skills/earnings-analysis",
        "financial-analysis/skills/dcf-model",
        "financial-analysis/skills/comps-analysis",
        "financial-analysis/skills/3-statement-model",
    ],
    "quarterly-stock-update": [
        "equity-research/skills/earnings-analysis",
        "equity-research/skills/model-update",
        "financial-analysis/skills/3-statement-model",
    ],
    "case-comp": [
        "equity-research/skills/initiating-coverage",
        "financial-analysis/skills/dcf-model",
        "financial-analysis/skills/comps-analysis",
        "investment-banking/skills/pitch-deck",
    ],
    "macro-update": [
        "financial-analysis/skills/competitive-analysis",
    ],
    "commodity-report": [
        "financial-analysis/skills/competitive-analysis",
    ],
    "weekly-commodity-update": [
        "financial-analysis/skills/competitive-analysis",
    ],
}

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

try:
    import pypdfium2 as pdfium
except Exception:  # pragma: no cover
    pdfium = None


def _merge_instructions(custom: str, intake_answers: dict) -> str:
    if not intake_answers:
        return custom
    intake_block = answers_to_instructions(intake_answers)
    return f"{intake_block}\n\n{custom}".strip() if custom else intake_block


def now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def slugify(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")


def entity_id(prefix: str, name: str) -> str:
    return f"{prefix}_{datetime.now().strftime('%Y%m%d%H%M%S')}_{slugify(name)[:36]}"


def sort_desc(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return sorted(items, key=lambda item: item.get("created_at", ""), reverse=True)


def get_state() -> dict[str, Any]:
    state = read_state()
    return {
        "data_root": DATA_ROOT.as_posix(),
        "finder_notes_path": (Path(__file__).resolve().parents[1] / "notes" / "report-finder-recipes.md").as_posix(),
        "executor_name": "Codex CLI",
        "executor_available": executor_exists(),
        "templates": sort_desc(state["templates"]),
        "projects": sort_desc(state["projects"]),
        "jobs": sort_desc(state["jobs"]),
        "update_definitions": sort_desc(state["update_definitions"]),
        "families": list_families(),
        "providers": list_providers(),
    }


def read_json_body(handler: BaseHTTPRequestHandler) -> dict[str, Any]:
    content_length = int(handler.headers.get("Content-Length", "0"))
    raw = handler.rfile.read(content_length) if content_length else b"{}"
    return json.loads(raw.decode("utf-8") or "{}")


def send_json(handler: BaseHTTPRequestHandler, payload: Any, status: int = 200) -> None:
    encoded = json.dumps(payload).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json")
    handler.send_header("Content-Length", str(len(encoded)))
    handler.send_header("Access-Control-Allow-Origin", "*")
    handler.send_header("Access-Control-Allow-Headers", "Content-Type")
    handler.send_header("Access-Control-Allow-Methods", "GET,POST,OPTIONS")
    handler.end_headers()
    handler.wfile.write(encoded)


def send_error_json(handler: BaseHTTPRequestHandler, message: str, status: int = 400) -> None:
    send_json(handler, {"detail": message}, status)


def find_record(items: list[dict[str, Any]], record_id: str) -> dict[str, Any] | None:
    for item in items:
        if item["id"] == record_id:
            return item
    return None


def ensure_workspace_dirs(workspace: Path) -> dict[str, Path]:
    mapping = {
        "source_uploads": workspace / "source" / "uploads",
        "source_downloads": workspace / "source" / "downloads",
        "source_urls": workspace / "source" / "urls",
        "source_transcripts": workspace / "source" / "transcripts",
        "source_filings": workspace / "source" / "filings",
        "extracted_text": workspace / "extracted" / "text",
        "extracted_images": workspace / "extracted" / "images",
        "extracted_tables": workspace / "extracted" / "tables",
        "extracted_metadata": workspace / "extracted" / "metadata",
        "context": workspace / "context",
        "templates": workspace / "templates" / "selected",
        "generated_charts": workspace / "generated" / "charts",
        "generated_diagrams": workspace / "generated" / "diagrams",
        "generated_excel": workspace / "generated" / "excel",
        "result": workspace / "result",
        "logs": workspace / "logs",
    }
    for path in mapping.values():
        path.mkdir(parents=True, exist_ok=True)
    return mapping


def extract_pdf_text(source: Path) -> str:
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(str(source))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        return ""


def extract_pdf_images(source: Path, dest_dir: Path) -> list[str]:
    if pdfium is None:
        return []
    try:
        pdf = pdfium.PdfDocument(str(source))
        saved: list[str] = []
        for i in range(len(pdf)):
            page = pdf[i]
            bitmap = page.render(scale=150 / 72)
            img = bitmap.to_pil()
            out_path = dest_dir / f"{source.stem}_page{i + 1}.png"
            img.save(str(out_path))
            saved.append(str(out_path))
        return saved
    except Exception:
        return []


def extract_text_like(source: Path) -> str:
    if source.is_dir():
        return ""
    lower = source.suffix.lower()
    if lower == ".pdf":
        return extract_pdf_text(source)
    if lower in {".txt", ".md", ".csv"}:
        return source.read_text(encoding="utf-8", errors="ignore")
    if lower == ".docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(source))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception:
            pass
        # fallback to XML approach
        try:
            with zipfile.ZipFile(source) as archive:
                snippets: list[str] = []
                for member in archive.namelist():
                    if member.endswith(".xml") and "word/" in member:
                        snippets.append(archive.read(member).decode("utf-8", errors="ignore"))
                return "\n".join(snippets[:20])
        except Exception:
            return ""
    if lower in {".pptx", ".xlsx"}:
        try:
            with zipfile.ZipFile(source) as archive:
                snippets: list[str] = []
                for member in archive.namelist():
                    if member.endswith(".xml") and ("ppt/" in member or "xl/" in member):
                        snippets.append(archive.read(member).decode("utf-8", errors="ignore"))
                return "\n".join(snippets[:20])
        except Exception:
            return ""
    return ""


def build_prompt(job: dict[str, Any], template: dict[str, Any] | None) -> str:
    lines = [
        f"You are preparing a final {job['family']} deliverable in {job['output_format']} format.",
        f"Job kind: {job['kind']}",
        f"Objective: {job['objective']}",
        f"Workspace root: {job['workspace_path']}",
        f"Write completed deliverables only into: {job['result_path']}",
        (
            "Workspace layout — all paths are relative to the workspace root:\n"
            "  source/uploads/         — user-uploaded files (broker PDFs, research reports, spreadsheets, DOCX)\n"
            "  source/filings/         — SEC filings, regulatory documents\n"
            "  source/downloads/       — files fetched from URLs\n"
            "  source/transcripts/     — earnings call and conference transcripts\n"
            "  extracted/text/         — text extracted from every source file (<stem>.txt)\n"
            "  extracted/images/       — page renders from PDFs (<stem>_page<N>.png) and cropped visuals\n"
            "  extracted/tables/       — tabular data parsed from source files\n"
            "  extracted/metadata/     — assets.json manifest of all staged files\n"
            "  context/                — prompt.md, intake.json, objective.md, questions.md, answers.md\n"
            "  templates/selected/     — any user-supplied template files\n"
            "  generated/charts/       — charts and annotated plots you produce (PNG/SVG)\n"
            "  generated/diagrams/     — flow diagrams, architecture visuals\n"
            "  generated/excel/        — valuation workbooks and data tables\n"
            "  result/                 — FINAL deliverable output only (write your finished file here)\n"
            "  logs/                   — status.json and launch-command.txt"
        ),
        (
            "Source-file pipeline:\n"
            "  1. Broker PDFs / research reports → staged in source/uploads/\n"
            "     Text extracts already written to extracted/text/<filename>.txt\n"
            "     Page renders already written to extracted/images/<filename>_page<N>.png\n"
            "  2. Read text extracts first for efficient content discovery.\n"
            "  3. Reference page renders when you need charts, tables, or visual data from the source PDFs.\n"
            "  4. Build any new charts with matplotlib (save PNG to generated/charts/) before assembling the output.\n"
            "  5. Assemble the final document:\n"
            "     - DOCX: use python-docx\n"
            "     - PPTX: use python-pptx\n"
            "     - PDF: use reportlab or weasyprint\n"
            "  6. Save the completed file to result/ — that is the only directory the user sees."
        ),
        "If information is missing, write follow-up questions into context/questions.md and wait for answers in context/answers.md.",
        "Do not produce partial artifacts unless the prompt explicitly asks for one.",
    ]
    if job.get("valuation_required"):
        lines.append(
            "A valuation workbook is required. Build it in generated/excel/ using openpyxl, "
            "then reference it from the final document. Include DCF, comparable company analysis, "
            "and sensitivity tables."
        )
    if template:
        lib = Path(template.get("library_path", ""))
        if lib.exists() and lib.suffix == ".md":
            lines.append(
                f"A structured template file has been staged at templates/selected/{lib.name}. "
                "Read it carefully and follow every instruction in it, including chart, Excel, and document assembly steps."
            )
        else:
            lines.append(f"Follow the selected template style: {template['name']} ({template['family']}).")
    staged_skills = job.get("_staged_skills", [])
    if staged_skills:
        lines.append(
            "Institutional skill packages are staged at templates/selected/skills/. "
            "Each skill directory contains: SKILL.md (entry point), references/ (detailed task instructions), "
            "assets/ (quality checklists, report templates), and scripts/ (validation helpers). "
            "Read the SKILL.md entry points first, then load the specific reference file for the task you are executing:\n"
            + "\n".join(f"  - templates/selected/{s}" for s in staged_skills)
        )
    family_guidance = get_prompt_guidance(job["family"])
    if family_guidance:
        lines.append(family_guidance)
    if job.get("custom_instructions"):
        lines.append(f"Custom instructions:\n{job['custom_instructions']}")
    if job.get("question_prompts"):
        lines.append("Questions to answer in the final output:")
        lines.extend(f"- {question}" for question in job["question_prompts"])
    return "\n\n".join(lines)


def stage_template(template: dict[str, Any] | None, destination: Path) -> None:
    if not template:
        return
    source = Path(template["library_path"])
    if source.exists():
        copy_any(source, destination / source.name)


def stage_skills(family: str, destination: Path) -> list[str]:
    """Copy relevant skill trees (SKILL.md + all references + scripts) into templates/selected/."""
    staged: list[str] = []
    skill_dirs = _FAMILY_SKILL_DIRS.get(family, [])
    destination.mkdir(parents=True, exist_ok=True)
    for rel_dir in skill_dirs:
        skill_root = SKILLS_DIR / rel_dir
        if not skill_root.exists():
            continue
        # Copy every file under this skill dir, preserving relative paths
        for src in skill_root.rglob("*"):
            if src.is_file() and not src.name.startswith("."):
                rel = src.relative_to(skill_root)
                prefix = rel_dir.replace("/", "_").replace("-", "_")
                dest_rel = f"skills/{prefix}/{rel}"
                dest_file = destination / dest_rel
                dest_file.parent.mkdir(parents=True, exist_ok=True)
                dest_file.write_bytes(src.read_bytes())
                if src.name == "SKILL.md":
                    staged.append(dest_rel)
    return staged


def prepare_sources(job: dict[str, Any], workspace_dirs: dict[str, Path]) -> list[dict[str, str]]:
    assets: list[dict[str, str]] = []
    pdfs_to_render: list[tuple[Path, Path]] = []

    for raw_path in job["source_paths"]:
        source = Path(raw_path).expanduser()
        if not source.exists():
            continue
        target = workspace_dirs["source_uploads"] / source.name
        copy_any(source, target)
        # Text extraction is fast — do it synchronously so the prompt can reference it
        extracted = extract_text_like(source)
        if extracted:
            (workspace_dirs["extracted_text"] / f"{source.stem}.txt").write_text(extracted, encoding="utf-8")
        # Image extraction is slow (renders every PDF page) — defer to background thread
        if source.suffix.lower() == ".pdf":
            pdfs_to_render.append((source, workspace_dirs["extracted_images"]))
        assets.append({"source": source.as_posix(), "copied_to": target.as_posix()})

    if pdfs_to_render:
        def _render_pages() -> None:
            for src, dest in pdfs_to_render:
                extract_pdf_images(src, dest)

        threading.Thread(target=_render_pages, daemon=True).start()

    if job["urls"]:
        (workspace_dirs["source_urls"] / "urls.txt").write_text("\n".join(job["urls"]), encoding="utf-8")
    return assets


def create_workspace(job: dict[str, Any], template: dict[str, Any] | None) -> dict[str, Any]:
    workspace = Path(job["workspace_path"])
    dirs = ensure_workspace_dirs(workspace)
    assets = prepare_sources(job, dirs)
    stage_template(template, dirs["templates"])
    staged_skills = stage_skills(job["family"], dirs["templates"])
    job["_staged_skills"] = staged_skills

    write_json(
        dirs["context"] / "intake.json",
        {
            "title": job["title"],
            "kind": job["kind"],
            "family": job["family"],
            "output_format": job["output_format"],
            "providers": job["provider_names"],
            "urls": job["urls"],
            "source_paths": job["source_paths"],
            "valuation_required": job["valuation_required"],
        },
    )
    provider_data = {"providers": job["provider_names"], "status": "stubbed"}
    if job["provider_names"] and job.get("source_paths"):
        try:
            provider_data["results"] = fetch_provider_data(job["provider_names"], [])
            provider_data["status"] = "fetched"
        except Exception:
            pass
    write_json(dirs["context"] / "provider-data.json", provider_data)
    write_json(dirs["extracted_metadata"] / "assets.json", assets)
    prompt = build_prompt(job, template)
    (dirs["context"] / "objective.md").write_text(job["objective"], encoding="utf-8")
    (dirs["context"] / "questions.md").write_text(
        "\n".join(f"- {question}" for question in job["question_prompts"]) or "No pending questions yet.",
        encoding="utf-8",
    )
    (dirs["context"] / "answers.md").write_text("", encoding="utf-8")
    (dirs["context"] / "prompt.md").write_text(prompt, encoding="utf-8")
    write_json(dirs["context"] / "run-config.json", {
        "workspace_path": job["workspace_path"],
        "result_path": job["result_path"],
        "prompt_path": (dirs["context"] / "prompt.md").as_posix(),
        "provider_flags": job["provider_names"],
        "executor": "codex",
        "status": job["status"],
        "family": job["family"],
        "output_format": job["output_format"],
        "valuation_required": job.get("valuation_required", False),
    })
    (dirs["context"] / "run-codex.txt").write_text(
        build_command(workspace, dirs["context"] / "prompt.md"), encoding="utf-8"
    )
    (dirs["logs"] / "status.json").write_text(
        json.dumps({"status": job["status"], "updated_at": job["updated_at"]}, indent=2), encoding="utf-8"
    )
    job["prompt_preview"] = prompt
    return job


def create_job_record(payload: dict[str, Any]) -> dict[str, Any]:
    state = read_state()
    template = find_record(state["templates"], payload.get("template_id")) if payload.get("template_id") else None
    project = find_record(state["projects"], payload.get("project_id")) if payload.get("project_id") else None

    job_id = entity_id("job", payload["title"])
    workspace_root = JOB_DIR / job_id
    record = {
        "id": job_id,
        "title": payload["title"],
        "kind": payload["kind"],
        "family": payload["family"],
        "objective": payload["objective"],
        "output_format": payload.get("output_format", "pptx"),
        "project_id": payload.get("project_id"),
        "template_id": payload.get("template_id"),
        "provider_names": payload.get("provider_names", []),
        "source_paths": payload.get("source_paths", []),
        "urls": payload.get("urls", []),
        "custom_instructions": _merge_instructions(
            payload.get("custom_instructions", ""),
            payload.get("intake_answers", {}),
        ),
        "question_prompts": payload.get("question_prompts", []),
        # templates that contain Excel instructions force valuation on
        "valuation_required": bool(payload.get("valuation_required"))
            or (template is not None and "xlsx" in template.get("notes", "").lower()),
        "cadence": payload.get("cadence"),
        "status": "ready",
        "workspace_path": workspace_root.as_posix(),
        "result_path": (workspace_root / "result").as_posix(),
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "prompt_preview": "",
        "question_log": [],
    }
    if project:
        linked_jobs = Path(project["root_path"]) / "linked-jobs"
        linked_jobs.mkdir(parents=True, exist_ok=True)
        (linked_jobs / f"{job_id}.txt").write_text(workspace_root.as_posix(), encoding="utf-8")

    create_workspace(record, template)
    state["jobs"].append(record)
    write_state(state)
    return record


def handle_template_register(payload: dict[str, Any]) -> tuple[dict[str, Any], int]:
    source_path = Path(payload["source_path"]).expanduser()
    if not source_path.exists():
        return {"detail": "Template source path does not exist"}, 400
    state = read_state()
    template_id = entity_id("tmpl", payload["name"])
    template_root = TEMPLATE_DIR / template_id
    template_root.mkdir(parents=True, exist_ok=True)
    library_path = template_root / source_path.name
    copy_any(source_path, library_path)
    record = {
        "id": template_id,
        "name": payload["name"],
        "family": payload.get("family", "general"),
        "output_formats": payload.get("output_formats", []),
        "source_path": source_path.as_posix(),
        "library_path": library_path.as_posix(),
        "notes": payload.get("notes", ""),
        "created_at": now_iso(),
    }
    state["templates"].append(record)
    write_state(state)
    return record, 200


def handle_project_create(payload: dict[str, Any]) -> tuple[dict[str, Any], int]:
    state = read_state()
    project_id = entity_id("proj", payload["name"])
    project_root = PROJECT_DIR / project_id
    project_root.mkdir(parents=True, exist_ok=True)
    record = {
        "id": project_id,
        "name": payload["name"],
        "root_path": project_root.as_posix(),
        "created_at": now_iso(),
    }
    (project_root / "README.md").write_text(
        f"# {record['name']}\n",
        encoding="utf-8",
    )
    state["projects"].append(record)
    write_state(state)
    return record, 200


def handle_job_launch(job_id: str, executor: str = "codex") -> tuple[dict[str, Any], int]:
    state = read_state()
    job = find_record(state["jobs"], job_id)
    if not job:
        return {"detail": "Job not found"}, 404
    if not executor_exists(executor):
        return {"detail": f"{executor} is not installed on this machine"}, 400
    workspace = Path(job["workspace_path"])
    prompt_path = workspace / "context" / "prompt.md"
    command = launch_job(workspace, prompt_path, job_id, executor)
    job["status"] = "agent_running"
    job["executor"] = executor
    job["updated_at"] = now_iso()
    (workspace / "logs" / "launch-command.txt").write_text(command, encoding="utf-8")
    (workspace / "logs" / "status.json").write_text(
        json.dumps({"status": job["status"], "updated_at": job["updated_at"]}, indent=2), encoding="utf-8"
    )
    write_state(state)
    return {"status": "launched", "command": command}, 200


def handle_job_logs(job_id: str) -> tuple[Any, int]:
    state = read_state()
    job = find_record(state["jobs"], job_id)
    if not job:
        return {"detail": "Job not found"}, 404
    log_path = Path(job["workspace_path"]) / "logs" / "codex.log"
    return {"log": read_job_log(log_path)}, 200


def handle_job_questions(job_id: str) -> tuple[Any, int]:
    state = read_state()
    job = find_record(state["jobs"], job_id)
    if not job:
        return {"detail": "Job not found"}, 404
    q_path = Path(job["workspace_path"]) / "context" / "questions.md"
    content = q_path.read_text(encoding="utf-8") if q_path.exists() else ""
    return {"questions": content}, 200


def handle_get_settings() -> tuple[dict[str, Any], int]:
    s = get_settings()
    # Never expose the raw API key — just confirm it's set
    return {"api_key_set": bool(s.get("openai_api_key")), "executor": s.get("executor", "codex")}, 200


def handle_save_settings(payload: dict[str, Any]) -> tuple[dict[str, Any], int]:
    s = get_settings()
    if "openai_api_key" in payload:
        s["openai_api_key"] = payload["openai_api_key"]
    if "executor" in payload:
        s["executor"] = payload["executor"]
    save_settings(s)
    return {"ok": True}, 200


def handle_intake_questions(payload: dict[str, Any]) -> tuple[dict[str, Any], int]:
    questions = get_intake_questions(
        objective=payload.get("objective", ""),
        family=payload.get("family", "equity-research"),
        use_ai=True,
    )
    return {"questions": questions}, 200


def handle_finder_run(payload: dict[str, Any]) -> tuple[dict[str, Any], int]:
    from backend.browser_worker import run_recipe_background, get_worker_status
    recipe = payload.get("recipe", {})
    download_dir = Path(payload.get("download_dir", str(DATA_ROOT / "downloads")))
    recipe_id = entity_id("recipe", recipe.get("site", "finder"))
    run_recipe_background(recipe, download_dir, recipe_id)
    return {"recipe_id": recipe_id, "status": "running"}, 200


def handle_finder_log(recipe_id: str) -> tuple[dict[str, Any], int]:
    from backend.browser_worker import get_worker_log, get_worker_status
    return {"log": get_worker_log(recipe_id), "status": get_worker_status(recipe_id)}, 200


def handle_job_qa(job_id: str, payload: dict[str, Any]) -> tuple[dict[str, Any], int]:
    state = read_state()
    job = find_record(state["jobs"], job_id)
    if not job:
        return {"detail": "Job not found"}, 404
    template = find_record(state["templates"], job.get("template_id")) if job.get("template_id") else None
    workspace = Path(job["workspace_path"])
    with (workspace / "context" / "answers.md").open("a", encoding="utf-8") as handle:
        handle.write(f"\n- {payload['content']}\n")
    job["question_log"].append({"role": "user", "content": payload["content"], "timestamp": now_iso()})
    job["updated_at"] = now_iso()
    job["prompt_preview"] = build_prompt(job, template)
    (workspace / "context" / "prompt.md").write_text(job["prompt_preview"], encoding="utf-8")
    write_state(state)
    return job, 200


def handle_update_definition(payload: dict[str, Any]) -> tuple[dict[str, Any], int]:
    state = read_state()
    record = {
        "id": entity_id("upd", payload["name"]),
        "name": payload["name"],
        "cadence": payload.get("cadence", "adhoc"),
        "family": payload.get("family", "macro-update"),
        "output_format": payload.get("output_format", "pdf"),
        "instruments": payload.get("instruments", []),
        "template_id": payload.get("template_id"),
        "created_at": now_iso(),
    }
    state["update_definitions"].append(record)
    write_state(state)
    return record, 200


def handle_update_run(definition_id: str) -> tuple[dict[str, Any], int]:
    state = read_state()
    definition = find_record(state["update_definitions"], definition_id)
    if not definition:
        return {"detail": "Update definition not found"}, 404
    return (
        create_job_record(
            {
                "kind": "update",
                "title": f"{definition['name']} {datetime.now().strftime('%Y-%m-%d')}",
                "family": definition["family"],
                "objective": f"Create a {definition['cadence']} update covering: {', '.join(definition['instruments'])}. Include annotated price charts, custom diagrams, and a fully polished final output.",
                "output_format": definition["output_format"],
                "template_id": definition.get("template_id"),
                "provider_names": ["premium-primary"],
                "question_prompts": [
                    "What changed across the tracked instruments?",
                    "Which moves are narrative noise versus actionable positioning shifts?",
                ],
                "cadence": definition["cadence"],
            }
        ),
        200,
    )


class RequestHandler(BaseHTTPRequestHandler):
    def do_OPTIONS(self) -> None:  # noqa: N802
        send_json(self, {"ok": True}, 200)

    def do_GET(self) -> None:  # noqa: N802
        path = urlparse(self.path).path
        if path == "/health":
            send_json(self, {"status": "ok"})
            return
        if path == "/app-state":
            send_json(self, get_state())
            return
        if path == "/families":
            send_json(self, list_families())
            return
        if path == "/providers":
            send_json(self, list_providers())
            return
        if path == "/settings":
            response, status = handle_get_settings()
            send_json(self, response, status)
            return

        job_logs = re.fullmatch(r"/jobs/([^/]+)/logs", path)
        if job_logs:
            response, status = handle_job_logs(job_logs.group(1))
            send_json(self, response, status)
            return

        job_questions = re.fullmatch(r"/jobs/([^/]+)/questions", path)
        if job_questions:
            response, status = handle_job_questions(job_questions.group(1))
            send_json(self, response, status)
            return

        finder_log = re.fullmatch(r"/finder/([^/]+)/log", path)
        if finder_log:
            response, status = handle_finder_log(finder_log.group(1))
            send_json(self, response, status)
            return

        send_error_json(self, "Route not found", 404)

    def do_POST(self) -> None:  # noqa: N802
        path = urlparse(self.path).path
        payload = read_json_body(self)

        if path == "/templates/register":
            response, status = handle_template_register(payload)
            send_json(self, response, status)
            return
        if path == "/projects":
            response, status = handle_project_create(payload)
            send_json(self, response, status)
            return
        if path == "/jobs":
            send_json(self, create_job_record(payload))
            return
        if path == "/update-definitions":
            response, status = handle_update_definition(payload)
            send_json(self, response, status)
            return
        if path == "/settings":
            response, status = handle_save_settings(payload)
            send_json(self, response, status)
            return
        if path == "/finder/run":
            response, status = handle_finder_run(payload)
            send_json(self, response, status)
            return
        if path == "/intake/questions":
            response, status = handle_intake_questions(payload)
            send_json(self, response, status)
            return

        job_launch = re.fullmatch(r"/jobs/([^/]+)/launch", path)
        if job_launch:
            executor = payload.get("executor", "codex")
            response, status = handle_job_launch(job_launch.group(1), executor)
            send_json(self, response, status)
            return

        job_qa = re.fullmatch(r"/jobs/([^/]+)/qa", path)
        if job_qa:
            response, status = handle_job_qa(job_qa.group(1), payload)
            send_json(self, response, status)
            return

        update_run = re.fullmatch(r"/update-definitions/([^/]+)/run", path)
        if update_run:
            response, status = handle_update_run(update_run.group(1))
            send_json(self, response, status)
            return

        send_error_json(self, "Route not found", 404)

    def log_message(self, format: str, *args: Any) -> None:  # noqa: A003
        return


def _fire_due_updates() -> None:
    """Check for scheduled update definitions that are due and create jobs for them."""
    today = date.today()
    state = read_state()
    changed = False
    for defn in state["update_definitions"]:
        cadence = defn.get("cadence", "adhoc")
        if cadence == "adhoc":
            continue
        last_run = defn.get("last_run_date")
        if cadence == "daily":
            due = not last_run or last_run != str(today)
        elif cadence == "weekly":
            last = date.fromisoformat(last_run) if last_run else None
            due = not last or (today - last).days >= 7
        else:
            continue
        if due:
            try:
                create_job_record({
                    "kind": "update",
                    "title": f"{defn['name']} {today}",
                    "family": defn["family"],
                    "objective": (
                        f"Create a {defn['cadence']} {defn['family']} update "
                        f"covering: {', '.join(defn['instruments'])}."
                    ),
                    "output_format": defn["output_format"],
                    "template_id": defn.get("template_id"),
                    "provider_names": ["stub"],
                    "source_paths": [],
                    "urls": [],
                    "custom_instructions": "",
                    "question_prompts": [],
                    "valuation_required": False,
                    "cadence": cadence,
                })
                defn["last_run_date"] = str(today)
                changed = True
                print(f"[scheduler] fired: {defn['name']}")
            except Exception as exc:
                print(f"[scheduler] error firing {defn['name']}: {exc}")
    if changed:
        write_state(state)


def _scheduler_loop() -> None:
    """Background thread: wake every 60 s and check for due update definitions."""
    while True:
        time.sleep(60)
        try:
            _fire_due_updates()
        except Exception as exc:
            print(f"[scheduler] loop error: {exc}")


_EQUITY_RESEARCH_TEMPLATE_MD = """\
# Equity Research Report — Full Pipeline Template

## Purpose
This template drives a full institutional-grade equity research report.
It is company-agnostic: replace every `[COMPANY]` / `[TICKER]` placeholder
with the subject of the current job.

---

## Required output files (both must appear in result/)
1. `result/report.docx`   — full research report (python-docx)
2. `result/valuation.xlsx` — standalone valuation workbook (openpyxl, NO template required)

---

## Report sections (in order)

### 1. Executive Summary
- One-paragraph investment recommendation (Buy / Hold / Sell)
- Price target with percentage upside/downside from current price
- Three bullet-point key catalysts

### 2. Company Overview
- Business description: what the company does, key products/segments, end markets
- Revenue breakdown by segment (table)
- Key financial metrics snapshot (market cap, EV, EV/EBITDA, P/E, revenue CAGR)

### 3. Latest Earnings Analysis
- Beat/miss table: for each KPI (Revenue, Gross Margin, Operating Margin, EPS, EBITDA)
  include columns: Actual | Street Consensus | Beat/Miss | % vs. Consensus
- Forward guidance vs. prior Street consensus
- Management commentary highlights (verbatim where significant)

### 4. Investment Thesis
Three numbered pillars, each with: thesis statement → evidence → quantification

### 5. Competitive Positioning
- Peer comparison table: Company | Revenue | Rev Growth | Gross Margin | EV/EBITDA | P/E
- Moat assessment (switching costs, IP, scale, network effects)
- Key threats and differentiation

### 6. Valuation & Price Target
- Comparable company (comps) analysis — use the peer table from section 5
- DCF model summary: WACC, terminal growth rate, implied price
- Bull / base / bear scenario table with price targets and key assumptions
- Final price target derivation

### 7. Key Risks
Ranked table: Risk | Probability (H/M/L) | Impact (H/M/L) | Mitigant

### 8. Catalysts & Timeline
Table: Catalyst | Expected Date | Potential Impact

---

## Charts to build (matplotlib, save as PNG to generated/charts/)

1. `revenue_growth.png`     — bar chart: quarterly revenue with YoY growth line overlay
2. `margin_trends.png`      — multi-line chart: gross margin %, operating margin %, net margin %
3. `valuation_comps.png`    — horizontal bar chart: EV/EBITDA across peer group
4. `beat_miss_history.png`  — bar chart: EPS beat/miss vs. consensus over last 8 quarters
5. `price_target_bridge.png` — waterfall or lollipop chart: bear / base / bull scenarios

Use a clean, minimal style (white background, navy/charcoal bars, gold accents for highlights).
Label all axes, include a legend, save at 150 dpi minimum.

---

## Excel Valuation Workbook (result/valuation.xlsx)
Build with openpyxl. Create these sheets:

### Sheet 1: DCF
| Row | Column A | Column B → Column H |
|-----|----------|---------------------|
| 1   | WACC assumption | (value) |
| 2   | Terminal growth rate | (value) |
| 3   | Projection year | Year 1 … Year 7 |
| 4   | Revenue ($M) | formula-driven |
| 5   | EBIT margin | assumption |
| 6   | NOPAT | =Revenue × EBIT margin |
| 7   | D&A | assumption |
| 8   | CapEx | assumption |
| 9   | Change in NWC | assumption |
| 10  | Free Cash Flow | =NOPAT+D&A-CapEx-ΔNWC |
| 11  | Discount factor | =1/(1+WACC)^year |
| 12  | PV of FCF | =FCF × discount factor |
| 13  | Terminal Value | Gordon Growth on Year 7 FCF |
| 14  | Implied Equity Value per Share | |

### Sheet 2: Comps
Columns: Company | EV ($M) | Revenue | EBITDA | EV/Revenue | EV/EBITDA | P/E | NTM P/E
At least 5 comparable companies. Include a median/mean row.

### Sheet 3: Scenarios
Rows: Bear | Base | Bull
Columns: Revenue CAGR | EBIT Margin | Terminal Multiple | Equity Value/Share | vs. Current Price

Style the workbook: navy header rows (hex #1E2D6F, white font), alternating row shading,
bold totals. No external template needed — build all formatting in code.

---

## Document assembly (python-docx)
- Use Heading 1 for section titles, Heading 2 for subsections
- Insert each chart PNG using `doc.add_picture()` with appropriate width (5–6 inches)
- Caption every chart and table
- Apply a clean sans-serif body font (Calibri 11pt)
- Page margins: 1 inch all sides
- Include a header with [COMPANY] — Equity Research and page numbers in the footer
- Save as `result/report.docx`

---

## Workflow checklist
- [ ] Read all text from extracted/text/ before writing anything
- [ ] Build all 5 charts and save to generated/charts/
- [ ] Build valuation.xlsx and save to result/valuation.xlsx
- [ ] Assemble report.docx referencing the charts
- [ ] Save report.docx to result/report.docx
- [ ] Write any blocking questions to context/questions.md before stopping
"""

_EQUITY_RESEARCH_TEMPLATE_VERSION = "v2"  # bump to force re-seed


def _seed_builtin_equity_template() -> None:
    """Seed (or update) the built-in full-pipeline equity research template."""
    state = read_state()

    existing = next(
        (t for t in state["templates"] if t.get("id") == "tmpl_builtin_equity_full"), None
    )
    if existing and existing.get("version") == _EQUITY_RESEARCH_TEMPLATE_VERSION:
        return  # already up to date

    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    template_root = TEMPLATE_DIR / "tmpl_builtin_equity_full"
    template_root.mkdir(parents=True, exist_ok=True)
    md_path = template_root / "EQUITY_RESEARCH_TEMPLATE.md"
    md_path.write_text(_EQUITY_RESEARCH_TEMPLATE_MD, encoding="utf-8")

    record = {
        "id": "tmpl_builtin_equity_full",
        "name": "Equity Research — Full Pipeline",
        "family": "equity-research",
        "output_formats": ["docx"],
        "source_path": md_path.as_posix(),
        "library_path": md_path.as_posix(),
        "version": _EQUITY_RESEARCH_TEMPLATE_VERSION,
        "notes": (
            "Full pipeline template: PDF extraction → 5 matplotlib charts → "
            "DCF + comps Excel workbook → python-docx report assembly. "
            "Outputs: result/report.docx + result/valuation.xlsx."
        ),
        "created_at": now_iso(),
    }

    if existing:
        # Update in-place
        idx = state["templates"].index(existing)
        record["created_at"] = existing.get("created_at", record["created_at"])
        state["templates"][idx] = record
    else:
        state["templates"].append(record)

    write_state(state)
    print(f"[startup] seeded equity research template {_EQUITY_RESEARCH_TEMPLATE_VERSION}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port", type=int, default=8765)
    args = parser.parse_args()
    set_settings_path(DATA_ROOT / "provider-settings.json")
    ensure_storage()
    _seed_builtin_equity_template()
    threading.Thread(target=_scheduler_loop, daemon=True, name="scheduler").start()
    server = ThreadingHTTPServer((args.host, args.port), RequestHandler)
    print(f"backend listening on http://{args.host}:{args.port}")
    server.serve_forever()


if __name__ == "__main__":
    main()
