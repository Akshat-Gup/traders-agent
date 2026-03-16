from __future__ import annotations

import json
import platform
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path
from typing import Any

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

try:
    import pypdfium2 as pdfium
except Exception:  # pragma: no cover
    pdfium = None

ROOT = Path(__file__).resolve().parents[1]
SKILLS_DIR = ROOT / "skills"
OCR_SCRIPT = ROOT / "scripts" / "ocr_pages.swift"

FAMILY_SKILL_REFERENCES: dict[str, list[str]] = {
    "equity-research": [
        "equity-research/skills/initiating-coverage/references/task1-company-research.md",
        "equity-research/skills/initiating-coverage/references/task2-financial-modeling.md",
        "equity-research/skills/initiating-coverage/references/task3-valuation.md",
        "equity-research/skills/initiating-coverage/references/task4-chart-generation.md",
        "equity-research/skills/initiating-coverage/references/task5-report-assembly.md",
        "equity-research/skills/initiating-coverage/references/valuation-methodologies.md",
        "equity-research/skills/initiating-coverage/assets/quality-checklist.md",
        "equity-research/skills/initiating-coverage/assets/report-template.md",
        "equity-research/skills/earnings-analysis/references/report-structure.md",
        "equity-research/skills/earnings-analysis/references/workflow.md",
        "equity-research/skills/earnings-analysis/references/best-practices.md",
    ],
}

STYLE_GUIDE = """# Research Style Guide

## Writing requirements
- Use a dry, direct tone.
- Do not use em dashes.
- Do not include meta-commentary about the writing process.
- Prefer bullet points when they improve scanability, but every bullet must be a complete sentence.
- Define less common finance abbreviations on first use.
- Do not assume the audience knows internal sell-side shorthand.

## Depth targets
- Executive summary: 250 to 400 words plus 3 to 5 full-sentence bullets.
- Each major body section: at least 3 substantive paragraphs or an equivalent table plus commentary.
- Risk, catalyst, and valuation sections: include explicit assumptions, not just conclusions.
- Target useful depth per page rather than compressed summaries.

## Evidence requirements
- Ground every key claim in the staged source files.
- Pull quotes from extracted text when they materially strengthen the argument.
- When a statement comes from a source PDF, include the source file name and page number in working notes or captions.
- Use the rendered page images when tables, charts, or scanned content matter.

## Formatting pass
- Run a final pass for ASCII punctuation, consistent headings, consistent bullet indentation, and complete sentences.
- Remove marketing language, filler, and repeated thesis claims.
- If the report uses acronyms such as ARR, FCF, or TAM, define them unless the meaning is obvious from context.

## Example bullets
- Demand remained concentrated in custom AI accelerator programs, which supports near-term growth but increases customer concentration risk.
- The valuation still depends on execution against the higher-revenue scenario, so downside support is a function of cash flow durability rather than narrative momentum.
- Management commentary on renewals was constructive, but the source pack still leaves open questions on pricing durability beyond the initial contract cycle.
"""


def _copy_any(source: Path, destination: Path) -> None:
    if source.is_dir():
        shutil.copytree(source, destination, dirs_exist_ok=True)
        return
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, destination)


def _ensure_workspace_dirs(workspace_root: Path) -> dict[str, Path]:
    dirs = {
        "source_uploads": workspace_root / "source" / "uploads",
        "source_pdfs": workspace_root / "source" / "pdfs",
        "source_urls": workspace_root / "source" / "urls",
        "context": workspace_root / "context",
        "templates": workspace_root / "templates" / "selected",
        "extracted_text": workspace_root / "extracted" / "text",
        "extracted_images": workspace_root / "extracted" / "images",
        "extracted_metadata": workspace_root / "extracted" / "metadata",
        "generated_charts": workspace_root / "generated" / "charts",
        "generated_diagrams": workspace_root / "generated" / "diagrams",
        "generated_excel": workspace_root / "generated" / "excel",
        "result": workspace_root / "result",
        "logs": workspace_root / "logs",
    }
    for path in dirs.values():
        path.mkdir(parents=True, exist_ok=True)
    return dirs


def _iter_input_files(source: Path) -> list[tuple[Path, Path]]:
    if source.is_file():
        return [(source, Path(source.name))]
    files: list[tuple[Path, Path]] = []
    for file_path in sorted(source.rglob("*")):
        if file_path.is_file():
            files.append((file_path, file_path.relative_to(source.parent)))
    return files


def _extract_text_from_docx(path: Path) -> str:
    try:
        from docx import Document

        document = Document(str(path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    except Exception:
        return ""


def _extract_text_from_zip_xml(path: Path, folder_prefix: str) -> str:
    try:
        with zipfile.ZipFile(path) as archive:
            snippets: list[str] = []
            for member in archive.namelist():
                if member.endswith(".xml") and member.startswith(folder_prefix):
                    snippets.append(archive.read(member).decode("utf-8", errors="ignore"))
            return "\n".join(snippets[:20])
    except Exception:
        return ""


def _extract_text_like(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".csv"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        text = _extract_text_from_docx(path)
        return text or _extract_text_from_zip_xml(path, "word/")
    if suffix == ".pptx":
        return _extract_text_from_zip_xml(path, "ppt/")
    if suffix == ".xlsx":
        return _extract_text_from_zip_xml(path, "xl/")
    return ""


def _render_pdf_pages(source: Path, destination: Path) -> list[Path]:
    if pdfium is None:
        return []
    try:
        pdf = pdfium.PdfDocument(str(source))
        rendered: list[Path] = []
        for index in range(len(pdf)):
            page = pdf[index]
            output_path = destination / f"{source.stem}_page_{index + 1}.png"
            page.render(scale=2.1).to_pil().save(str(output_path))
            rendered.append(output_path)
        return rendered
    except Exception:
        return []


def _extract_pdf_page_text(source: Path) -> list[str]:
    if PdfReader is None:
        return []
    try:
        reader = PdfReader(str(source))
        return [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception:
        return []


def _needs_ocr(page_texts: list[str]) -> bool:
    text = "".join(page_texts).strip()
    if not text:
        return True
    alnum_count = sum(1 for char in text if char.isalnum())
    return alnum_count < 160


def _ocr_pages_with_swift(images: list[Path]) -> list[str]:
    if not images or platform.system() != "Darwin" or not OCR_SCRIPT.exists():
        return ["" for _ in images]
    cmd = ["xcrun", "swift", str(OCR_SCRIPT), *[str(path) for path in images]]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        return ["" for _ in images]
    try:
        payload = json.loads(result.stdout or "{}")
    except json.JSONDecodeError:
        return ["" for _ in images]
    text_by_path = {item["path"]: item.get("text", "") for item in payload.get("items", [])}
    return [text_by_path.get(str(path), "") for path in images]


def _prepare_pdf(source: Path, staged_pdf: Path, dirs: dict[str, Path]) -> dict[str, Any]:
    page_images = _render_pdf_pages(source, dirs["extracted_images"])
    direct_page_texts = _extract_pdf_page_text(source)
    text_source = "pdf-text"
    page_texts = direct_page_texts
    if _needs_ocr(direct_page_texts) and page_images:
        ocr_page_texts = _ocr_pages_with_swift(page_images)
        if any(text.strip() for text in ocr_page_texts):
            page_texts = ocr_page_texts
            text_source = "ocr"

    text_path = dirs["extracted_text"] / f"{source.stem}.txt"
    text_path.write_text("\n\n".join(text for text in page_texts if text.strip()), encoding="utf-8")

    page_manifest = {
        "source_pdf": str(staged_pdf),
        "text_source": text_source,
        "pages": [
            {
                "page": index + 1,
                "text": page_texts[index] if index < len(page_texts) else "",
                "image_path": str(page_images[index]) if index < len(page_images) else None,
            }
            for index in range(max(len(page_images), len(page_texts)))
        ],
    }
    page_manifest_path = dirs["extracted_metadata"] / f"{source.stem}.pages.json"
    page_manifest_path.write_text(json.dumps(page_manifest, indent=2), encoding="utf-8")

    return {
        "kind": "pdf",
        "text_path": str(text_path),
        "page_manifest_path": str(page_manifest_path),
        "page_images": [str(path) for path in page_images],
        "text_source": text_source,
    }


def _stage_template(template: dict[str, Any] | None, destination: Path) -> dict[str, Any] | None:
    if not template:
        return None
    source = Path(template["library_path"]).expanduser()
    if not source.exists():
        return None
    if source.is_dir():
        staged_path = destination / source.name
    else:
        staged_path = destination / source.name
    _copy_any(source, staged_path)
    return {
        "id": template.get("id"),
        "name": template.get("name"),
        "staged_path": str(staged_path),
        "output_formats": template.get("output_formats", []),
        "notes": template.get("notes", ""),
    }


def _stage_skill_references(family: str, destination: Path) -> list[str]:
    staged: list[str] = []
    for relative_path in FAMILY_SKILL_REFERENCES.get(family, []):
        source = SKILLS_DIR / relative_path
        if not source.exists():
            continue
        destination_path = destination / "skill-references" / relative_path
        destination_path.parent.mkdir(parents=True, exist_ok=True)
        destination_path.write_bytes(source.read_bytes())
        staged.append(str(destination_path.relative_to(destination.parents[1])))
    return staged


def _expected_outputs(output_format: str, valuation_required: bool) -> list[str]:
    outputs = [f"result/report.{output_format}"]
    if valuation_required:
        outputs.append("result/valuation.xlsx")
    return outputs


def _build_prompt(job: dict[str, Any], template_info: dict[str, Any] | None, staged_refs: list[str]) -> str:
    expected_outputs = _expected_outputs(job["output_format"], bool(job.get("valuation_required")))
    prompt_lines = [
        f"You are preparing a final {job['family']} deliverable in {job['output_format'].upper()} format.",
        f"Job kind: {job['kind']}",
        f"Objective: {job['objective']}",
        f"Workspace root: {job['workspace_path']}",
        f"Write finished deliverables only to {job['result_path']}.",
        "Read context/style-guide.md before writing or editing any report content.",
        "Use the extracted PDF text and page images as primary evidence. The source pack is already staged locally.",
        "Source layout:",
        "- source/uploads/ contains every uploaded file or folder copied into the workspace.",
        "- source/pdfs/ contains normalized PDF copies that Codex can read locally if needed.",
        "- extracted/text/ contains text extracted from each source file.",
        "- extracted/metadata/*.pages.json contains per-page text and page-image references for PDFs.",
        "- extracted/images/ contains rendered PDF pages for scanned documents, quotes, tables, and charts.",
        "- generated/charts/ and generated/diagrams/ are where you should save any custom visuals you create.",
        "- templates/selected/ contains the selected report template and staged equity research reference material.",
        "- result/ is the only location for the final deliverables that the app will open for the user.",
        "Execution requirements:",
        "- Edit the staged Python template when it helps you assemble the report faster or more cleanly.",
        "- Generate custom charts and diagrams as needed, save them as PNG files, and insert them into the report.",
        "- Back every major claim with the uploaded PDFs or other staged sources.",
        "- Use quotes sparingly but precisely, and keep them grounded in the extracted source text.",
        "- If a PDF is image-only, rely on the OCR text in extracted/text/ plus the rendered page images.",
        "- Run a final formatting pass to remove em dashes, filler, and meta-commentary.",
        f"Expected outputs: {', '.join(expected_outputs)}",
    ]
    if template_info:
        prompt_lines.extend(
            [
                f"Selected template: {template_info['name']}",
                f"Template path: {Path(template_info['staged_path']).relative_to(Path(job['workspace_path']))}",
                "The staged template is executable Python. You may edit it in-place and then run it from the workspace root.",
            ]
        )
    if job.get("valuation_required"):
        prompt_lines.extend(
            [
                "Valuation is enabled.",
                "- Produce result/valuation.xlsx alongside the report.",
                "- Keep the workbook source-backed and aligned with the written valuation section.",
            ]
        )
    if staged_refs:
        prompt_lines.append("Use the staged equity research references as working guidance:")
        prompt_lines.extend([f"- {ref}" for ref in staged_refs])
        prompt_lines.append(
            "Treat those reference files as methodology support. Do not stop the run to ask which sub-task to perform."
        )
    if job.get("custom_instructions"):
        prompt_lines.append(f"Additional instructions:\n{job['custom_instructions']}")
    return "\n\n".join(prompt_lines)


def _prepare_sources(job: dict[str, Any], dirs: dict[str, Path]) -> list[dict[str, Any]]:
    assets: list[dict[str, Any]] = []
    for raw_path in job.get("source_paths", []):
        source = Path(raw_path).expanduser()
        if not source.exists():
            continue
        for file_path, relative_path in _iter_input_files(source):
            staged_path = dirs["source_uploads"] / relative_path
            _copy_any(file_path, staged_path)
            asset: dict[str, Any] = {
                "source_path": str(file_path),
                "staged_path": str(staged_path),
                "relative_path": str(relative_path),
                "suffix": file_path.suffix.lower(),
            }
            if file_path.suffix.lower() == ".pdf":
                normalized_pdf = dirs["source_pdfs"] / file_path.name
                _copy_any(file_path, normalized_pdf)
                asset.update(_prepare_pdf(file_path, normalized_pdf, dirs))
                asset["normalized_pdf_path"] = str(normalized_pdf)
            else:
                extracted_text = _extract_text_like(file_path)
                if extracted_text.strip():
                    text_path = dirs["extracted_text"] / f"{file_path.stem}.txt"
                    text_path.write_text(extracted_text, encoding="utf-8")
                    asset["text_path"] = str(text_path)
                    asset["text_source"] = "file-text"
            assets.append(asset)
    if job.get("urls"):
        (dirs["source_urls"] / "urls.txt").write_text("\n".join(job["urls"]) + "\n", encoding="utf-8")
    return assets


def prepare_workspace(payload: dict[str, Any]) -> dict[str, Any]:
    job = payload["job"]
    template = payload.get("template")
    workspace_root = Path(job["workspace_path"]).expanduser()
    dirs = _ensure_workspace_dirs(workspace_root)
    assets = _prepare_sources(job, dirs)
    template_info = _stage_template(template, dirs["templates"])
    staged_refs = _stage_skill_references(job["family"], dirs["templates"])

    prompt = _build_prompt(job, template_info, staged_refs)
    (dirs["context"] / "objective.md").write_text(job["objective"], encoding="utf-8")
    (dirs["context"] / "answers.md").write_text("", encoding="utf-8")
    (dirs["context"] / "style-guide.md").write_text(STYLE_GUIDE, encoding="utf-8")
    (dirs["context"] / "prompt.md").write_text(prompt, encoding="utf-8")
    (dirs["context"] / "questions.md").write_text("", encoding="utf-8")
    (dirs["logs"] / "status.json").write_text(
        json.dumps({"status": job["status"], "updated_at": job["updated_at"]}, indent=2),
        encoding="utf-8",
    )
    (dirs["context"] / "run-config.json").write_text(
        json.dumps(
            {
                "family": job["family"],
                "output_format": job["output_format"],
                "valuation_required": bool(job.get("valuation_required")),
                "expected_outputs": _expected_outputs(job["output_format"], bool(job.get("valuation_required"))),
                "workspace_root": str(workspace_root),
                "result_path": str(dirs["result"]),
                "template_path": template_info["staged_path"] if template_info else None,
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    (dirs["extracted_metadata"] / "assets.json").write_text(json.dumps(assets, indent=2), encoding="utf-8")

    return {
        "prompt_preview": prompt,
        "assets_count": len(assets),
        "template": template_info,
        "staged_skill_references": staged_refs,
    }


def main() -> int:
    payload = json.loads(sys.stdin.read() or "{}")
    result = prepare_workspace(payload)
    sys.stdout.write(json.dumps(result))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
