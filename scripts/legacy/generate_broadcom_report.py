from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pypdfium2 as pdfium
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
TMP_DIR = ROOT / "tmp" / "docs"
SOURCE_IMG_DIR = ROOT / "tmp" / "pdfs" / "source_pages"
CROP_DIR = TMP_DIR / "crops"
CHART_DIR = TMP_DIR / "charts"
PDF_DIR = ROOT / "output" / "pdf"
DOC_DIR = ROOT / "output" / "doc"
PDF_PATH = PDF_DIR / "broadcom_equity_research_report.pdf"
DOCX_PATH = DOC_DIR / "broadcom_equity_research_report.docx"

TITLE = "Broadcom Inc. (NASDAQ: AVGO)"
SUBTITLE = "Comprehensive Earnings Update for Private Bankers"
REPORT_DATE = "14 March 2026"
PRICE_DATE = "04 March 2026"

CURRENT_PRICE = 317.53
TARGET_PRICE = 500.00
UPSIDE = TARGET_PRICE / CURRENT_PRICE - 1.0
MARKET_CAP_B = 1577.8
ENTERPRISE_VALUE_B = 1613.0
NET_DEBT_B = 51.9
DIVIDEND_YIELD = 0.007
DIVIDEND_Q = 0.65

NAVY = "#10324A"
BLUE = "#2F6EA0"
TEAL = "#4F96AE"
LIGHT = "#DBE8F3"
LIGHTER = "#EEF4F8"
GOLD = "#B68A3A"
RED = "#C45454"
GRAY = "#5B6670"
DARK = "#1B1F24"

Q1_ACTUALS = {
    "revenue": 19.3,
    "ai_revenue": 8.4,
    "semi": 12.5,
    "software": 6.8,
    "gross_margin": 77.0,
    "op_margin": 66.4,
    "ebitda_margin": 68.0,
    "eps": 2.05,
    "cfo": 8.3,
    "capex": 0.25,
    "buybacks": 7.9,
    "dividends": 3.0,
    "cash": 14.2,
}

STREET = {
    "revenue": 19.2,
    "ai_revenue": 8.2,
    "semi": 12.3,
    "software": 6.9,
    "gross_margin": 76.8,
    "op_margin": 65.5,
    "ebitda_margin_q2": 67.1,
    "eps": 2.03,
    "q2_revenue": 20.5,
    "q2_ai_revenue": 9.3,
}

Q2_GUIDE = {
    "revenue": 22.0,
    "ai_revenue": 10.7,
    "non_ai_semi": 4.1,
    "software": 7.2,
    "ebitda_margin": 68.0,
}

ANNUAL_MODEL = {
    "FY25A": {"semis": 36.9, "software": 27.0, "revenue": 63.9, "eps": 6.82, "fcf": 26.9},
    "FY26E": {"semis": 63.0, "software": 29.5, "revenue": 92.5, "eps": 12.80, "fcf": 34.5},
    "FY27E": {"semis": 87.0, "software": 34.0, "revenue": 121.0, "eps": 20.00, "fcf": 45.0},
}

BROKER_PTS = [
    ("HSBC", 535),
    ("J.P. Morgan", 500),
    ("Jefferies", 500),
    ("House view", 500),
    ("Goldman Sachs", 480),
    ("UBS", 475),
    ("Morgan Stanley", 470),
    ("BofA Securities", 450),
]

BROKER_EPS_27 = [
    ("UBS", 22.76),
    ("House view", 20.00),
    ("J.P. Morgan", 19.61),
    ("Morgan Stanley", 17.98),
    ("HSBC", 16.72),
]

FY27_AI_FRAMES = [
    ("Mgmt floor", 100),
    ("JPM", 120),
    ("UBS", 130),
    ("Jefferies upside", 200),
]

BROKER_GRID = [
    ["Broker", "PT", "Key message"],
    ["J.P. Morgan", "$500", "AI revenue >$65bn in FY26, ~$120bn in FY27 (mgmt floor >$100bn); OpenAI qualified; networking mix rising; top semi pick."],
    ["UBS", "$475", "Margin debate improved; >$130bn FY27 AI revenue plausible; semi business at 20x P/E, one turn above NVDA."],
    ["Morgan Stanley", "$470", "Visibility improved, rack-margin fears receding; COT not settled but constructive; risk/reward favors upside."],
    ["Goldman Sachs", "$480 (CL)", "Lowest inference cost thesis; $60bn/$130bn/$170bn AI revenue FY26/FY27/FY28; supply secured thru FY28."],
    ["Jefferies", "$500", "Revenue per GW rising to $20-25bn; >$100bn is a floor; networking outgrowing ASICs; post-P/E crash is cheap."],
    ["HSBC", "$535", "Highest PT. ASIC estimates 44-59% above consensus. Networking TAM expansion beyond hyperscalers."],
    ["BofA", "$450", "Most cautious. AVGO content ~$10-15bn/GW vs. NVDA $25-35bn. Google may diversify TPUv8 to MediaTek."],
]

ESTIMATE_MATRIX = [
    ["Metric", "FY25A", "FY26E", "FY27E"],
    ["Revenue ($bn)", "63.9", "92.5", "121.0"],
    ["Revenue growth", "23.9%", "44.8%", "30.8%"],
    ["Semiconductor revenue ($bn)", "36.9", "63.0", "87.0"],
    ["Software revenue ($bn)", "27.0", "29.5", "34.0"],
    ["AI revenue ($bn)", "19.0", "65.0", "110.0"],
    ["Adj. EBITDA margin", "67.6%", "68.2%", "68.0%"],
    ["Adj. EPS ($)", "6.82", "12.80", "20.00"],
    ["Free cash flow ($bn)", "26.9", "34.5", "45.0"],
]

SOURCE_MATRIX = [
    ["Source", "Use in this report"],
    ["J.P. Morgan full take (Harlan Sur, 5 Mar 2026)", "Primary factual base for quarter, AI revenue path, customer ramps, FY26-FY27 EPS framing, and inventory analysis."],
    ["UBS strong results + SOTP deep dive (Timothy Arcuri, 5 Mar & 23 Feb 2026)", "Margin interpretation, AI networking mix, SOTP valuation, software renewal risk, and higher-end FY27 AI revenue debate."],
    ["Morgan Stanley AI momentum builds (Joseph Moore, 5 Mar 2026)", "Risk/reward framework, COT debate framing, interconnect technology outlook, and scenario mapping."],
    ["Goldman Sachs first take / follow-up (Toshiya Hari, 4 Mar 2026)", "Lowest-inference-cost thesis, out-year AI revenue estimates, and Conviction List rationale."],
    ["Jefferies strong 2027 outlook (Blayne Curtis, 4 Mar 2026)", "Revenue-per-GW analysis, networking share upside case, and post-P/E-crash valuation argument."],
    ["HSBC potent mix of ASIC and networking (Frank Lee, 24 Nov 2025)", "High-end ASIC, CoWoS, networking TAM, and Marvell competitive comparison."],
    ["BofA Securities (5 Mar 2026)", "Content-per-GW competitive comparison (AVGO vs. NVDA vs. AMD), Google multi-vendor risk, and cautious counterpoint."],
    ["JPM Asian Tech key takeaways (5 Mar 2026)", "MediaTek Zebrafish / Humufish TPU programs and COT risk framing."],
]

THESIS_MATRIX = [
    ["Thesis pillar", "What supports it now", "What would weaken it"],
    ["AI platform breadth", "Six-customer framework, stronger networking mix, and multiple programs beyond Google.", "A material delay in Meta, OpenAI or Anthropic ramps."],
    ["Rising earnings power", "Broker estimates moved up, Q2 guide beat, and FY27 AI revenue floor is now explicit.", "Guide cuts or lower AI revenue conversion into EPS."],
    ["Cash flow support", "Software stability, large buybacks, and modest capex intensity relative to operating cash flow.", "Software churn or a sharp increase in capital intensity."],
]

DEBATE_MATRIX = [
    ["Debate", "Current broker read-through", "House take"],
    ["How big can FY27 AI revenue get?", "Mgmt >$100bn; JPM $120bn; UBS >$130bn; Jefferies sees upside far above that at $20-25bn/GW.", "Use $110bn base; leave upside for execution."],
    ["Will rack mix hurt margins?", "Concern eased materially. MS no longer models rack margin pressure. UBS sees more networking in Anthropic mix.", "Watch closely, but no longer a thesis-breaker."],
    ["How real is the COT / insourcing risk?", "JPM: 'multiple years' away. BofA: Google diversifying to MediaTek starting TPUv8. MS: 'not settled' but constructive.", "Networking and SerDes harder to insource than XPU logic. Medium-term risk, not near-term."],
    ["Does AVGO content/GW matter vs. NVDA?", "BofA: AVGO $10-15bn vs. NVDA $25-35bn/GW. UBS: 'similar range to AMD.' Jefferies: rising to $20-25bn.", "Lower content/GW offset by higher margins and dual monetization (XPU + networking)."],
    ["Is valuation full at ~25x FY27?", "Depends on whose FY27 EPS you use. UBS: semi business at 20x, one turn above NVDA.", "Not full if Broadcom delivers near AI and capital-return expectations."],
]

CLIENT_FIT = [
    ["Client-use appendix", "Practical interpretation"],
    ["Who this suits", "Growth-oriented clients who want AI exposure backed by cash flow and capital returns."],
    ["Who should be careful", "Clients needing low volatility or immediate certainty on hyperscaler spending trends."],
    ["Best holding period", "12-24 months so FY27 deployment visibility can develop."],
    ["What would make us more bullish", "Evidence that networking mix and customer breadth continue to expand without margin leakage."],
    ["What would make us less bullish", "AI demand pauses, software renewals deteriorate, or buybacks slow materially."],
]

MONITORING_FRAME = [
    ["What to monitor next", "Why it matters", "What would be positive"],
    ["AI revenue mix", "Shows whether networking continues to widen the moat and improve quality of growth.", "Networking remains close to or above 40% of AI revenue."],
    ["Customer breadth", "Stock will rerate more easily if investors believe AVGO is not a one-customer story.", "OpenAI, Meta, Anthropic commentary becomes more concrete."],
    ["COT / MediaTek progress", "Zebrafish TPU at Google is the most tangible COT risk cited by analysts (BofA, JPM Asian Tech).", "Zebrafish delayed or limited in scope; no new COT announcements."],
    ["Marvell competitive wins", "If MRVL wins a second major hyperscaler beyond Amazon, AVGO's ASIC moat weakens.", "MRVL remains Amazon-concentrated; no new hyperscaler wins."],
    ["Software renewal health", "VMware renewals are the main cushion under the valuation if AI spending volatility rises.", "Renewals stabilize and churn concerns prove manageable."],
    ["Capital returns", "Buybacks and dividends make the stock easier to own in diversified client portfolios.", "Repurchases remain aggressive without hurting leverage discipline."],
]

ROBOTICS_LINK = [
    ["Robotics AI layer", "Why it matters", "Broadcom exposure"],
    ["Compute infrastructure", "Large robotics foundation models need the same high-end clusters used for frontier LLM training.", "Direct via custom XPU / ASIC programs and related silicon IP."],
    ["High-speed networking", "Simulation, sensor fusion and video-heavy robot training move enormous amounts of data between chips and racks.", "Direct and significant via switching, interconnect and networking silicon."],
    ["Data layer", "Proprietary video, factory and sensor data are major moats, but difficult to access in public markets.", "Mostly indirect; Broadcom benefits when data volumes force more infrastructure spend."],
    ["Simulation / digital twins", "Robotics developers increasingly train in sim before deployment in the physical world.", "Indirect; more simulation means more compute and networking demand."],
]

# --- Competitive landscape data (BofA 5 Mar 2026, JPM Asian Tech, UBS, HSBC) ---
COMPETITIVE_MATRIX = [
    ["Dimension", "Broadcom (AVGO)", "NVIDIA (NVDA)", "AMD", "Marvell (MRVL)"],
    ["Silicon type", "Custom ASIC/XPU", "General-purpose GPU", "General-purpose GPU", "Custom ASIC"],
    ["Content per GW", "$10-15bn (BofA est.)", "$25-35bn (GPU+system+SW)", "$15-20bn (GPU+CPU+NIC)", "$8-12bn (est.)"],
    ["Key customers", "GOOGL, META, Anthropic, OpenAI +2", "All hyperscalers + enterprise", "MSFT, META, cloud", "AMZN (Trainium)"],
    ["Networking", "TH5/TH6, Jericho, SerDes (mkt leader)", "NVLink, Spectrum-X, InfiniBand", "Pensando DPU (limited)", "1.6T DSP (co-leader)"],
    ["Vertical integration", "Design→TSMC fab→packaging→SerDes", "Full stack incl. CUDA software", "GPU+CPU+NIC bundle", "Design→fab→packaging"],
    ["FY27 AI revenue est.", ">$100bn (mgmt); $120-130bn (Street)", "~$200bn+ (consensus)", "~$12-15bn (est.)", "~$8-10bn (est.)"],
    ["Key competitive risk", "COT / customer insourcing", "Customer diversification to ASIC", "Share loss to custom silicon", "AMZN concentration; COT"],
]

CONTENT_PER_GW = [
    ("NVDA\n(GPU+system+SW+CPU+net)", 25, 35),
    ("AMD\n(GPU+CPU+NIC)", 15, 20),
    ("AVGO\n(XPU silicon+net switches)", 10, 15),
    ("MRVL\n(ASIC silicon, est.)", 8, 12),
]

CUSTOMER_GW_FY27 = [
    ("Google (TPU v8)", 3.5),
    ("Anthropic", 3.0),
    ("Meta (Athena/Iris)", 1.5),
    ("OpenAI", 1.0),
    ("Customer 5", 0.5),
    ("Customer 6", 0.5),
]

ASIC_VS_GPU_ECONOMICS = [
    ["Factor", "Custom ASIC (Broadcom)", "General-purpose GPU (NVIDIA)"],
    ["Inference cost", "Lower (workload-optimized)", "Higher (general overhead)"],
    ["Design lead time", "18-24 months per gen", "Standard product cycle"],
    ["Switching cost", "Very high (multi-year roadmaps)", "Lower (commodity procurement)"],
    ["Software ecosystem", "Customer-managed", "CUDA ecosystem (deep moat)"],
    ["Flexibility", "Single workload optimized", "Multi-workload capable"],
    ["Total cost of ownership", "Lower at scale for inference", "Lower for mixed/dev workloads"],
]

COT_RISK_MATRIX = [
    ["Customer", "COT status", "Timeline risk", "Broadcom view (Q1 FY26 call)"],
    ["Google", "Internal teams exist; MediaTek collab on TPU Zebrafish (4Q26 ramp)", "Medium-term", "No meaningful COT share loss for 'many years'; technical barriers high"],
    ["Meta", "Early internal silicon efforts alongside Broadcom ASICs", "Long-term", "Engagements 'multi-year and strategic'; multiple product lines ramping"],
    ["OpenAI", "Newly engaged with AVGO; exploring multiple options", "Long-term", "Qualified as 6th XPU customer; 1+ GW FY27 deployment"],
    ["Anthropic", "Uses GOOGL infra; no known internal silicon effort", "Minimal", "Ramping to 3 GW FY27; rack assembly via ODM model"],
]


def ensure_dirs() -> None:
    for path in (TMP_DIR, SOURCE_IMG_DIR, CROP_DIR, CHART_DIR, PDF_DIR, DOC_DIR):
        path.mkdir(parents=True, exist_ok=True)


def style_matplotlib() -> None:
    plt.style.use("default")
    plt.rcParams.update(
        {
            "figure.facecolor": "white",
            "axes.facecolor": "white",
            "axes.edgecolor": "#C9D5DF",
            "axes.labelcolor": DARK,
            "text.color": DARK,
            "axes.titleweight": "bold",
            "font.size": 10,
            "font.family": "DejaVu Serif",
        }
    )


def save_barh_chart(path: Path, title: str, subtitle: str, labels: list[str], values: list[float], colors_list: list[str]) -> None:
    n = len(labels)
    fig_h = max(3.0, 0.52 * n + 1.2)
    fig, ax = plt.subplots(figsize=(7.0, fig_h))
    y = np.arange(n)
    ax.barh(y, values, color=colors_list[:n])
    ax.set_yticks(y, labels, fontsize=9)
    ax.invert_yaxis()
    ax.set_title(title, loc="left", fontsize=13, color=NAVY, fontweight="bold", pad=24)
    ax.text(0.0, 1.06, subtitle, transform=ax.transAxes, fontsize=8, color=GRAY)
    ax.grid(axis="x", color="#DCE4EA", linewidth=0.7)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for idx, value in enumerate(values):
        ax.text(value + max(values) * 0.01, idx, f"${value:.0f}", va="center", fontsize=9, color=NAVY)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_grouped_bar_chart(path: Path, title: str, subtitle: str, categories: list[str], left_vals: list[float], right_vals: list[float], left_label: str, right_label: str) -> None:
    x = np.arange(len(categories))
    width = 0.36
    fig, ax = plt.subplots(figsize=(7.0, 3.4))
    ax.bar(x - width / 2, left_vals, width, label=left_label, color=BLUE)
    ax.bar(x + width / 2, right_vals, width, label=right_label, color=TEAL)
    ax.set_xticks(x, categories, fontsize=9)
    ax.set_title(title, loc="left", fontsize=13, color=NAVY, fontweight="bold", pad=24)
    ax.text(0.0, 1.06, subtitle, transform=ax.transAxes, fontsize=8, color=GRAY)
    ax.grid(axis="y", color="#DCE4EA", linewidth=0.7)
    ax.legend(frameon=False, loc="upper left", fontsize=8.5)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for bars in ax.containers:
        ax.bar_label(bars, fmt="%.1f", padding=3, fontsize=8.5)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_stacked_bar_chart(path: Path, title: str, subtitle: str, categories: list[str], lower_vals: list[float], upper_vals: list[float], lower_label: str, upper_label: str) -> None:
    x = np.arange(len(categories))
    fig, ax = plt.subplots(figsize=(7.0, 3.4))
    ax.bar(x, lower_vals, label=lower_label, color=BLUE, width=0.55)
    ax.bar(x, upper_vals, bottom=lower_vals, label=upper_label, color=TEAL, width=0.55)
    ax.set_xticks(x, categories, fontsize=9)
    ax.set_title(title, loc="left", fontsize=13, color=NAVY, fontweight="bold", pad=24)
    ax.text(0.0, 1.06, subtitle, transform=ax.transAxes, fontsize=8, color=GRAY)
    ax.grid(axis="y", color="#DCE4EA", linewidth=0.7)
    ax.legend(frameon=False, loc="upper left", fontsize=8.5)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for idx, total in enumerate(np.array(lower_vals) + np.array(upper_vals)):
        ax.text(idx, total + max(total * 0.015, 0.12), f"{total:.1f}", ha="center", fontsize=8.5, color=NAVY)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_scenario_chart(path: Path, title: str, subtitle: str, labels: list[str], values: list[float], current_price: float) -> None:
    fig, ax = plt.subplots(figsize=(7.0, 3.4))
    colors_list = [RED, BLUE, GOLD]
    ax.bar(labels, values, color=colors_list, width=0.55)
    ax.axhline(current_price, color="#666666", linestyle="--", linewidth=1.2, label=f"Current ${current_price:.0f}")
    ax.set_title(title, loc="left", fontsize=13, color=NAVY, fontweight="bold", pad=24)
    ax.text(0.0, 1.06, subtitle, transform=ax.transAxes, fontsize=8, color=GRAY)
    ax.grid(axis="y", color="#DCE4EA", linewidth=0.7)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.legend(frameon=False, loc="upper left", fontsize=8.5)
    for idx, value in enumerate(values):
        ax.text(idx, value + 8, f"${value:.0f}", ha="center", fontsize=9, color=NAVY)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_content_per_gw_chart(path: Path) -> None:
    """Horizontal range bar chart: content per GW by vendor (BofA framework)."""
    fig, ax = plt.subplots(figsize=(7.0, 3.2))
    labels = [x[0] for x in CONTENT_PER_GW]
    lows = [x[1] for x in CONTENT_PER_GW]
    highs = [x[2] for x in CONTENT_PER_GW]
    y = np.arange(len(labels))
    bar_colors = ["#5B6670", "#5B6670", BLUE, TEAL]
    for i in range(len(labels)):
        ax.barh(y[i], highs[i] - lows[i], left=lows[i], height=0.55, color=bar_colors[i], alpha=0.85)
        ax.text(highs[i] + 0.5, y[i], f"${lows[i]}-{highs[i]}bn", va="center", fontsize=9, color=NAVY)
    ax.set_yticks(y, labels, fontsize=8.5)
    ax.invert_yaxis()
    ax.set_xlabel("Content per GW ($bn)", fontsize=9)
    ax.set_title("Revenue content per gigawatt of deployment", loc="left", fontsize=13, color=NAVY, fontweight="bold", pad=24)
    ax.text(0.0, 1.07, "AVGO captures less per GW than NVDA but at higher margins on silicon+networking (BofA, 5 Mar 2026)", transform=ax.transAxes, fontsize=8, color=GRAY)
    ax.grid(axis="x", color="#DCE4EA", linewidth=0.7)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.set_xlim(0, 42)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_customer_gw_chart(path: Path) -> None:
    """Horizontal bar chart: estimated FY27 GW allocation by customer."""
    fig, ax = plt.subplots(figsize=(7.0, 3.2))
    labels = [x[0] for x in CUSTOMER_GW_FY27]
    vals = [x[1] for x in CUSTOMER_GW_FY27]
    y = np.arange(len(labels))
    clist = [NAVY, BLUE, TEAL, GOLD, GRAY, GRAY]
    ax.barh(y, vals, color=clist, height=0.6)
    ax.set_yticks(y, labels, fontsize=8.5)
    ax.invert_yaxis()
    ax.set_xlabel("Gigawatts of compute capacity", fontsize=9)
    ax.set_title("Estimated FY27 compute deployment by XPU customer", loc="left", fontsize=13, color=NAVY, fontweight="bold", pad=24)
    ax.text(0.0, 1.07, f"Total ~{sum(vals):.1f} GW across six customers; mgmt guided 'close to 10 GW' (JPM, 5 Mar 2026)", transform=ax.transAxes, fontsize=8, color=GRAY)
    ax.grid(axis="x", color="#DCE4EA", linewidth=0.7)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for i, v in enumerate(vals):
        ax.text(v + 0.05, i, f"{v:.1f} GW", va="center", fontsize=9, color=NAVY)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_networking_mix_chart(path: Path) -> None:
    """Bar chart showing networking % of AI revenue rising over time."""
    fig, ax = plt.subplots(figsize=(7.0, 3.2))
    quarters = ["Q3 FY25", "Q4 FY25", "Q1 FY26", "Q2 FY26G"]
    net_pct = [25, 30, 33, 40]
    asic_pct = [75, 70, 67, 60]
    x = np.arange(len(quarters))
    ax.bar(x, asic_pct, label="XPU / ASIC compute", color=BLUE, width=0.55)
    ax.bar(x, net_pct, bottom=asic_pct, label="AI networking", color=GOLD, width=0.55)
    ax.set_xticks(x, quarters, fontsize=9)
    ax.set_ylabel("% of AI semiconductor revenue", fontsize=9)
    ax.set_title("AI revenue mix: networking share rising", loc="left", fontsize=13, color=NAVY, fontweight="bold", pad=24)
    ax.text(0.0, 1.07, "Networking at ~40% of AI revenue in Q2 guide vs. ~25% in Q3 FY25 (MS, Jefferies, 5 Mar 2026)", transform=ax.transAxes, fontsize=8, color=GRAY)
    ax.legend(frameon=False, loc="lower left", fontsize=8.5)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for i, n in enumerate(net_pct):
        ax.text(i, 100 + 1.5, f"{n}%", ha="center", fontsize=9, color=NAVY, fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_capital_chart(path: Path) -> None:
    labels = ["CFO", "Capex", "Buybacks", "Dividends"]
    values = [Q1_ACTUALS["cfo"], Q1_ACTUALS["capex"], Q1_ACTUALS["buybacks"], Q1_ACTUALS["dividends"]]
    fig, ax = plt.subplots(figsize=(7.0, 3.2))
    ax.bar(labels, values, color=[BLUE, TEAL, GOLD, RED], width=0.55)
    ax.set_title("Q1 FY26 cash generation and capital return", loc="left", fontsize=13, color=NAVY, fontweight="bold", pad=24)
    ax.text(0.0, 1.06, "Cash from operations easily covered repurchases and dividends", transform=ax.transAxes, fontsize=8, color=GRAY)
    ax.grid(axis="y", color="#DCE4EA", linewidth=0.7)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for idx, value in enumerate(values):
        ax.text(idx, value + 0.22, f"${value:.2f}bn" if value < 1 else f"${value:.1f}bn", ha="center", fontsize=9, color=NAVY)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def render_pdf_page(pdf_path: Path, page_no: int, out_path: Path) -> None:
    pdf = pdfium.PdfDocument(str(pdf_path))
    pdf[page_no - 1].render(scale=2.2).to_pil().save(out_path)


def ensure_source_page(name: str, pdf_rel: str, page_no: int) -> Path:
    out_path = SOURCE_IMG_DIR / name
    if not out_path.exists():
        render_pdf_page(ROOT / pdf_rel, page_no, out_path)
    return out_path


def crop_source_charts() -> None:
    jpm_page = ensure_source_page("jpm_page2.png", "report/jp-broadcom-fulltake.pdf", 2)
    ms_page = ensure_source_page("ms_page4.png", "report/VisionAlpha_Report_1773474642089.pdf", 4)
    ubs_page = ensure_source_page("ubs_p4.png", "report/VisionAlpha_Report_1773474623069.pdf", 4)
    hsbc_page = ensure_source_page("hsbc_p4.png", "report/VisionAlpha_Report_1773474758155.pdf", 4)
    jeff_page = ensure_source_page("jeff_va_p4.png", "report/VisionAlpha_Report_1773474668082.pdf", 4)
    jpm_crop = CROP_DIR / "jpm_crop_metrics.png"
    ms_crop = CROP_DIR / "ms_crop_risk_reward.png"
    ubs_crop = CROP_DIR / "ubs_crop_tpu.png"
    hsbc_crop = CROP_DIR / "hsbc_crop_revisions.png"
    jeff_crop = CROP_DIR / "jeff_crop_charts.png"
    if not jpm_crop.exists():
        img = Image.open(jpm_page)
        img.crop((120, 380, 1480, 2600)).save(jpm_crop)
    if not ms_crop.exists():
        img = Image.open(ms_page)
        img.crop((120, 620, 1940, 2850)).save(ms_crop)
    if not ubs_crop.exists():
        img = Image.open(ubs_page)
        img.crop((90, 120, 2340, 2380)).save(ubs_crop)
    if not hsbc_crop.exists():
        img = Image.open(hsbc_page)
        img.crop((90, 220, 2330, 3180)).save(hsbc_crop)
    if not jeff_crop.exists():
        img = Image.open(jeff_page)
        img.crop((80, 230, 2290, 2540)).save(jeff_crop)


def generate_charts() -> dict[str, Path]:
    style_matplotlib()
    charts = {
        "targets": CHART_DIR / "price_targets.png",
        "rev_ai": CHART_DIR / "revenue_ai_progression.png",
        "q1_mix": CHART_DIR / "q1_revenue_mix.png",
        "ai_mix": CHART_DIR / "ai_mix_q1_q2.png",
        "ai_frames": CHART_DIR / "fy27_ai_frameworks.png",
        "cash": CHART_DIR / "cash_returns.png",
        "annual_mix": CHART_DIR / "annual_mix.png",
        "scenario": CHART_DIR / "valuation_scenarios.png",
        "broker_eps": CHART_DIR / "broker_eps_fy27.png",
        "content_gw": CHART_DIR / "content_per_gw.png",
        "customer_gw": CHART_DIR / "customer_gw_fy27.png",
        "net_mix": CHART_DIR / "networking_mix_progression.png",
    }
    save_barh_chart(
        charts["targets"],
        "Price targets sampled from sell-side coverage",
        "Broadcom remains broadly rated bullish after the March quarter",
        [x[0] for x in BROKER_PTS],
        [x[1] for x in BROKER_PTS],
        [GOLD, BLUE, BLUE, NAVY, TEAL, TEAL, LIGHT, "#B8CBD9"],
    )
    save_grouped_bar_chart(
        charts["rev_ai"],
        "Revenue and AI revenue are stepping up together",
        "Q1 FY26 actuals and Q2 FY26 guidance versus prior-year quarter",
        ["Q1 FY25A", "Q1 FY26A", "Q2 FY26G"],
        [14.9, 19.3, 22.0],
        [4.1, 8.4, 10.7],
        "Total revenue",
        "AI revenue",
    )
    save_stacked_bar_chart(
        charts["q1_mix"],
        "Q1 FY26 revenue mix",
        "AI is already the largest piece of the Broadcom story",
        ["Q1 FY26"],
        [5.63 + 2.77, ][:1],
        [4.12 + 6.80, ][:1],
        "AI revenue",
        "Non-AI semi + software",
    )
    save_stacked_bar_chart(
        charts["ai_mix"],
        "AI mix is shifting toward networking",
        "One-third of AI revenue in Q1; management says 40% in Q2",
        ["Q1 FY26A", "Q2 FY26G"],
        [5.63, 6.42],
        [2.77, 4.28],
        "AI compute / ASIC",
        "AI networking",
    )
    save_barh_chart(
        charts["ai_frames"],
        "FY27 AI revenue frameworks across sell-side coverage",
        "Management floor sits materially below the most bullish broker views",
        [x[0] for x in FY27_AI_FRAMES],
        [x[1] for x in FY27_AI_FRAMES],
        [BLUE, TEAL, GOLD, NAVY],
    )
    save_capital_chart(charts["cash"])
    save_stacked_bar_chart(
        charts["annual_mix"],
        "Our revenue mix view: AI semis drive growth, software anchors cash flow",
        "Amounts in $bn; semis include both AI and non-AI semiconductor solutions",
        list(ANNUAL_MODEL.keys()),
        [ANNUAL_MODEL[k]["semis"] for k in ANNUAL_MODEL],
        [ANNUAL_MODEL[k]["software"] for k in ANNUAL_MODEL],
        "Semiconductor solutions",
        "Infrastructure software",
    )
    save_scenario_chart(
        charts["scenario"],
        "12-month valuation scenarios",
        "Base case uses 25x FY27E adjusted EPS of $20.00",
        ["Bear", "Base", "Bull"],
        [360, 500, 575],
        CURRENT_PRICE,
    )
    save_barh_chart(
        charts["broker_eps"],
        "FY27 EPS framing across key brokers",
        "Wide range reflects disagreement on how fast AI revenue converts into earnings",
        [x[0] for x in BROKER_EPS_27],
        [x[1] for x in BROKER_EPS_27],
        [GOLD, NAVY, BLUE, TEAL, LIGHT],
    )
    save_content_per_gw_chart(charts["content_gw"])
    save_customer_gw_chart(charts["customer_gw"])
    save_networking_mix_chart(charts["net_mix"])
    return charts


def header_footer(canvas, doc) -> None:
    canvas.saveState()
    w, h = letter
    canvas.setStrokeColor(colors.HexColor(LIGHT))
    canvas.line(doc.leftMargin, h - 0.42 * inch, w - doc.rightMargin, h - 0.42 * inch)
    canvas.setFont("Helvetica-Bold", 8.5)
    canvas.setFillColor(colors.HexColor(NAVY))
    canvas.drawString(doc.leftMargin, h - 0.34 * inch, "Broadcom Inc. - Equity Research Update")
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor(GRAY))
    canvas.drawRightString(w - doc.rightMargin, h - 0.34 * inch, REPORT_DATE)
    canvas.line(doc.leftMargin, 0.44 * inch, w - doc.rightMargin, 0.44 * inch)
    canvas.drawString(doc.leftMargin, 0.30 * inch, "Prepared from broker research pack and company filings.")
    canvas.drawRightString(w - doc.rightMargin, 0.30 * inch, f"Page {canvas.getPageNumber()}")
    canvas.restoreState()


def make_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="Body",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=7.6,
            leading=8.8,
            textColor=colors.HexColor(DARK),
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Tiny",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=6.8,
            leading=7.8,
            textColor=colors.HexColor(DARK),
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ERBullet",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=7.5,
            leading=8.6,
            leftIndent=12,
            firstLineIndent=-8,
            textColor=colors.HexColor(DARK),
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TitleLarge",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=17,
            leading=19,
            textColor=colors.HexColor(NAVY),
            alignment=TA_LEFT,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Subtitle",
            parent=styles["Heading2"],
            fontName="Helvetica",
            fontSize=9.0,
            leading=10.5,
            textColor=colors.HexColor(GRAY),
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Section",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12.0,
            leading=13.2,
            textColor=colors.HexColor(NAVY),
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CenterNote",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=7.0,
            leading=8.0,
            textColor=colors.HexColor(GRAY),
            alignment=TA_CENTER,
            spaceAfter=2,
        )
    )
    return styles


def info_table(rows: list[list[str]], col_widths: list[float], highlight_col: int | None = None) -> Table:
    table = Table(rows, colWidths=col_widths, hAlign="LEFT")
    style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(NAVY)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 7.4),
        ("TOPPADDING", (0, 0), (-1, 0), 4),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 4),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(LIGHTER)]),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 1), (-1, -1), 7.2),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C7D4DE")),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 1), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 3),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]
    if highlight_col is not None:
        style.extend(
            [
                ("BACKGROUND", (highlight_col, 1), (highlight_col, -1), colors.HexColor(LIGHT)),
                ("FONTNAME", (highlight_col, 1), (highlight_col, -1), "Helvetica-Bold"),
            ]
        )
    table.setStyle(TableStyle(style))
    return table


def image_flow(path: Path, width: float, height: float, caption: str | None = None, styles=None):
    img = RLImage(str(path), width=width, height=height)
    if caption:
        return [img, Paragraph(caption, styles["CenterNote"])]
    return [img]


def side_by_side(left_flowables: list, right_flowables: list, widths: tuple[float, float] = (3.45 * inch, 3.45 * inch)) -> Table:
    tbl = Table([[left_flowables, right_flowables]], colWidths=list(widths), hAlign="LEFT")
    tbl.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    return tbl


def make_pdf(charts: dict[str, Path]) -> None:
    styles = make_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.44 * inch,
        rightMargin=0.44 * inch,
        topMargin=0.56 * inch,
        bottomMargin=0.52 * inch,
    )
    story = []

    def add_title(title: str, subtitle: str) -> None:
        story.append(Paragraph(title, styles["TitleLarge"]))
        story.append(Paragraph(subtitle, styles["Subtitle"]))

    add_title(TITLE, f"Q1 FY26 Earnings Update | Rating: Outperform | PT: ${TARGET_PRICE:.0f} | Price ({PRICE_DATE}): ${CURRENT_PRICE:.2f}")
    rating_table = info_table(
        [
            ["Rating box", "Value"],
            ["Recommendation", "Outperform"],
            ["Current price", f"${CURRENT_PRICE:.2f}"],
            ["Target price", f"${TARGET_PRICE:.0f}"],
            ["Implied upside", f"{UPSIDE:.1%}"],
            ["Market cap", f"${MARKET_CAP_B:,.0f}bn"],
            ["Enterprise value", f"${ENTERPRISE_VALUE_B:,.0f}bn"],
            ["Quarterly dividend", f"${DIVIDEND_Q:.2f}"],
        ],
        [1.6 * inch, 1.45 * inch],
        highlight_col=1,
    )
    company_snapshot = [
        Paragraph("<b>Company overview.</b> Broadcom is a $1.6tn semiconductor and infrastructure software company. The semiconductor segment designs custom AI accelerator chips (XPUs/ASICs) for hyperscale customers and sells high-speed networking silicon (switches, SerDes, DSPs) that connects those chips inside data-center clusters. The infrastructure software segment, centered on VMware, provides enterprise infrastructure with 93% gross margins and ~$27bn annual recurring revenue.", styles["Body"]),
        Paragraph("<b>Q1 FY26 results: beat and raise.</b> Revenue $19.3bn (+0.5% vs. Street), AI revenue $8.4bn (+2.4%), adj. EPS $2.05 (+1.0%). Q2 guided to $22.0bn revenue (+7.3% vs. Street) and $10.7bn AI revenue (+15.1%). Operating margin 66.4% (+90bps vs. consensus). $8.3bn operating cash flow, $7.9bn buybacks, $10bn new repurchase authorization. JPM (Harlan Sur, 5 Mar) raised PT to $500; Goldman (Toshiya Hari, 4 Mar) added to Conviction List.", styles["Body"]),
        Paragraph("<b>Key change after the print.</b> Management disclosed >$100bn FY27 AI semiconductor revenue across ~10 GW of compute for six XPU customers. All critical components secured through FY28. Rack-margin dilution concerns substantially eased. Networking rising from ~33% to ~40% of AI revenue.", styles["Body"]),
    ]
    story.append(side_by_side([rating_table], company_snapshot))
    story.append(Spacer(1, 0.04 * inch))
    summary_table = info_table(
        [
            ["Q1 FY26 dashboard", "Actual", "Street", "Variance"],
            ["Revenue", "$19.3bn", "$19.2bn", "+0.5%"],
            ["AI revenue", "$8.4bn", "$8.2bn", "+2.4%"],
            ["Gross margin", "77.0%", "76.8%", "+20bps"],
            ["Operating margin", "66.4%", "65.5%", "+90bps"],
            ["EBITDA margin", "68.0%", "66.3%", "+170bps"],
            ["Adj. EPS", "$2.05", "$2.03", "+1.0%"],
            ["Q2 revenue guide", "$22.0bn", "$20.5bn", "+7.3%"],
            ["Q2 AI revenue guide", "$10.7bn", "$9.3bn", "+15.1%"],
        ],
        [1.8 * inch, 1.0 * inch, 1.0 * inch, 0.9 * inch],
        highlight_col=1,
    )
    summary_bullets = [
        Paragraph("■ <b>Guidance, not the quarter, was the inflection.</b> Q1 beat modestly, but Q2 AI revenue guide of $10.7bn was 15% above consensus. This shifts the debate from quarterly optionality to multi-year deployment visibility. JPM (Harlan Sur, 5 Mar) estimates >$65bn FY26 AI revenue, on track for >$100bn in FY27. Goldman added AVGO to the Conviction List.", styles["ERBullet"]),
        Paragraph("■ <b>Margin fears defused.</b> Gross margin 77.0% (vs. 76.8% Street), operating margin 66.4% (vs. 65.5%). Management explicit that rack shipments will not dilute margins. UBS (Timothy Arcuri, 5 Mar) notes Anthropic mix may include more high-margin networking plus ODM rack assembly. MS (Joseph Moore, 5 Mar) no longer models incremental margin pressure from racks.", styles["ERBullet"]),
        Paragraph("■ <b>Customer base broadening across six XPU programs.</b> Google (anchor, TPU v8 Sunfish), Anthropic (~3 GW FY27), OpenAI (newly qualified, 1+ GW), Meta (Athena/Iris), plus two unnamed (likely SoftBank/Arm, ByteDance per Goldman and Jefferies). This directly addresses the concentration bear case that has weighed on the multiple.", styles["ERBullet"]),
        Paragraph("■ <b>Supply chain secured through FY28.</b> Leading-edge wafers, HBM, substrates, and advanced packaging locked in. Removes risk that supplier margin demands erode Broadcom's profitability. Chip inventories at $2.96bn (vs. $2.64bn consensus, up ~30% QoQ) signal demand confidence per JPM.", styles["ERBullet"]),
    ]
    story.append(side_by_side([summary_table], image_flow(charts["targets"], 3.35 * inch, 2.1 * inch, "Figure 1. Sell-side price target comparison.", styles)))
    story.append(Spacer(1, 0.03 * inch))
    story.extend(summary_bullets)
    story.append(PageBreak())

    add_title("Latest Earnings Update", "Q1 FY26 (January quarter) | Beat and raise across all key metrics")
    beat_table = info_table(
        [
            ["Q1 FY26 detail", "Actual", "Street est.", "Beat / miss", "Commentary"],
            ["Total revenue", "$19.3bn", "$19.2bn", "+0.5%", "Beat driven by semi solutions"],
            ["AI revenue", "$8.4bn", "$8.2bn", "+2.4%", "XPU + networking above consensus"],
            ["Semi solutions", "$12.5bn", "$12.3bn", "+1.6%", "AI offset soft legacy markets"],
            ["Infra software", "$6.8bn", "$6.9bn", "-1.4%", "Stable, slightly light on renewals"],
            ["Gross margin", "77.0%", "76.8%", "+20bps", "Above despite rack shipment fears"],
            ["Operating margin", "66.4%", "65.5%", "+90bps", "Better cost flow-through on scale"],
            ["EBITDA margin", "68.0%", "66.3%", "+170bps", "Best margin leverage in 4 quarters"],
            ["Adj. EPS", "$2.05", "$2.03", "+1.0%", "Quality beat: op leverage > buybacks"],
            ["Op. cash flow", "$8.3bn", "n/a", "n/a", "Capex only $250m; FCF ~$8.1bn"],
        ],
        [1.35 * inch, 0.82 * inch, 0.82 * inch, 0.72 * inch, 3.15 * inch],
    )
    story.append(beat_table)
    story.append(Spacer(1, 0.03 * inch))
    story.append(
        side_by_side(
            image_flow(charts["rev_ai"], 3.35 * inch, 1.85 * inch, "Figure 2. Revenue and AI revenue progression, FY25-FY27E.", styles),
            image_flow(charts["q1_mix"], 3.35 * inch, 1.85 * inch, "Figure 3. Q1 FY26 revenue composition.", styles),
        )
    )
    for line in [
        "■ <b>Guidance was the inflection, not the beat.</b> Q2 AI revenue guided to $10.7bn, 15.1% above consensus of $9.3bn. Total revenue guided to $22.0bn vs. $20.5bn Street. Goldman (Toshiya Hari, 4 Mar) notes this 'is well above GS at $9.3bn and the Street at $9.3bn,' and expects the stock to 'trade meaningfully higher.' This moved the narrative from quarterly optionality to multi-year deployment visibility.",
        "■ <b>Inventory buildup is demand-pull, not concern.</b> Chip inventories reached $2.96bn vs. $2.64bn consensus, up ~30% QoQ. JPM (Harlan Sur, 5 Mar) interprets this as bullish at the start of a fiscal year, signaling Broadcom expects strong forward demand. Historical inventory builds have preceded revenue acceleration in prior AI ramp cycles.",
        "■ <b>Networking within AI is inflecting.</b> Networking rose from ~25% in Q3 FY25 to ~33% in Q1 FY26 and is guided to ~40% in Q2 FY26. This matters because networking carries higher gross margins than XPU silicon and is less customer-concentrated. Jefferies (Blayne Curtis, 4 Mar) notes Broadcom's Tomahawk 6 and 1.6T DSP leadership are 'furthering market share gains' in networking, with Tomahawk 7 (200T capacity) ramping next year.",
        "■ <b>Supply chain secured through FY28.</b> Management confirmed that leading-edge wafers, HBM, substrates, and CoWoS advanced packaging are locked in through FY28. Goldman identifies this as 'a significant positive' that 'addresses a key investor concern regarding potential gross margin erosion.' HSBC (Frank Lee, 24 Nov) had earlier noted CoWoS allocation expanding from ~150k to ~250k wafers.",
        "■ <b>Capital returns remained aggressive.</b> $7.9bn in buybacks, $3.0bn in dividends, plus a new $10bn repurchase authorization. FCF conversion above 95%. JPM estimates buybacks could contribute ~$0.15-0.20 annual EPS accretion.",
    ]:
        story.append(Paragraph(line, styles["ERBullet"]))
    story.append(PageBreak())

    add_title("AI Engine, Customer Roadmap, And Competitive Position", "XPU/ASIC enablement + networking silicon + how Broadcom competes against NVIDIA, AMD, and Marvell")
    story.append(
        side_by_side(
            image_flow(charts["customer_gw"], 3.35 * inch, 1.75 * inch, "Figure 4. Estimated FY27 GW allocation by XPU customer (JPM, GS, 5 Mar 2026).", styles),
            image_flow(charts["net_mix"], 3.35 * inch, 1.75 * inch, "Figure 5. Networking share of AI revenue rising toward 40% (MS, Jefferies).", styles),
        )
    )
    story.append(Spacer(1, 0.02 * inch))
    story.append(Paragraph("<b>What Broadcom does in AI.</b> Broadcom's AI semiconductor business has two components. First, XPU/ASIC enablement: Broadcom designs custom accelerator chips for hyperscale customers. These are application-specific integrated circuits (ASICs) optimized for specific workloads rather than general-purpose GPUs. Google's TPU, Meta's Athena, and forthcoming OpenAI chips are all designed with Broadcom IP. Second, networking silicon: Broadcom sells the switch chips (Tomahawk 5/6, Jericho3-AI), serializer/deserializer (SerDes) IP, and optical DSPs that connect accelerators within clusters. CEO Hock Tan noted at an investor call that for every $1 spent on compute silicon, $0.15-0.20 is spent on networking.", styles["Body"]))
    for line in [
        "■ <b>Google (anchor customer).</b> Runs its entire TPU stack on Broadcom XPU designs. TPU v5e, v6 (Trillium), and v8 (Sunfish) all use Broadcom silicon and SerDes. UBS (Timothy Arcuri, 5 Mar) estimates ~3.7M TPU units shipping in CY2026 and ~6M in CY2027. Google accounted for the majority of Broadcom's AI revenue historically. BofA flags that Google may move from exclusive Broadcom sourcing to multi-vendor starting with TPUv8, with MediaTek collaborating on the Zebrafish program (ramping 4Q26) per JPM Asian Tech.",
        "■ <b>Anthropic (~3 GW FY27).</b> Uses Google Cloud infrastructure but requires dedicated Broadcom-designed silicon. JPM (Harlan Sur, 5 Mar) estimates Anthropic deployment triples YoY to 3 GW in FY27. Initial rack-margin dilution concerns eased after management clarified ODM assembly structure. UBS notes that the mix of Anthropic shipments may include more high-margin networking than initially feared.",
        "■ <b>Meta (Athena/Iris programs).</b> Developing custom training and inference ASICs with Broadcom. JPM and Goldman (Toshiya Hari, 4 Mar) both note Meta remains committed to the Broadcom ASIC roadmap despite concurrent NVIDIA purchases. Meta's scale (est. >$50bn CY26 capex) provides meaningful incremental volume for Broadcom's XPU pipeline.",
        "■ <b>OpenAI (newly qualified, 1+ GW).</b> Committed ~$10bn to develop custom chips with Broadcom per JPM. This is significant because OpenAI was previously perceived as entirely NVIDIA-dependent. Goldman notes this engagement reduces customer concentration risk.",
        "■ <b>Customers 5 and 6.</b> Not named by management; Goldman and Jefferies (Blayne Curtis, 4 Mar) suggest SoftBank/Arm and ByteDance. Expected to 'more than double' shipments in FY27.",
    ]:
        story.append(Paragraph(line, styles["ERBullet"]))
    story.append(Spacer(1, 0.02 * inch))
    story.append(Paragraph("<b>Competitive positioning: AVGO vs. NVDA, AMD, and MRVL.</b> Broadcom does not compete head-to-head with NVIDIA in selling GPUs. Instead, it provides an alternative path for hyperscalers that want custom silicon optimized for their specific workloads. BofA's 5 March note quantifies the difference: NVIDIA captures $25-35bn of content per gigawatt of deployment (GPU + system + software + CPU + end-to-end networking), AMD captures $15-20bn (GPU + CPU + NICs), and Broadcom captures $10-15bn (XPU silicon + networking switches only). Broadcom's lower content per GW reflects its narrower scope, but this is offset by higher gross margins on the silicon it does sell and by the expanding networking attach rate. Marvell (MRVL) is the closest ASIC competitor, primarily serving Amazon's Trainium program, with estimated $8-12bn content per GW. HSBC (Frank Lee, 24 Nov) notes Broadcom has matched Marvell's DSP roadmap with its 1.6T solution.", styles["Body"]))
    story.append(Paragraph("<b>Robotics as AI demand amplifier.</b> Training large robotics foundation models (world models, multimodal perception, simulation-heavy RL) requires the same dense accelerator clusters used for frontier LLMs but with disproportionately higher data-movement requirements. Simulation-based robot training generates large video and sensor-data volumes, increasing demand for high-bandwidth networking. This strengthens the infrastructure bottleneck thesis: as robotics workloads scale, value accrues to compute and networking layers before it reaches application software.", styles["Body"]))
    story.append(PageBreak())

    add_title("Guidance, Margins, And Capital Returns", "Q2 FY26 guide confirmed scaling AI revenue without margin sacrifice")
    guide_table = info_table(
        [
            ["Q2 FY26 guidance", "Management", "Prior Street", "Beat %", "Key read-through"],
            ["Revenue", "$22.0bn", "$20.5bn", "+7.3%", "Well above consensus; semis driving"],
            ["AI revenue", "$10.7bn", "$9.3bn", "+15.1%", "Both XPU and networking accelerating"],
            ["Non-AI semi", "$4.1bn", "~$4.0bn", "~flat", "No further deterioration in legacy"],
            ["Software", "$7.2bn", "~$7.0bn", "+2.9%", "Stable recurring cash engine"],
            ["EBITDA margin", "68.0%", "67.1%", "+90bps", "Rack margin fears eased materially"],
        ],
        [1.35 * inch, 0.92 * inch, 0.92 * inch, 0.65 * inch, 3.0 * inch],
    )
    story.append(guide_table)
    story.append(Spacer(1, 0.02 * inch))
    story.append(side_by_side(
        image_flow(charts["cash"], 3.35 * inch, 1.85 * inch, "Figure 6. Capital return composition, Q1 FY26.", styles),
        [
            Paragraph("<b>Margin structure.</b> Gross margin held at 77.0% vs. 76.8% consensus. Operating margin expanded to 66.4% vs. 65.5% Street. The key concern entering the print was that Anthropic rack shipments and pass-through content would dilute margins. Management stated these shipments use ODM assembly with Broadcom retaining high-margin silicon and networking content. MS (Joseph Moore, 5 Mar) notes it 'no longer models incremental margin pressure from racks.' UBS (Timothy Arcuri, 5 Mar) raised margin assumptions post-print, noting Anthropic mix 'may now include even more high-margin networking.'", styles["Body"]),
            Paragraph("<b>Supply chain locked through FY28.</b> Broadcom confirmed all critical components (leading-edge wafers, HBM, substrates, CoWoS advanced packaging) are secured through FY28. Goldman (4 Mar) calls this 'a significant positive as it addresses a key investor concern regarding potential gross margin erosion across the AI semiconductor space.' HSBC (Frank Lee, 24 Nov) notes CoWoS capacity expanding from ~150k to ~250k wafers. This removes the risk that TSMC, SK Hynix, or packaging suppliers capture margin on Broadcom's AI growth.", styles["Body"]),
        ],
    ))
    story.append(Spacer(1, 0.02 * inch))
    for line in [
        "■ <b>Cash generation remains exceptional.</b> Q1 operating cash flow $8.3bn on capex of only $250m, implying ~$8.1bn free cash flow. FCF margin approximately 42%. FCF/share annualized ~$7.70. This underpins the dividend ($12/yr annualized) and buyback program with significant headroom. UBS estimates FCF grows to >$45bn by FY27.",
        "■ <b>Capital return program expanded.</b> $7.9bn buybacks + $3.0bn dividends in Q1. New $10bn repurchase authorization. TTM shareholder returns exceed $20bn. Net debt/EBITDA declining toward 2.0x. JPM estimates buybacks could contribute ~$0.15-0.20 annual EPS accretion.",
        "■ <b>Software is the cash backstop, not the growth driver.</b> Infrastructure software at $6.8bn revenue, ~93% gross margin, ~$27bn ARR. VMware migration to VCF subscription model is largely complete. UBS's 23 Feb deep dive (Timothy Arcuri) flags renewal churn as the key software risk: 3-year contracts begin renewing FY27-FY28, and AI coding tools may push more workloads from on-prem to cloud, reducing VCF demand.",
        "■ <b>AI revenue trajectory accelerating.</b> From $4.0bn in Q3 FY25 to $8.4bn in Q1 FY26 to $10.7bn guided in Q2 FY26. Annualized run-rate exiting Q2 would be ~$43bn. JPM models >$65bn FY26 total AI revenue. Goldman projects $60bn/$130bn/$170bn FY26/FY27/FY28. Jefferies notes $20-25bn revenue per GW in CY27 implies potential above $200bn at 10 GW scale.",
    ]:
        story.append(Paragraph(line, styles["ERBullet"]))
    story.append(PageBreak())

    add_title("Core Investment Thesis", "Three pillars and one amplifier supporting an Outperform rating at $500 PT")
    for block in [
        "<b>Pillar 1: AI infrastructure platform with dual monetization.</b> Broadcom monetizes both custom compute silicon (XPUs/ASICs) and high-speed networking for every large AI cluster deployment. This dual exposure creates multiple revenue streams per customer and reduces dependence on any single chip generation. The company manages full ASIC design and validation, orchestrates fabrication with TSMC at 3nm/5nm nodes, handles advanced packaging (2.5D/3D stacking, CoWoS), and integrates proprietary SerDes IP for ultra-high-speed chip-to-chip communication. Unlike NVIDIA, which captures $25-35bn per GW through a full-stack approach (GPU + system + software + CPU + networking), Broadcom captures $10-15bn per GW (BofA, 5 Mar) but at higher margins on the silicon and networking it does sell. UBS's SOTP analysis shows the semiconductor business alone trading at 20x P/E, 23x EV/FCF, and 17x EV/EBITDA, 'only one turn more expensive than NVDA' despite the AI growth inflection.",
        "<b>Pillar 2: Earnings power rising faster than consensus reflects.</b> FY27 consensus EPS has moved from ~$17 to ~$20+ after the Q1 print. JPM (Harlan Sur, 5 Mar) models >$120bn FY27 AI revenue and $19.61 EPS. UBS (Timothy Arcuri) projects >$130bn AI revenue and $22.76 EPS. Jefferies (Blayne Curtis, 4 Mar) notes revenue per GW of $20-25bn in CY27 implies a $200bn theoretical upper bound and 'struggles to see how these compute names will remain this cheap.' Even on conservative assumptions (FY27E EPS $20.00, 25x multiple), the stock works to $500. The revision cycle is still early relative to FY27 deployment ramp, and Goldman projects AI revenue continuing to grow into FY28 ($170bn).",
        "<b>Pillar 3: Software and capital returns provide structural downside support.</b> The VMware-centered infrastructure software segment generates ~$27bn ARR at ~93% gross margins. Combined with >$30bn annual FCF capacity and aggressive buybacks ($10bn new authorization), the stock has a cash-flow floor that pure-play AI names lack. UBS's SOTP assigns ~$150-180/share to software alone. In a drawdown, software cash flows provide valuation support at roughly $250-300/share on standalone multiples. The risk, flagged by UBS on 23 Feb, is that 3-year VMware renewal cycles beginning FY27-FY28 may produce churn that erodes this floor.",
        "<b>Amplifier: Robotics and simulation workloads deepen the networking bottleneck thesis.</b> As robotics workloads scale (world models, simulation-heavy RL, sensor fusion), they stress exactly the infrastructure Broadcom monetizes: dense accelerator clusters, high-bandwidth networking, and data movement. Data center spending must front-run robotics monetization, so value accrues to the chip and networking layer first. For every $1 spent on chips, $0.15-0.20 is spent on networking (CEO Hock Tan). As simulation-based robot training generates enormous video and sensor-data volumes, the networking TAM expands disproportionately. This is incremental to the core LLM-driven demand thesis and specifically strengthens the case for Broadcom's networking business, which is already inflecting from ~25% to ~40% of AI revenue.",
    ]:
        story.append(Paragraph(block, styles["Body"]))
    story.append(Spacer(1, 0.03 * inch))
    story.append(info_table(THESIS_MATRIX, [1.45 * inch, 3.0 * inch, 2.45 * inch]))
    story.append(PageBreak())

    add_title("Competitive Landscape And Business Structure", "ASIC vs. GPU economics, COT risk, and the software valuation backstop")
    story.append(
        side_by_side(
            image_flow(charts["content_gw"], 3.35 * inch, 1.75 * inch, "Figure 7. Content per GW by vendor (BofA, UBS, Jefferies estimates).", styles),
            image_flow(charts["ai_frames"], 3.35 * inch, 1.75 * inch, "Figure 8. FY27 AI revenue frameworks across sell-side coverage.", styles),
        )
    )
    story.append(Spacer(1, 0.02 * inch))
    story.append(Paragraph("<b>Why hyperscalers are shifting toward custom ASICs.</b> NVIDIA's general-purpose GPUs dominated the initial phase of AI infrastructure buildout because the CUDA software ecosystem reduced time-to-deployment. As training and inference workloads mature, however, hyperscalers increasingly seek custom silicon that is optimized for their specific architectures. Custom ASICs designed by Broadcom (or Marvell) offer lower inference cost per token, lower power consumption per operation, and the elimination of NVIDIA's margin stack. Goldman Sachs (Toshiya Hari, 4 Mar) emphasizes that Broadcom's leadership in custom silicon 'enables the lowest inference cost for its hyperscaler customers' and that Broadcom is 'delivering ongoing cost reductions on pace with market leader NVIDIA.' This is the core of the ASIC bull case: as workloads scale into production inference at massive volumes, the economics increasingly favor purpose-built silicon over general-purpose GPUs.", styles["Body"]))
    story.append(info_table(ASIC_VS_GPU_ECONOMICS, [1.35 * inch, 2.75 * inch, 2.75 * inch]))
    story.append(Spacer(1, 0.02 * inch))
    story.append(Paragraph("<b>Customer-owned tooling (COT) risk: the key bear debate.</b> The most important competitive risk for Broadcom is that hyperscaler customers eventually internalize chip design. On the Q1 FY26 earnings call, management stated that COT is 'not an imminent risk' and that meaningful share loss is 'multiple years' away at the earliest. JPM's Asian Tech team (5 Mar) is more cautious, noting that MediaTek is collaborating with Google on the TPU Zebrafish project (slated to ramp in 4Q26) and a next-generation program (Humufish). BofA (5 Mar) flags that Google may move away from exclusivity starting with the FY27 TPUv8 generation. Morgan Stanley (Joseph Moore, 5 Mar) notes the debate 'is not settled' but that management commentary was constructive, emphasizing technical barriers in silicon design, high-speed SerDes, advanced packaging, and large-scale cluster networking. The key nuance: even if COT emerges over time, Broadcom's networking silicon and SerDes IP create revenue streams that are harder to insource than the XPU logic itself.", styles["Body"]))
    story.append(info_table(COT_RISK_MATRIX, [1.0 * inch, 2.2 * inch, 0.8 * inch, 2.85 * inch]))
    story.append(Spacer(1, 0.02 * inch))
    story.append(
        side_by_side(
            [
                Paragraph("<b>Infrastructure software (~35% of revenue).</b> Centered on VMware (acquired Nov 2023 for $69bn). ~$6.8bn quarterly revenue at ~93% gross margins. ~$27bn ARR. VMware has been migrated to a VCF subscription model. UBS (Timothy Arcuri, 23 Feb) values software at ~$150-180/share standalone. Risk: three-year enterprise contracts begin renewing FY27-FY28. AI coding tools may accelerate workload migration to cloud, reducing on-prem VMware demand. The software segment is not the growth driver but provides a cash-flow floor absent in pure-play AI names.", styles["Body"]),
            ],
            image_flow(charts["annual_mix"], 3.35 * inch, 1.75 * inch, "Figure 9. Annual revenue mix: semis growing share as AI scales.", styles),
        )
    )
    story.append(PageBreak())

    add_title("Valuation And Price Target", "PT $500 | 25x FY27E EPS of $20.00 | Base case in line with broker consensus midpoint")
    valuation_table = info_table(
        [
            ["Scenario", "FY27E EPS", "P/E", "Implied price", "Key assumptions"],
            ["Bear", "$15.00", "24x", "$360", "AI revenue ~$80bn FY27; software churn; capex pause; GM compresses to 74%"],
            ["Base", "$20.00", "25x", "$500", "AI revenue ~$110bn FY27; stable margins; software flat; 6 customers deploy"],
            ["Bull", "$23.00", "25x", "$575", "AI revenue >$130bn FY27; networking >40% mix; 7th customer; software re-rates"],
            ["Street high (HSBC)", "$22.50", "24x", "$535", "ASIC + networking double tailwind; highest PT on the Street"],
        ],
        [0.85 * inch, 0.85 * inch, 0.55 * inch, 0.85 * inch, 3.85 * inch],
        highlight_col=3,
    )
    story.append(valuation_table)
    story.append(Spacer(1, 0.02 * inch))
    story.append(
        side_by_side(
            image_flow(charts["scenario"], 3.35 * inch, 1.85 * inch, "Figure 10. Bear / base / bull valuation scenarios.", styles),
            image_flow(charts["broker_eps"], 3.35 * inch, 1.85 * inch, "Figure 11. FY27E EPS estimates across brokers.", styles),
        )
    )
    for line in [
        "■ <b>Methodology.</b> We use a forward P/E framework on FY27E adjusted EPS, consistent with JPM ($500 PT, ~25x per Harlan Sur), UBS ($475, ~21x C2027 per Timothy Arcuri), and Morgan Stanley ($470, ~24x per Joseph Moore). Goldman (Toshiya Hari, 4 Mar) uses 30x normalized EPS of $16, arriving at $480. A SOTP approach yields a similar range: UBS values the semiconductor business at 20x P/E and infrastructure software at 15-18x FCF, arriving at ~$475 blended.",
        "■ <b>Multiple justification.</b> 25x forward earnings sits at a ~10% premium to the semiconductor peer median but at a discount to pure-play AI beneficiaries trading at 30-40x. UBS notes AVGO's semi business is 'only one turn more expensive than NVDA,' which trades at ~19x P/E on CY27. The premium is justified by AI revenue duration (>$100bn FY27 with six customers), networking optionality, and software cash-flow support. The discount to high-growth AI names reflects the software segment's lower growth and customer-concentration overhang.",
        "■ <b>Broker range context.</b> The Street range spans $470 (MS) to $535 (HSBC). Our $500 base case sits near the midpoint. JPM uses >$120bn FY27 AI revenue; Goldman on CL is the most aggressive on out-year estimates ($130bn/$170bn FY27/FY28). Jefferies (Blayne Curtis, 4 Mar) argues the >$100bn floor is conservative and notes $20-25bn/GW in CY27 implies potential for $200bn at scale. Post the recent P/E reset (multiple compressed from >35x to ~25x on software weakness), Jefferies views AVGO as 'cheap' relative to the AI growth trajectory.",
        "■ <b>Why not higher.</b> The stock already carries significant AI expectations. Software multiple compression (VMware renewal risk flagged by UBS 23 Feb) could offset semiconductor upside. BofA (5 Mar) flags that AVGO content per GW is limited at $10-15bn vs. NVDA's $25-35bn, and that Google may diversify suppliers starting TPUv8. A 25x multiple requires continued evidence of customer breadth and margin stability. We prefer to let execution drive re-rating.",
    ]:
        story.append(Paragraph(line, styles["ERBullet"]))
    story.append(Spacer(1, 0.02 * inch))
    story.append(RLImage(str(CROP_DIR / "jeff_crop_charts.png"), width=6.8 * inch, height=1.85 * inch))
    story.append(Paragraph("Figure 12. Jefferies (Blayne Curtis, 4 Mar 2026): historical earnings surprise, forward P/E, margins, and short interest.", styles["CenterNote"]))
    story.append(PageBreak())

    add_title("Key Risks And Near-Term Catalysts", "Duration risk replaces discovery risk as the primary stock debate")
    risks_table = info_table(
        [
            ["Risk", "Description", "Severity", "What to monitor"],
            ["Hyperscaler capex normalization", "AI capex is being priced for multi-year duration. A pause or deceleration in cloud capex budgets compresses the revenue trajectory and re-rates the multiple. Jefferies (4 Mar) notes 'the AI spend overhang will still linger' and this call 'didn't necessarily answer' the sustainability question into C28.", "High", "Quarterly capex disclosures from GOOGL, META, AMZN, MSFT"],
            ["Customer insourcing (COT)", "Hyperscalers may internalize more chip design over time. BofA (5 Mar) flags Google moving to multi-vendor starting TPUv8 gen. JPM Asian Tech notes MediaTek on Zebrafish (4Q26) and Humufish programs. Even AVGO management acknowledges it won't have '100% share at any customer indefinitely' (MS, 5 Mar).", "Medium-High", "MediaTek ramp timeline, Google TPUv8 vendor split, design-win cadence"],
            ["Gross margin dilution", "Rack-level shipments and HBM pass-through content could lower blended margins as XPU mix rises. MS (5 Mar) still expects 'some longer-term mix pressure at the consolidated level as semis grow faster than software and as XPU mix rises (given pass-through elements like HBM).'", "Medium", "Networking % of AI revenue, incremental GM on new deployments"],
            ["Software renewal churn", "VMware 3-year contracts begin renewing FY27-FY28. UBS software deep dive (23 Feb) flags three risks: 1) customer churn as deals renew, 2) growth headwinds as VCF upsell is lapped, 3) AI coding tools pushing on-prem workloads to cloud.", "Medium", "VMware renewal rates, VCF net-new ARR, enterprise IT budgets"],
            ["Competitive risk from Marvell", "Marvell serves Amazon's Trainium program and has matched Broadcom's 1.6T DSP roadmap (HSBC, 24 Nov). If Marvell gains share with additional hyperscalers, Broadcom's ASIC revenue concentration worsens.", "Low-Med", "MRVL customer wins, DSP market share data"],
            ["Valuation risk", "At ~25x FY27E P/E, the stock needs consistent execution. Goldman (4 Mar) identifies 'slowdown in AI infrastructure spending' and 'share loss in custom compute franchise' as key downside risks.", "Medium", "Quarterly beat/miss pattern, forward guide trajectory"],
        ],
        [1.2 * inch, 2.55 * inch, 0.65 * inch, 2.45 * inch],
    )
    story.append(risks_table)
    story.append(Spacer(1, 0.03 * inch))
    catalyst_table = info_table(
        [
            ["Near-term catalyst", "Timeline", "Expected impact"],
            ["Q2 FY26 earnings", "Jun 2026", "Confirms $10.7bn AI revenue guide; validates networking mix shift to ~40%"],
            ["Customer deployment disclosures", "Ongoing", "OpenAI/Meta/Anthropic ramp evidence reduces concentration bear case"],
            ["Networking revenue inflection", "H2 FY26", "Tomahawk 5, Jericho3-AI, and SerDes attach rates drive margin-accretive growth"],
            ["Buyback execution", "Ongoing", "$10bn new auth; $0.15-0.20 annual EPS accretion potential (JPM est.)"],
            ["FY27 AI revenue framework update", "Q3/Q4 FY26", "Any upward revision to >$100bn floor would be a significant re-rating catalyst"],
            ["Robotics/simulation workload scaling", "CY26-27", "Incremental demand for networking bandwidth from video-heavy training workloads"],
        ],
        [1.8 * inch, 0.75 * inch, 4.35 * inch],
    )
    story.append(catalyst_table)
    story.append(Spacer(1, 0.02 * inch))
    story.append(side_by_side(
        image_flow(CROP_DIR / "ms_crop_risk_reward.png", 3.35 * inch, 2.4 * inch, "Figure 13. Morgan Stanley risk/reward framework.", styles),
        [
            Paragraph("<b>What would increase conviction:</b>", styles["Body"]),
            Paragraph("■ Q2 AI revenue at or above $10.7bn guide with networking >40% of mix", styles["ERBullet"]),
            Paragraph("■ OpenAI deployment confirmation beyond qualification stage", styles["ERBullet"]),
            Paragraph("■ No visible MediaTek/Zebrafish ramp in 4Q26 timeline", styles["ERBullet"]),
            Paragraph("■ Gross margin sustained above 76% despite rack shipment scaling", styles["ERBullet"]),
            Paragraph("<b>What would reduce conviction:</b>", styles["Body"]),
            Paragraph("■ Hyperscaler capex guidance cuts or deferrals in upcoming earnings", styles["ERBullet"]),
            Paragraph("■ Google announcing multi-vendor TPUv8 with meaningful MediaTek share", styles["ERBullet"]),
            Paragraph("■ VMware churn rates above 15% in early renewal cohorts", styles["ERBullet"]),
            Paragraph("■ Marvell winning a second major hyperscaler ASIC program beyond Amazon", styles["ERBullet"]),
        ],
    ))
    story.append(PageBreak())

    add_title("Street Read-Through / Broker Debate", "Consensus is constructive, but sell-side coverage reveals meaningful disagreement on magnitude and duration")
    story.append(info_table(BROKER_GRID, [1.35 * inch, 0.8 * inch, 4.85 * inch]))
    story.append(Spacer(1, 0.02 * inch))
    for line in [
        "■ <b>JPMorgan ($500 PT, OW; Harlan Sur, 5 Mar 2026).</b> Most detailed on customer framework. Identifies six XPU customers including OpenAI as newly qualified. Models >$120bn FY27 AI revenue and $19.61 FY27 EPS. Estimates $12-15bn revenue per GW 'conservatively.' Highlights inventory build ($2.96bn vs. $2.64bn consensus) as demand-pull signal. Component security through FY28 removes margin overhang. Asian Tech team separately notes MediaTek Zebrafish COT risk at Google.",
        "■ <b>UBS ($475 PT, Buy; Timothy Arcuri, 5 Mar 2026).</b> Strongest on margin analysis and SOTP. Values semi business at 20x P/E, 23x EV/FCF, 'only one turn more expensive than NVDA.' Estimates ~6M TPU in FY27 and revenue per GW 'in a similar range to AMD's ~$15B.' Models >$130bn FY27 AI rev, $22.76 EPS. 23 Feb deep dive flags VMware renewal churn risk. Argues >$100bn AI floor 'still seems very conservative.'",
        "■ <b>Morgan Stanley ($470 PT, OW; Joseph Moore, 5 Mar 2026).</b> Risk/reward focused. Uses ~$20bn/GW framework, $120bn FY27 AI rev as base. COT debate 'is not settled' but AVGO's SerDes, packaging, and networking advantages create barriers. Expects copper-based interconnect (200G/400G SerDes) before optical adoption, favoring Broadcom. No longer models incremental rack-margin pressure.",
        "■ <b>Goldman Sachs ($480 PT, Buy/CL; Toshiya Hari, 4 Mar 2026).</b> Most aggressive out-years: $60bn/$130bn/$170bn AI rev FY26/FY27/FY28. Core thesis: Broadcom 'enables the lowest inference cost for hyperscaler customers.' Uses 30x normalized $16 EPS. Key downside risks: AI spending slowdown, custom compute share loss, VMware competition.",
        "■ <b>Jefferies ($500, Franchise Pick; Blayne Curtis, 4 Mar 2026).</b> Most bullish on theoretical upside. Revenue per GW rising from $13bn CY25 to '$20-25bn/GW in 2027, implying potential for AI revenue to actually reach over $200bn.' Networking outgrowing ASICs. China 'likely a tailwind' for networking. 'Struggles to see how these compute names will remain this cheap.'",
        "■ <b>HSBC ($535, Buy; Frank Lee, 24 Nov 2025).</b> Highest PT. ASIC estimates 44-59% above VA consensus. Notes Broadcom matched Marvell's 1.6T DSP roadmap. CoWoS capacity 150k→250k wafers. Argues FY26 EPS may approach consensus FY27 EPS.",
        "■ <b>BofA ($450, Buy; 5 Mar 2026).</b> Most cautious competitive framing. Content per GW: AVGO ~$10-15bn vs. NVDA $25-35bn vs. AMD $15-20bn. Flags Google diversifying away from AVGO exclusivity starting TPUv8 (MediaTek). Notes 'revisions mostly FY27-28.' Important counterweight to the bull consensus.",
    ]:
        story.append(Paragraph(line, styles["ERBullet"]))
    story.append(Spacer(1, 0.02 * inch))
    story.append(info_table(DEBATE_MATRIX, [1.75 * inch, 2.55 * inch, 2.7 * inch]))
    story.append(
        side_by_side(
            image_flow(CROP_DIR / "ubs_crop_tpu.png", 3.35 * inch, 1.75 * inch, "Figure 14. UBS TPU unit estimates (5 Mar 2026).", styles),
            image_flow(CROP_DIR / "hsbc_crop_revisions.png", 3.35 * inch, 1.75 * inch, "Figure 15. HSBC consensus revision data (24 Nov 2025).", styles),
        )
    )
    story.append(PageBreak())

    add_title("Detailed Estimates And Portfolio Framework", "House estimates, position guidance, and monitoring framework")
    story.append(info_table(ESTIMATE_MATRIX, [2.0 * inch, 1.25 * inch, 1.25 * inch, 1.25 * inch], highlight_col=2))
    story.append(Spacer(1, 0.03 * inch))
    story.append(info_table(CLIENT_FIT, [1.8 * inch, 5.2 * inch]))
    story.append(Spacer(1, 0.02 * inch))
    story.append(info_table(MONITORING_FRAME, [1.55 * inch, 2.8 * inch, 2.55 * inch]))
    story.append(Spacer(1, 0.02 * inch))
    for line in [
        "■ <b>Position sizing guidance.</b> AVGO is appropriate as a 2-4% position in a diversified equity portfolio or 4-8% in a concentrated growth/tech allocation. The stock carries single-name risk from AI revenue concentration and customer-insourcing uncertainty (BofA flags Google TPUv8 diversification), but the software cash-flow base and aggressive capital return program provide structural support absent in NVDA or AMD.",
        "■ <b>Pair/hedge considerations.</b> Long AVGO can be paired with underweight positions in higher-multiple, lower-cash-flow AI names. For clients who want AI infrastructure exposure but are concerned about ASIC-specific COT risk, pairing AVGO with NVDA captures both the general-purpose GPU and custom silicon sides of the hyperscaler compute decision.",
        "■ <b>Rating-change triggers (upgrade to Strong Buy).</b> FY27 AI revenue visibility exceeding $130bn with stable margins above 76% gross. MediaTek Zebrafish delayed or descoped. Two or more new XPU customer deployments beyond current six. Software renewal rates demonstrating <10% churn.",
        "■ <b>Rating-change triggers (downgrade to Neutral).</b> AI revenue guidance miss of >10%. Gross margin declining below 74% on mix deterioration. Google publicly confirming meaningful MediaTek share in TPUv8. Marvell winning a second hyperscaler ASIC program. VMware churn exceeding 20%.",
    ]:
        story.append(Paragraph(line, styles["ERBullet"]))
    story.append(Spacer(1, 0.02 * inch))
    story.append(Paragraph("Broadcom occupies a structurally advantaged position in the AI infrastructure stack, monetizing both custom compute silicon and high-speed networking across a broadening hyperscaler customer base. The competitive landscape is not without risk: customer-owned tooling efforts (MediaTek Zebrafish at Google per BofA and JPM Asian Tech), Marvell's ambitions, and NVIDIA's full-stack dominance all constrain upside. But the Q1 FY26 print and Q2 guide confirm AI revenue is scaling without margin sacrifice, and supply security through FY28 provides a degree of earnings visibility that is rare in the semiconductor sector. We maintain Outperform with a $500 price target.", styles["Body"]))

    doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 8.0, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    if color is not None:
        run.font.color.rgb = color


def add_docx_table(document: Document, rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = Inches(width)
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.cell(r, c)
            if r == 0:
                set_cell_text(cell, value, bold=True, font_size=7.8, color=RGBColor(0xFF, 0xFF, 0xFF))
                shade_cell(cell, "10324A")
            else:
                set_cell_text(cell, value, bold=False, font_size=7.8)
                if r % 2 == 0:
                    shade_cell(cell, "EEF4F8")
    document.add_paragraph()


def set_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.56)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.44)
    section.right_margin = Inches(0.44)


def add_header(section, title: str) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    run = p.add_run(f"Broadcom Inc. - Equity Research Update | {title}")
    run.bold = True
    run.font.name = "Helvetica"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x10, 0x32, 0x4A)
    # Add bottom border line
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "DBE8F3")
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)


def add_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    run = p.add_run("Prepared from broker research pack and company filings.")
    run.font.name = "Helvetica"
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x5B, 0x66, 0x70)
    # Add top border
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "DBE8F3")
    pBdr.append(top)
    p._p.get_or_add_pPr().append(pBdr)


def add_docx_page(document: Document, header_title: str) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    set_page(section)
    add_header(section, header_title)
    add_footer(section)


def add_section_heading(document: Document, text: str, level: int = 1) -> None:
    """Add a navy-colored heading matching the PDF style."""
    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Helvetica"
    run.font.color.rgb = RGBColor(0x10, 0x32, 0x4A)
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(11)
    else:
        run.font.size = Pt(9.5)


def add_subtitle(document: Document, text: str) -> None:
    """Add a gray subtitle matching the PDF style."""
    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Helvetica"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x5B, 0x66, 0x70)


def add_body(document: Document, text: str) -> None:
    """Add body text matching the PDF Body style."""
    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1B, 0x1F, 0x24)


def add_bullet(document: Document, text: str) -> None:
    """Add a bullet point matching the PDF ERBullet style."""
    p = document.add_paragraph(style="List Bullet")
    # Clear existing runs and add formatted one
    for run in p.runs:
        run.clear()
    p.clear()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(8.2)
    run.font.color.rgb = RGBColor(0x1B, 0x1F, 0x24)


def add_caption(document: Document, text: str) -> None:
    """Add a centered caption below a figure."""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Helvetica"
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x5B, 0x66, 0x70)


def add_side_by_side_images(document: Document, left_path: Path, right_path: Path,
                             left_caption: str = "", right_caption: str = "") -> None:
    """Place two images side by side using a borderless table."""
    # Add a small spacer paragraph before the table
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(4)
    spacer.paragraph_format.space_after = Pt(2)
    tbl = document.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove borders
    tbl_pr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border_el = OxmlElement(f"w:{border_name}")
        border_el.set(qn("w:val"), "none")
        border_el.set(qn("w:sz"), "0")
        border_el.set(qn("w:space"), "0")
        border_el.set(qn("w:color"), "auto")
        borders.append(border_el)
    tbl_pr.append(borders)
    # Set column widths
    left_cell = tbl.cell(0, 0)
    right_cell = tbl.cell(0, 1)
    left_cell.width = Inches(3.7)
    right_cell.width = Inches(3.7)
    # Add images
    left_p = left_cell.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_run = left_p.add_run()
    left_run.add_picture(str(left_path), width=Inches(3.5))
    if left_caption:
        cap_p = left_cell.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(3)
        cap_run = cap_p.add_run(left_caption)
        cap_run.font.name = "Helvetica"
        cap_run.font.size = Pt(6.5)
        cap_run.font.color.rgb = RGBColor(0x5B, 0x66, 0x70)
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_run = right_p.add_run()
    right_run.add_picture(str(right_path), width=Inches(3.5))
    if right_caption:
        cap_p = right_cell.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(3)
        cap_run = cap_p.add_run(right_caption)
        cap_run.font.name = "Helvetica"
        cap_run.font.size = Pt(6.5)
        cap_run.font.color.rgb = RGBColor(0x5B, 0x66, 0x70)


def make_docx(charts: dict[str, Path]) -> None:
    document = Document()
    set_page(document.sections[0])
    add_header(document.sections[0], "Summary")
    add_footer(document.sections[0])
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(8.5)
    normal.font.color.rgb = RGBColor(0x1B, 0x1F, 0x24)
    # Style list bullets
    if "List Bullet" in document.styles:
        lb = document.styles["List Bullet"]
        lb.font.name = "Times New Roman"
        lb.font.size = Pt(8.2)
        lb.font.color.rgb = RGBColor(0x1B, 0x1F, 0x24)
        lb.paragraph_format.space_before = Pt(1)
        lb.paragraph_format.space_after = Pt(2)

    # ── Page 1: Title / Summary ──
    add_section_heading(document, TITLE, level=1)
    add_subtitle(document, f"Q1 FY26 Earnings Update | Rating: Outperform | PT: ${TARGET_PRICE:.0f} | Price ({PRICE_DATE}): ${CURRENT_PRICE:.2f}")

    add_body(document,
        f"Broadcom is a $1.6tn semiconductor and infrastructure software company. The semiconductor segment designs custom AI "
        f"accelerator chips (XPUs/ASICs) for hyperscale customers and sells high-speed networking silicon (switches, SerDes, DSPs) "
        f"that connects those chips inside data-center clusters. The infrastructure software segment, centered on VMware, provides "
        f"enterprise infrastructure with 93% gross margins and ~$27bn annual recurring revenue."
    )
    add_body(document,
        f"Q1 FY26 results: beat and raise. Revenue $19.3bn (+0.5% vs. Street), AI revenue $8.4bn (+2.4%), adj. EPS $2.05 (+1.0%). "
        f"Q2 guided to $22.0bn revenue (+7.3% vs. Street) and $10.7bn AI revenue (+15.1%). Operating margin 66.4% (+90bps vs. consensus). "
        f"$8.3bn operating cash flow, $7.9bn buybacks, $10bn new repurchase authorization. JPM (Harlan Sur, 5 Mar) raised PT to $500; "
        f"Goldman (Toshiya Hari, 4 Mar) added to Conviction List."
    )
    add_body(document,
        f"Key change after the print. Management disclosed >$100bn FY27 AI semiconductor revenue across ~10 GW of compute for six XPU "
        f"customers. All critical components secured through FY28. Rack-margin dilution concerns substantially eased. Networking rising "
        f"from ~33% to ~40% of AI revenue."
    )
    add_docx_table(document,
        [
            ["Rating box", "Value"],
            ["Recommendation", "Outperform"],
            ["Current price", f"${CURRENT_PRICE:.2f}"],
            ["Target price", f"${TARGET_PRICE:.0f}"],
            ["Implied upside", f"{UPSIDE:.1%}"],
            ["Market cap", f"${MARKET_CAP_B:,.0f}bn"],
            ["Enterprise value", f"${ENTERPRISE_VALUE_B:,.0f}bn"],
            ["Quarterly dividend", f"${DIVIDEND_Q:.2f}"],
        ],
        [1.8, 1.5],
    )
    add_docx_table(document,
        [
            ["Q1 FY26 dashboard", "Actual", "Street", "Variance"],
            ["Revenue", "$19.3bn", "$19.2bn", "+0.5%"],
            ["AI revenue", "$8.4bn", "$8.2bn", "+2.4%"],
            ["Gross margin", "77.0%", "76.8%", "+20bps"],
            ["Operating margin", "66.4%", "65.5%", "+90bps"],
            ["EBITDA margin", "68.0%", "66.3%", "+170bps"],
            ["Adj. EPS", "$2.05", "$2.03", "+1.0%"],
            ["Q2 revenue guide", "$22.0bn", "$20.5bn", "+7.3%"],
            ["Q2 AI revenue guide", "$10.7bn", "$9.3bn", "+15.1%"],
        ],
        [1.8, 1.0, 1.0, 0.9],
    )
    document.add_picture(str(charts["targets"]), width=Inches(7.2))
    add_caption(document, "Figure 1. Sell-side price target comparison.")
    for bullet in [
        "Guidance, not the quarter, was the inflection. Q2 AI revenue guide of $10.7bn was 15% above consensus. This shifts the debate from quarterly optionality to multi-year deployment visibility. JPM (Harlan Sur, 5 Mar) estimates >$65bn FY26 AI revenue, on track for >$100bn in FY27. Goldman added AVGO to the Conviction List.",
        "Margin fears defused. Gross margin 77.0% (vs. 76.8% Street), operating margin 66.4% (vs. 65.5%). Management explicit that rack shipments will not dilute margins. UBS (Timothy Arcuri, 5 Mar) notes Anthropic mix may include more high-margin networking plus ODM rack assembly. MS (Joseph Moore, 5 Mar) no longer models incremental margin pressure from racks.",
        "Customer base broadening across six XPU programs. Google (anchor, TPU v8 Sunfish), Anthropic (~3 GW FY27), OpenAI (newly qualified, 1+ GW), Meta (Athena/Iris), plus two unnamed (likely SoftBank/Arm, ByteDance per Goldman and Jefferies). This directly addresses the concentration bear case.",
        "Supply chain secured through FY28. Leading-edge wafers, HBM, substrates, and advanced packaging locked in. Removes risk that supplier margin demands erode Broadcom's profitability. Chip inventories at $2.96bn (vs. $2.64bn consensus, up ~30% QoQ) signal demand confidence per JPM.",
    ]:
        add_bullet(document, bullet)

    # ── Page 2: Earnings Update ──
    add_docx_page(document, "Latest earnings")
    add_section_heading(document, "Latest Earnings Update")
    add_subtitle(document, "Q1 FY26 (January quarter) | Beat and raise across all key metrics")
    add_docx_table(document,
        [
            ["Q1 FY26 detail", "Actual", "Street est.", "Beat / miss", "Commentary"],
            ["Total revenue", "$19.3bn", "$19.2bn", "+0.5%", "Beat driven by semi solutions"],
            ["AI revenue", "$8.4bn", "$8.2bn", "+2.4%", "XPU + networking above consensus"],
            ["Semi solutions", "$12.5bn", "$12.3bn", "+1.6%", "AI offset soft legacy markets"],
            ["Infra software", "$6.8bn", "$6.9bn", "-1.4%", "Stable, slightly light on renewals"],
            ["Gross margin", "77.0%", "76.8%", "+20bps", "Above despite rack shipment fears"],
            ["Operating margin", "66.4%", "65.5%", "+90bps", "Better cost flow-through on scale"],
            ["EBITDA margin", "68.0%", "66.3%", "+170bps", "Best margin leverage in 4 quarters"],
            ["Adj. EPS", "$2.05", "$2.03", "+1.0%", "Quality beat: op leverage > buybacks"],
            ["Op. cash flow", "$8.3bn", "n/a", "n/a", "Capex only $250m; FCF ~$8.1bn"],
        ],
        [1.35, 0.82, 0.82, 0.72, 3.0],
    )
    add_side_by_side_images(document,
        charts["rev_ai"], charts["q1_mix"],
        "Figure 2. Revenue and AI revenue progression.", "Figure 3. Q1 FY26 revenue composition.",
    )
    for bullet in [
        "Guidance was the inflection, not the beat. Q2 AI revenue guided to $10.7bn, 15.1% above consensus of $9.3bn. Total revenue guided to $22.0bn vs. $20.5bn Street. Goldman (Toshiya Hari, 4 Mar) notes this 'is well above GS at $9.3bn and the Street at $9.3bn,' and expects the stock to 'trade meaningfully higher.'",
        "Inventory buildup is demand-pull, not concern. Chip inventories reached $2.96bn vs. $2.64bn consensus, up ~30% QoQ. JPM (Harlan Sur, 5 Mar) interprets this as bullish at the start of a fiscal year, signaling Broadcom expects strong forward demand.",
        "Networking within AI is inflecting. Networking rose from ~25% in Q3 FY25 to ~33% in Q1 FY26 and is guided to ~40% in Q2 FY26. Networking carries higher gross margins than XPU silicon and is less customer-concentrated. Jefferies (Blayne Curtis, 4 Mar) notes Tomahawk 6 and 1.6T DSP leadership are 'furthering market share gains.'",
        "Supply chain secured through FY28. Leading-edge wafers, HBM, substrates, and CoWoS advanced packaging locked in. Goldman identifies this as 'a significant positive.' HSBC (Frank Lee, 24 Nov) noted CoWoS allocation expanding from ~150k to ~250k wafers.",
        "Capital returns remained aggressive. $7.9bn in buybacks, $3.0bn in dividends, plus a new $10bn repurchase authorization. FCF conversion above 95%. JPM estimates buybacks could contribute ~$0.15-0.20 annual EPS accretion.",
    ]:
        add_bullet(document, bullet)

    # ── Page 3: AI Engine / Customers / Competition ──
    add_docx_page(document, "AI engine & competition")
    add_section_heading(document, "AI Engine, Customer Roadmap, And Competitive Position")
    add_subtitle(document, "XPU/ASIC enablement + networking silicon + how Broadcom competes against NVIDIA, AMD, and Marvell")
    add_side_by_side_images(document,
        charts["customer_gw"], charts["net_mix"],
        "Figure 4. Estimated FY27 GW allocation by XPU customer.", "Figure 5. Networking share of AI revenue rising toward 40%.",
    )
    add_body(document,
        "What Broadcom does in AI. Broadcom's AI semiconductor business has two components. First, XPU/ASIC enablement: Broadcom "
        "designs custom accelerator chips for hyperscale customers. These are application-specific integrated circuits (ASICs) "
        "optimized for specific workloads rather than general-purpose GPUs. Google's TPU, Meta's Athena, and forthcoming OpenAI chips "
        "are all designed with Broadcom IP. Second, networking silicon: Broadcom sells the switch chips (Tomahawk 5/6, Jericho3-AI), "
        "serializer/deserializer (SerDes) IP, and optical DSPs that connect accelerators within clusters. CEO Hock Tan noted at an "
        "investor call that for every $1 spent on compute silicon, $0.15-0.20 is spent on networking."
    )
    for bullet in [
        "Google (anchor customer). Runs its entire TPU stack on Broadcom XPU designs. TPU v5e, v6 (Trillium), and v8 (Sunfish) all use Broadcom silicon and SerDes. UBS (Timothy Arcuri, 5 Mar) estimates ~3.7M TPU units in CY2026 and ~6M in CY2027. BofA flags Google may move from exclusive Broadcom sourcing to multi-vendor starting TPUv8, with MediaTek collaborating on the Zebrafish program (ramping 4Q26) per JPM Asian Tech.",
        "Anthropic (~3 GW FY27). Uses Google Cloud infrastructure with dedicated Broadcom-designed silicon. JPM (Harlan Sur, 5 Mar) estimates Anthropic deployment triples YoY. Rack-margin dilution concerns eased after management clarified ODM assembly structure. UBS notes Anthropic mix may include more high-margin networking.",
        "Meta (Athena/Iris programs). Developing custom training and inference ASICs with Broadcom. JPM and Goldman (Toshiya Hari, 4 Mar) both note Meta remains committed to the Broadcom ASIC roadmap despite concurrent NVIDIA purchases. Est. >$50bn CY26 capex.",
        "OpenAI (newly qualified, 1+ GW). Committed ~$10bn to develop custom chips with Broadcom per JPM. Previously perceived as entirely NVIDIA-dependent. Goldman notes this engagement reduces customer concentration risk.",
        "Customers 5 and 6. Not named by management; Goldman and Jefferies (Blayne Curtis, 4 Mar) suggest SoftBank/Arm and ByteDance. Expected to 'more than double' shipments in FY27.",
    ]:
        add_bullet(document, bullet)
    add_body(document,
        "Competitive positioning: AVGO vs. NVDA, AMD, and MRVL. Broadcom does not compete head-to-head with NVIDIA in selling GPUs. "
        "Instead, it provides an alternative path for hyperscalers that want custom silicon. BofA's 5 March note quantifies: NVIDIA captures "
        "$25-35bn content per GW (GPU + system + software + CPU + networking), AMD captures $15-20bn (GPU + CPU + NICs), Broadcom captures "
        "$10-15bn (XPU silicon + networking switches). Lower content per GW is offset by higher margins and expanding networking attach. "
        "Marvell (MRVL) is the closest ASIC competitor, serving Amazon's Trainium, with ~$8-12bn content per GW. HSBC (Frank Lee, 24 Nov) "
        "notes Broadcom has matched Marvell's DSP roadmap with its 1.6T solution."
    )
    add_body(document,
        "Robotics as AI demand amplifier. Training large robotics foundation models (world models, multimodal perception, "
        "simulation-heavy RL) requires the same dense accelerator clusters used for frontier LLMs but with disproportionately higher "
        "data-movement requirements. Simulation-based robot training generates large video and sensor-data volumes, increasing demand "
        "for high-bandwidth networking. Value accrues to compute and networking layers before application software."
    )

    # ── Page 4: Guidance / Margins / Capital Returns ──
    add_docx_page(document, "Margins and cash")
    add_section_heading(document, "Guidance, Margins, And Capital Returns")
    add_subtitle(document, "Q2 FY26 guide confirmed scaling AI revenue without margin sacrifice")
    add_docx_table(document,
        [
            ["Q2 FY26 guidance", "Management", "Prior Street", "Beat %", "Key read-through"],
            ["Revenue", "$22.0bn", "$20.5bn", "+7.3%", "Well above consensus; semis driving"],
            ["AI revenue", "$10.7bn", "$9.3bn", "+15.1%", "Both XPU and networking accelerating"],
            ["Non-AI semi", "$4.1bn", "~$4.0bn", "~flat", "No further deterioration in legacy"],
            ["Software", "$7.2bn", "~$7.0bn", "+2.9%", "Stable recurring cash engine"],
            ["EBITDA margin", "68.0%", "67.1%", "+90bps", "Rack margin fears eased materially"],
        ],
        [1.35, 0.92, 0.92, 0.65, 2.8],
    )
    document.add_picture(str(charts["cash"]), width=Inches(7.2))
    add_caption(document, "Figure 6. Capital return composition, Q1 FY26.")
    add_body(document,
        "Margin structure. Gross margin held at 77.0% vs. 76.8% consensus. Operating margin expanded to 66.4% vs. 65.5% Street. "
        "The key concern entering the print was that Anthropic rack shipments and pass-through content would dilute margins. Management "
        "stated these shipments use ODM assembly with Broadcom retaining high-margin silicon and networking content. MS (Joseph Moore, "
        "5 Mar) notes it 'no longer models incremental margin pressure from racks.' UBS (Timothy Arcuri, 5 Mar) raised margin "
        "assumptions post-print, noting Anthropic mix 'may now include even more high-margin networking.'"
    )
    add_body(document,
        "Supply chain locked through FY28. Broadcom confirmed all critical components (leading-edge wafers, HBM, substrates, CoWoS "
        "advanced packaging) are secured through FY28. Goldman (4 Mar) calls this 'a significant positive as it addresses a key "
        "investor concern regarding potential gross margin erosion.' HSBC (Frank Lee, 24 Nov) notes CoWoS capacity expanding from "
        "~150k to ~250k wafers."
    )
    for bullet in [
        "Cash generation remains exceptional. Q1 operating cash flow $8.3bn on capex of only $250m, implying ~$8.1bn free cash flow. FCF margin approximately 42%. UBS estimates FCF grows to >$45bn by FY27.",
        "Capital return program expanded. $7.9bn buybacks + $3.0bn dividends in Q1. New $10bn repurchase authorization. TTM shareholder returns exceed $20bn. Net debt/EBITDA declining toward 2.0x.",
        "Software is the cash backstop, not the growth driver. Infrastructure software at $6.8bn revenue, ~93% gross margin, ~$27bn ARR. UBS's 23 Feb deep dive (Timothy Arcuri) flags renewal churn as the key software risk: 3-year contracts begin renewing FY27-FY28.",
        "AI revenue trajectory accelerating. From $4.0bn in Q3 FY25 to $8.4bn in Q1 FY26 to $10.7bn guided in Q2 FY26. JPM models >$65bn FY26. Goldman projects $60bn/$130bn/$170bn FY26/FY27/FY28. Jefferies notes $20-25bn/GW in CY27 implies potential above $200bn.",
    ]:
        add_bullet(document, bullet)

    # ── Page 5: Core Investment Thesis ──
    add_docx_page(document, "Investment thesis")
    add_section_heading(document, "Core Investment Thesis")
    add_subtitle(document, "Three pillars and one amplifier supporting an Outperform rating at $500 PT")
    for para in [
        "Pillar 1: AI infrastructure platform with dual monetization. Broadcom monetizes both custom compute silicon (XPUs/ASICs) and high-speed networking for every large AI cluster deployment. This dual exposure creates multiple revenue streams per customer and reduces dependence on any single chip generation. The company manages full ASIC design and validation, orchestrates fabrication with TSMC at 3nm/5nm nodes, handles advanced packaging (2.5D/3D stacking, CoWoS), and integrates proprietary SerDes IP for ultra-high-speed chip-to-chip communication. Unlike NVIDIA, which captures $25-35bn per GW through a full-stack approach, Broadcom captures $10-15bn per GW (BofA, 5 Mar) but at higher margins. UBS's SOTP analysis shows the semiconductor business alone trading at 20x P/E, 23x EV/FCF, and 17x EV/EBITDA, 'only one turn more expensive than NVDA.'",
        "Pillar 2: Earnings power rising faster than consensus reflects. FY27 consensus EPS has moved from ~$17 to ~$20+ after the Q1 print. JPM (Harlan Sur, 5 Mar) models >$120bn FY27 AI revenue and $19.61 EPS. UBS (Timothy Arcuri) projects >$130bn AI revenue and $22.76 EPS. Jefferies (Blayne Curtis, 4 Mar) notes revenue per GW of $20-25bn in CY27 implies a $200bn theoretical upper bound. Even on conservative assumptions (FY27E EPS $20.00, 25x multiple), the stock works to $500. Goldman projects AI revenue continuing to grow into FY28 ($170bn).",
        "Pillar 3: Software and capital returns provide structural downside support. The VMware-centered infrastructure software segment generates ~$27bn ARR at ~93% gross margins. Combined with >$30bn annual FCF capacity and aggressive buybacks ($10bn new authorization), the stock has a cash-flow floor that pure-play AI names lack. UBS's SOTP assigns ~$150-180/share to software alone. The risk, flagged by UBS on 23 Feb, is that 3-year VMware renewal cycles beginning FY27-FY28 may produce churn.",
        "Amplifier: Robotics and simulation workloads deepen the networking bottleneck thesis. As robotics workloads scale (world models, simulation-heavy RL, sensor fusion), they stress exactly the infrastructure Broadcom monetizes. Data center spending must front-run robotics monetization, so value accrues to the chip and networking layer first. For every $1 spent on chips, $0.15-0.20 is spent on networking (CEO Hock Tan). This is incremental to the core LLM-driven demand thesis.",
    ]:
        add_body(document, para)
    add_docx_table(document, THESIS_MATRIX, [1.45, 3.0, 2.45])

    # ── Page 6: Competitive Landscape ──
    add_docx_page(document, "Competitive landscape")
    add_section_heading(document, "Competitive Landscape And Business Structure")
    add_subtitle(document, "ASIC vs. GPU economics, COT risk, and the software valuation backstop")
    add_side_by_side_images(document,
        charts["content_gw"], charts["ai_frames"],
        "Figure 7. Content per GW by vendor (BofA, UBS, Jefferies).", "Figure 8. FY27 AI revenue frameworks across sell-side coverage.",
    )
    add_body(document,
        "Why hyperscalers are shifting toward custom ASICs. NVIDIA's general-purpose GPUs dominated the initial phase of AI "
        "infrastructure buildout because the CUDA software ecosystem reduced time-to-deployment. As workloads mature, hyperscalers "
        "increasingly seek custom silicon optimized for their specific architectures. Custom ASICs offer lower inference cost per "
        "token, lower power consumption, and elimination of NVIDIA's margin stack. Goldman Sachs (Toshiya Hari, 4 Mar) emphasizes that "
        "Broadcom 'enables the lowest inference cost for its hyperscaler customers.' This is the core ASIC bull case: at production "
        "inference volumes, economics increasingly favor purpose-built silicon."
    )
    add_docx_table(document, ASIC_VS_GPU_ECONOMICS, [1.35, 2.75, 2.75])
    add_body(document,
        "Customer-owned tooling (COT) risk: the key bear debate. The most important competitive risk is that hyperscaler customers "
        "eventually internalize chip design. Management stated COT is 'not an imminent risk' and meaningful share loss is 'multiple "
        "years' away. JPM's Asian Tech team (5 Mar) notes MediaTek is collaborating with Google on the TPU Zebrafish project (4Q26 "
        "ramp) and next-gen Humufish. BofA (5 Mar) flags Google may move away from exclusivity starting TPUv8. Morgan Stanley (Joseph "
        "Moore, 5 Mar) notes the debate 'is not settled' but management commentary was constructive. Key nuance: networking silicon "
        "and SerDes IP create revenue streams harder to insource than XPU logic."
    )
    add_docx_table(document, COT_RISK_MATRIX, [1.0, 2.2, 0.8, 2.85])
    add_body(document,
        "Infrastructure software (~35% of revenue). Centered on VMware (acquired Nov 2023 for $69bn). ~$6.8bn quarterly revenue at "
        "~93% gross margins. ~$27bn ARR. UBS (Timothy Arcuri, 23 Feb) values software at ~$150-180/share standalone. Risk: three-year "
        "contracts begin renewing FY27-FY28. The software segment provides a cash-flow floor absent in pure-play AI names."
    )
    document.add_picture(str(charts["annual_mix"]), width=Inches(7.2))
    add_caption(document, "Figure 9. Annual revenue mix: semis growing share as AI scales.")

    # ── Page 7: Valuation ──
    add_docx_page(document, "Valuation")
    add_section_heading(document, "Valuation And Price Target")
    add_subtitle(document, f"PT ${TARGET_PRICE:.0f} | 25x FY27E EPS of $20.00 | Base case in line with broker consensus midpoint")
    add_docx_table(document,
        [
            ["Scenario", "FY27E EPS", "P/E", "Implied price", "Key assumptions"],
            ["Bear", "$15.00", "24x", "$360", "AI revenue ~$80bn FY27; software churn; capex pause; GM compresses to 74%"],
            ["Base", "$20.00", "25x", "$500", "AI revenue ~$110bn FY27; stable margins; software flat; 6 customers deploy"],
            ["Bull", "$23.00", "25x", "$575", "AI revenue >$130bn FY27; networking >40% mix; 7th customer; software re-rates"],
            ["Street high (HSBC)", "$22.50", "24x", "$535", "ASIC + networking double tailwind; highest PT on the Street"],
        ],
        [0.85, 0.85, 0.55, 0.85, 3.6],
    )
    add_side_by_side_images(document,
        charts["scenario"], charts["broker_eps"],
        "Figure 10. Bear / base / bull valuation scenarios.", "Figure 11. FY27E EPS estimates across brokers.",
    )
    for bullet in [
        "Methodology. We use a forward P/E framework on FY27E adjusted EPS, consistent with JPM ($500 PT, ~25x per Harlan Sur), UBS ($475, ~21x C2027 per Timothy Arcuri), and Morgan Stanley ($470, ~24x per Joseph Moore). Goldman (Toshiya Hari, 4 Mar) uses 30x normalized EPS of $16, arriving at $480.",
        "Multiple justification. 25x forward earnings sits at a ~10% premium to the semiconductor peer median but at a discount to pure-play AI beneficiaries trading at 30-40x. UBS notes AVGO's semi business is 'only one turn more expensive than NVDA,' which trades at ~19x P/E on CY27.",
        "Broker range context. The Street range spans $470 (MS) to $535 (HSBC). Our $500 base case sits near the midpoint. Jefferies (Blayne Curtis, 4 Mar) argues >$100bn is conservative and notes $20-25bn/GW implies potential for $200bn at scale. Post the P/E reset, Jefferies views AVGO as 'cheap.'",
        "Why not higher. Software multiple compression (VMware renewal risk per UBS 23 Feb) could offset semiconductor upside. BofA (5 Mar) flags AVGO content per GW is limited at $10-15bn vs. NVDA's $25-35bn, and Google may diversify suppliers starting TPUv8.",
    ]:
        add_bullet(document, bullet)
    document.add_picture(str(CROP_DIR / "jeff_crop_charts.png"), width=Inches(7.2))
    add_caption(document, "Figure 12. Jefferies (Blayne Curtis, 4 Mar 2026): historical earnings surprise, forward P/E, margins, and short interest.")

    # ── Page 8: Risks & Catalysts ──
    add_docx_page(document, "Risks and catalysts")
    add_section_heading(document, "Key Risks And Near-Term Catalysts")
    add_subtitle(document, "Duration risk replaces discovery risk as the primary stock debate")
    add_docx_table(document,
        [
            ["Risk", "Description", "Severity", "What to monitor"],
            ["Hyperscaler capex normalization", "AI capex priced for multi-year duration. A pause or deceleration compresses revenue trajectory and re-rates the multiple. Jefferies (4 Mar): 'the AI spend overhang will still linger.'", "High", "Quarterly capex from GOOGL, META, AMZN, MSFT"],
            ["Customer insourcing (COT)", "BofA (5 Mar): Google moving to multi-vendor starting TPUv8. JPM Asian Tech: MediaTek on Zebrafish (4Q26) and Humufish. Even mgmt acknowledges won't have '100% share at any customer indefinitely' (MS).", "Med-High", "MediaTek ramp, Google TPUv8 vendor split"],
            ["Gross margin dilution", "Rack-level shipments and HBM pass-through could lower blended margins. MS (5 Mar) expects 'some longer-term mix pressure at the consolidated level.'", "Medium", "Networking % of AI revenue, incremental GM"],
            ["Software renewal churn", "VMware 3-year contracts begin renewing FY27-FY28. UBS (23 Feb) flags customer churn, growth headwinds, and AI coding tools pushing on-prem to cloud.", "Medium", "VMware renewal rates, VCF net-new ARR"],
            ["Marvell competitive risk", "MRVL serves Amazon Trainium; matched AVGO 1.6T DSP roadmap (HSBC). Second hyperscaler win would broaden ASIC competitive pressure.", "Low-Med", "MRVL customer wins, DSP share data"],
            ["Valuation risk", "At ~25x FY27E, stock needs consistent execution. Goldman (4 Mar): 'slowdown in AI spending' and 'share loss in custom compute' as key risks.", "Medium", "Beat/miss pattern, forward guide trajectory"],
        ],
        [1.2, 2.55, 0.65, 2.3],
    )
    add_docx_table(document,
        [
            ["Near-term catalyst", "Timeline", "Expected impact"],
            ["Q2 FY26 earnings", "Jun 2026", "Confirms $10.7bn AI revenue guide; validates networking mix shift to ~40%"],
            ["Customer deployment disclosures", "Ongoing", "OpenAI/Meta/Anthropic ramp evidence reduces concentration bear case"],
            ["Networking revenue inflection", "H2 FY26", "Tomahawk 5, Jericho3-AI, SerDes attach rates drive margin-accretive growth"],
            ["Buyback execution", "Ongoing", "$10bn new auth; $0.15-0.20 annual EPS accretion potential (JPM est.)"],
            ["FY27 AI revenue framework update", "Q3/Q4 FY26", "Any upward revision to >$100bn floor is significant re-rating catalyst"],
            ["Robotics/simulation workload scaling", "CY26-27", "Incremental networking demand from video-heavy training workloads"],
        ],
        [1.8, 0.75, 4.15],
    )
    document.add_picture(str(CROP_DIR / "ms_crop_risk_reward.png"), width=Inches(7.2))
    add_caption(document, "Figure 13. Morgan Stanley risk/reward framework.")
    add_body(document, "What would increase conviction:")
    for bullet in [
        "Q2 AI revenue at or above $10.7bn guide with networking >40% of mix",
        "OpenAI deployment confirmation beyond qualification stage",
        "No visible MediaTek/Zebrafish ramp in 4Q26 timeline",
        "Gross margin sustained above 76% despite rack shipment scaling",
    ]:
        add_bullet(document, bullet)
    add_body(document, "What would reduce conviction:")
    for bullet in [
        "Hyperscaler capex guidance cuts or deferrals in upcoming earnings",
        "Google announcing multi-vendor TPUv8 with meaningful MediaTek share",
        "VMware churn rates above 15% in early renewal cohorts",
        "Marvell winning a second major hyperscaler ASIC program beyond Amazon",
    ]:
        add_bullet(document, bullet)

    # ── Page 9: Broker Debate ──
    add_docx_page(document, "Street read-through")
    add_section_heading(document, "Street Read-Through / Broker Debate")
    add_subtitle(document, "Consensus is constructive, but sell-side coverage reveals meaningful disagreement on magnitude and duration")
    add_docx_table(document, BROKER_GRID, [1.35, 0.8, 4.55])
    for bullet in [
        "JPMorgan ($500 PT, OW; Harlan Sur, 5 Mar 2026). Most detailed on customer framework. Models >$120bn FY27 AI revenue and $19.61 EPS. Estimates $12-15bn revenue per GW. Inventory build ($2.96bn vs. $2.64bn) as demand-pull signal. Components secured through FY28. Asian Tech flags MediaTek Zebrafish COT risk.",
        "UBS ($475 PT, Buy; Timothy Arcuri, 5 Mar 2026). Semi business at 20x P/E, 'only one turn more expensive than NVDA.' ~6M TPU in FY27. >$130bn AI rev, $22.76 EPS. 23 Feb deep dive flags VMware renewal churn risk. >$100bn AI floor 'still seems very conservative.'",
        "Morgan Stanley ($470 PT, OW; Joseph Moore, 5 Mar 2026). Risk/reward focused. ~$20bn/GW framework. COT 'not settled' but SerDes, packaging, and networking create barriers. Copper-based interconnect before optical adoption. No longer models rack-margin pressure.",
        "Goldman Sachs ($480 PT, Buy/CL; Toshiya Hari, 4 Mar 2026). Most aggressive out-years: $60bn/$130bn/$170bn AI rev FY26/FY27/FY28. 'Lowest inference cost' thesis. 30x normalized $16 EPS.",
        "Jefferies ($500, Franchise Pick; Blayne Curtis, 4 Mar 2026). Revenue/GW rising to $20-25bn. >$100bn is floor. Networking outgrowing ASICs. 'Struggles to see how these compute names will remain this cheap.'",
        "HSBC ($535, Buy; Frank Lee, 24 Nov 2025). Highest PT. ASIC estimates 44-59% above consensus. Matched Marvell 1.6T DSP. CoWoS 150k to 250k wafers.",
        "BofA ($450, Buy; 5 Mar 2026). Most cautious. AVGO $10-15bn/GW vs. NVDA $25-35bn. Google diversifying to MediaTek TPUv8. Key counterweight to bull consensus.",
    ]:
        add_bullet(document, bullet)
    add_docx_table(document, DEBATE_MATRIX, [1.75, 2.55, 2.4])
    add_side_by_side_images(document,
        CROP_DIR / "ubs_crop_tpu.png", CROP_DIR / "hsbc_crop_revisions.png",
        "Figure 14. UBS TPU unit estimates (5 Mar 2026).", "Figure 15. HSBC consensus revision data (24 Nov 2025).",
    )

    # ── Page 10: Appendix ──
    add_docx_page(document, "Appendix")
    add_section_heading(document, "Detailed Estimates And Portfolio Framework")
    add_subtitle(document, "House estimates, position guidance, and monitoring framework")
    add_docx_table(document, ESTIMATE_MATRIX, [2.0, 1.25, 1.25, 1.25])
    add_docx_table(document, CLIENT_FIT, [1.8, 4.9])
    add_docx_table(document, MONITORING_FRAME, [1.55, 2.8, 2.35])
    add_docx_table(document, ROBOTICS_LINK, [1.35, 2.5, 2.85])
    for bullet in [
        "Position sizing guidance. AVGO is appropriate as a 2-4% position in a diversified equity portfolio or 4-8% in a concentrated growth/tech allocation. Software cash-flow base and capital return program provide structural support absent in NVDA or AMD.",
        "Pair/hedge considerations. Long AVGO can be paired with underweight positions in higher-multiple, lower-cash-flow AI names. For COT-concerned clients, pairing AVGO with NVDA captures both GPU and custom silicon sides of the hyperscaler compute decision.",
        "Rating-change triggers (upgrade to Strong Buy). FY27 AI revenue visibility exceeding $130bn with stable margins above 76% gross. MediaTek Zebrafish delayed or descoped. Two or more new XPU customer deployments beyond current six. Software renewal rates <10% churn.",
        "Rating-change triggers (downgrade to Neutral). AI revenue guidance miss of >10%. Gross margin below 74%. Google confirming meaningful MediaTek share in TPUv8. Marvell winning second hyperscaler ASIC program. VMware churn exceeding 20%.",
    ]:
        add_bullet(document, bullet)
    add_body(document,
        "Broadcom occupies a structurally advantaged position in the AI infrastructure stack, monetizing both custom compute silicon "
        "and high-speed networking across a broadening hyperscaler customer base. The competitive landscape is not without risk: "
        "customer-owned tooling efforts (MediaTek Zebrafish at Google per BofA and JPM Asian Tech), Marvell's ambitions, and NVIDIA's "
        "full-stack dominance all constrain upside. But the Q1 FY26 print and Q2 guide confirm AI revenue is scaling without margin "
        "sacrifice, and supply security through FY28 provides a degree of earnings visibility that is rare in the semiconductor "
        "sector. We maintain Outperform with a $500 price target."
    )

    document.save(str(DOCX_PATH))


def main() -> None:
    ensure_dirs()
    crop_source_charts()
    charts = generate_charts()
    make_pdf(charts)
    make_docx(charts)
    print(PDF_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
