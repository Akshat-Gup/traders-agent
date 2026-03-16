from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pypdfium2 as pdfium
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
TMP_DIR = ROOT / "tmp" / "docs"
SOURCE_IMG_DIR = ROOT / "tmp" / "pdfs" / "source_pages"
CROP_DIR = TMP_DIR / "crops"
CHART_DIR = TMP_DIR / "charts"
PDF_DIR = ROOT / "output" / "pdf"
DOC_DIR = ROOT / "output" / "doc"
PDF_PATH = PDF_DIR / "broadcom_equity_research_report.pdf"
DOCX_PATH = DOC_DIR / "broadcom_equity_research_report.docx"

TITLE = "Broadcom Inc. (NASDAQ: AVGO)"
SUBTITLE = "Comprehensive Earnings Update for Private Bankers"
REPORT_DATE = "14 March 2026"
PRICE_DATE = "04 March 2026"

CURRENT_PRICE = 317.53
TARGET_PRICE = 500.00
UPSIDE = TARGET_PRICE / CURRENT_PRICE - 1.0
MARKET_CAP_B = 1577.8
ENTERPRISE_VALUE_B = 1613.0
NET_DEBT_B = 51.9
DIVIDEND_YIELD = 0.007
DIVIDEND_Q = 0.65

NAVY = "#10324A"
BLUE = "#2F6EA0"
TEAL = "#4F96AE"
LIGHT = "#DBE8F3"
LIGHTER = "#EEF4F8"
GOLD = "#B68A3A"
RED = "#C45454"
GRAY = "#5B6670"
DARK = "#1B1F24"

Q1_ACTUALS = {
    "revenue": 19.3,
    "ai_revenue": 8.4,
    "semi": 12.5,
    "software": 6.8,
    "gross_margin": 77.0,
    "op_margin": 66.4,
    "ebitda_margin": 68.0,
    "eps": 2.05,
    "cfo": 8.3,
    "capex": 0.25,
    "buybacks": 7.9,
    "dividends": 3.0,
    "cash": 14.2,
}

STREET = {
    "revenue": 19.2,
    "ai_revenue": 8.2,
    "semi": 12.3,
    "software": 6.9,
    "gross_margin": 76.8,
    "op_margin": 65.5,
    "ebitda_margin_q2": 67.1,
    "eps": 2.03,
    "q2_revenue": 20.5,
    "q2_ai_revenue": 9.3,
}

Q2_GUIDE = {
    "revenue": 22.0,
    "ai_revenue": 10.7,
    "non_ai_semi": 4.1,
    "software": 7.2,
    "ebitda_margin": 68.0,
}

ANNUAL_MODEL = {
    "FY25A": {"semis": 36.9, "software": 27.0, "revenue": 63.9, "eps": 6.82, "fcf": 26.9},
    "FY26E": {"semis": 63.0, "software": 29.5, "revenue": 92.5, "eps": 12.80, "fcf": 34.5},
    "FY27E": {"semis": 87.0, "software": 34.0, "revenue": 121.0, "eps": 20.00, "fcf": 45.0},
}

BROKER_PTS = [
    ("HSBC", 535),
    ("J.P. Morgan", 500),
    ("House view", 500),
    ("UBS", 475),
    ("Morgan Stanley", 470),
    ("Broad broker floor", 450),
]

BROKER_EPS_27 = [
    ("UBS", 22.76),
    ("House view", 20.00),
    ("J.P. Morgan", 19.61),
    ("Morgan Stanley", 17.98),
    ("HSBC", 16.72),
]

FY27_AI_FRAMES = [
    ("Mgmt floor", 100),
    ("JPM", 120),
    ("UBS", 130),
    ("Jefferies upside", 200),
]

BROKER_GRID = [
    ["Broker", "PT", "Key message"],
    ["J.P. Morgan", "$500", "AI revenue >$65bn in FY26 and >$100bn in FY27; OpenAI now qualified; networking mix rising."],
    ["UBS", "$475", "Margin debate improved; >$130bn FY27 AI revenue is plausible if TPU and networking scale together."],
    ["Morgan Stanley", "$470", "Visibility improved, rack-margin fears receding, and risk/reward still favors upside."],
    ["Goldman Sachs", "Buy", "Guidance and customer disclosures should settle key investor debates and support stock appreciation."],
    ["Jefferies", "Bullish / Top Pick", "10GW framework suggests current >$100bn FY27 AI view may still prove conservative."],
    ["HSBC", "$535", "Street is still underestimating ASIC plus networking upside and FY26 EPS could approach FY27 consensus."],
]

ESTIMATE_MATRIX = [
    ["Metric", "FY25A", "FY26E", "FY27E"],
    ["Revenue ($bn)", "63.9", "92.5", "121.0"],
    ["Revenue growth", "23.9%", "44.8%", "30.8%"],
    ["Semiconductor revenue ($bn)", "36.9", "63.0", "87.0"],
    ["Software revenue ($bn)", "27.0", "29.5", "34.0"],
    ["AI revenue ($bn)", "19.0", "65.0", "110.0"],
    ["Adj. EBITDA margin", "67.6%", "68.2%", "68.0%"],
    ["Adj. EPS ($)", "6.82", "12.80", "20.00"],
    ["Free cash flow ($bn)", "26.9", "34.5", "45.0"],
]

SOURCE_MATRIX = [
    ["Source", "Use in this report"],
    ["J.P. Morgan full take (5 Mar 2026)", "Primary factual base for quarter, AI revenue path, customer ramps, and FY26-FY27 EPS framing."],
    ["UBS strong results note (5 Mar 2026)", "Margin interpretation, AI networking mix, and higher-end FY27 AI revenue debate."],
    ["Morgan Stanley AI momentum builds (5 Mar 2026)", "Risk/reward framework, valuation ranges, and scenario mapping."],
    ["Goldman Sachs first take / follow-up", "Street reaction framing and debate resolution around customers and guidance."],
    ["Jefferies strong 2027 outlook", "Revenue-per-GW and networking share upside case."],
    ["HSBC potent mix of ASIC and networking upside", "High-end ASIC, CoWoS and networking TAM assumptions."],
    ["UBS software deep dive / SOTP preview", "Software renewal and valuation-backstop risk discussion."],
]

THESIS_MATRIX = [
    ["Thesis pillar", "What supports it now", "What would weaken it"],
    ["AI platform breadth", "Six-customer framework, stronger networking mix, and multiple programs beyond Google.", "A material delay in Meta, OpenAI or Anthropic ramps."],
    ["Rising earnings power", "Broker estimates moved up, Q2 guide beat, and FY27 AI revenue floor is now explicit.", "Guide cuts or lower AI revenue conversion into EPS."],
    ["Cash flow support", "Software stability, large buybacks, and modest capex intensity relative to operating cash flow.", "Software churn or a sharp increase in capital intensity."],
]

DEBATE_MATRIX = [
    ["Debate", "Current broker read-through", "House take"],
    ["How big can FY27 AI revenue get?", "Mgmt says >$100bn; JPM $120bn; UBS >$130bn; Jefferies sees upside far above that.", "Use $110bn in our base case and leave upside for execution."],
    ["Will rack mix hurt margins?", "Concern has eased materially after the quarter.", "Still watch closely, but no longer a thesis-breaker."],
    ["Does software deserve a premium multiple?", "Views diverge; UBS is the most explicit on churn and renewal risks.", "Software is a backstop, not the main driver of upside."],
    ["Is valuation already full?", "Depends on whose FY27 EPS you believe.", "Not full if Broadcom delivers near current AI and cash-return expectations."],
]

CLIENT_FIT = [
    ["Client-use appendix", "Practical interpretation"],
    ["Who this suits", "Growth-oriented clients who want AI exposure backed by cash flow and capital returns."],
    ["Who should be careful", "Clients needing low volatility or immediate certainty on hyperscaler spending trends."],
    ["Best holding period", "12-24 months so FY27 deployment visibility can develop."],
    ["What would make us more bullish", "Evidence that networking mix and customer breadth continue to expand without margin leakage."],
    ["What would make us less bullish", "AI demand pauses, software renewals deteriorate, or buybacks slow materially."],
]

MONITORING_FRAME = [
    ["What to monitor next", "Why it matters", "What would be positive"],
    ["AI revenue mix", "This will show whether networking continues to widen the moat and improve quality of growth.", "Networking remains close to or above 40% of AI revenue."],
    ["Customer breadth", "The stock will rerate more easily if investors believe Broadcom is not dependent on one customer family.", "OpenAI, Meta and Anthropic commentary becomes more concrete."],
    ["Software renewal health", "This is the main cushion under the valuation if AI spending volatility rises.", "Renewals stabilize and churn concerns prove manageable."],
    ["Capital returns", "Buybacks and dividends make the stock easier to own in diversified client portfolios.", "Repurchases remain aggressive without hurting leverage discipline."],
]

ROBOTICS_LINK = [
    ["Robotics AI layer", "Why it matters", "Broadcom exposure"],
    ["Compute infrastructure", "Large robotics foundation models need the same high-end clusters used for frontier LLM training.", "Direct via custom XPU / ASIC programs and related silicon IP."],
    ["High-speed networking", "Simulation, sensor fusion and video-heavy robot training move enormous amounts of data between chips and racks.", "Direct and significant via switching, interconnect and networking silicon."],
    ["Data layer", "Proprietary video, factory and sensor data are major moats, but difficult to access in public markets.", "Mostly indirect; Broadcom benefits when data volumes force more infrastructure spend."],
    ["Simulation / digital twins", "Robotics developers increasingly train in sim before deployment in the physical world.", "Indirect; more simulation means more compute and networking demand."],
]


def ensure_dirs() -> None:
    for path in (TMP_DIR, SOURCE_IMG_DIR, CROP_DIR, CHART_DIR, PDF_DIR, DOC_DIR):
        path.mkdir(parents=True, exist_ok=True)


def style_matplotlib() -> None:
    plt.style.use("default")
    plt.rcParams.update(
        {
            "figure.facecolor": "white",
            "axes.facecolor": "white",
            "axes.edgecolor": "#C9D5DF",
            "axes.labelcolor": DARK,
            "text.color": DARK,
            "axes.titleweight": "bold",
            "font.size": 10,
            "font.family": "DejaVu Serif",
        }
    )


def save_barh_chart(path: Path, title: str, subtitle: str, labels: list[str], values: list[float], colors_list: list[str]) -> None:
    fig, ax = plt.subplots(figsize=(8.4, 3.8))
    y = np.arange(len(labels))
    ax.barh(y, values, color=colors_list)
    ax.set_yticks(y, labels)
    ax.invert_yaxis()
    ax.set_title(title, loc="left", fontsize=15, color=NAVY, pad=18)
    ax.text(0.0, 1.03, subtitle, transform=ax.transAxes, fontsize=9, color=GRAY)
    ax.grid(axis="x", color="#DCE4EA", linewidth=0.8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for idx, value in enumerate(values):
        ax.text(value + max(values) * 0.01, idx, f"${value:.0f}", va="center", fontsize=10, color=NAVY)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_grouped_bar_chart(path: Path, title: str, subtitle: str, categories: list[str], left_vals: list[float], right_vals: list[float], left_label: str, right_label: str) -> None:
    x = np.arange(len(categories))
    width = 0.36
    fig, ax = plt.subplots(figsize=(8.4, 4.0))
    ax.bar(x - width / 2, left_vals, width, label=left_label, color=BLUE)
    ax.bar(x + width / 2, right_vals, width, label=right_label, color=TEAL)
    ax.set_xticks(x, categories)
    ax.set_title(title, loc="left", fontsize=15, color=NAVY, pad=18)
    ax.text(0.0, 1.03, subtitle, transform=ax.transAxes, fontsize=9, color=GRAY)
    ax.grid(axis="y", color="#DCE4EA", linewidth=0.8)
    ax.legend(frameon=False, loc="upper left")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for bars in ax.containers:
        ax.bar_label(bars, fmt="%.1f", padding=3, fontsize=9)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_stacked_bar_chart(path: Path, title: str, subtitle: str, categories: list[str], lower_vals: list[float], upper_vals: list[float], lower_label: str, upper_label: str) -> None:
    x = np.arange(len(categories))
    fig, ax = plt.subplots(figsize=(8.4, 4.0))
    ax.bar(x, lower_vals, label=lower_label, color=BLUE)
    ax.bar(x, upper_vals, bottom=lower_vals, label=upper_label, color=TEAL)
    ax.set_xticks(x, categories)
    ax.set_title(title, loc="left", fontsize=15, color=NAVY, pad=18)
    ax.text(0.0, 1.03, subtitle, transform=ax.transAxes, fontsize=9, color=GRAY)
    ax.grid(axis="y", color="#DCE4EA", linewidth=0.8)
    ax.legend(frameon=False, loc="upper left")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for idx, total in enumerate(np.array(lower_vals) + np.array(upper_vals)):
        ax.text(idx, total + max(total * 0.015, 0.12), f"{total:.1f}", ha="center", fontsize=9, color=NAVY)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_scenario_chart(path: Path, title: str, subtitle: str, labels: list[str], values: list[float], current_price: float) -> None:
    fig, ax = plt.subplots(figsize=(8.4, 4.0))
    colors_list = [RED, BLUE, GOLD]
    ax.bar(labels, values, color=colors_list)
    ax.axhline(current_price, color="#666666", linestyle="--", linewidth=1.2, label=f"Current ${current_price:.0f}")
    ax.set_title(title, loc="left", fontsize=15, color=NAVY, pad=18)
    ax.text(0.0, 1.03, subtitle, transform=ax.transAxes, fontsize=9, color=GRAY)
    ax.grid(axis="y", color="#DCE4EA", linewidth=0.8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.legend(frameon=False, loc="upper left")
    for idx, value in enumerate(values):
        ax.text(idx, value + 8, f"${value:.0f}", ha="center", fontsize=10, color=NAVY)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def save_capital_chart(path: Path) -> None:
    labels = ["CFO", "Capex", "Buybacks", "Dividends"]
    values = [Q1_ACTUALS["cfo"], Q1_ACTUALS["capex"], Q1_ACTUALS["buybacks"], Q1_ACTUALS["dividends"]]
    fig, ax = plt.subplots(figsize=(8.4, 4.0))
    ax.bar(labels, values, color=[BLUE, TEAL, GOLD, RED])
    ax.set_title("Q1 FY26 cash generation and capital return", loc="left", fontsize=15, color=NAVY, pad=18)
    ax.text(0.0, 1.03, "Cash from operations easily covered repurchases and dividends", transform=ax.transAxes, fontsize=9, color=GRAY)
    ax.grid(axis="y", color="#DCE4EA", linewidth=0.8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for idx, value in enumerate(values):
        ax.text(idx, value + 0.22, f"${value:.2f}bn" if value < 1 else f"${value:.1f}bn", ha="center", fontsize=10, color=NAVY)
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def render_pdf_page(pdf_path: Path, page_no: int, out_path: Path) -> None:
    pdf = pdfium.PdfDocument(str(pdf_path))
    pdf[page_no - 1].render(scale=2.2).to_pil().save(out_path)


def ensure_source_page(name: str, pdf_rel: str, page_no: int) -> Path:
    out_path = SOURCE_IMG_DIR / name
    if not out_path.exists():
        render_pdf_page(ROOT / pdf_rel, page_no, out_path)
    return out_path


def crop_source_charts() -> None:
    jpm_page = ensure_source_page("jpm_page2.png", "report/jp-broadcom-fulltake.pdf", 2)
    ms_page = ensure_source_page("ms_page4.png", "report/VisionAlpha_Report_1773474642089.pdf", 4)
    ubs_page = ensure_source_page("ubs_p4.png", "report/VisionAlpha_Report_1773474623069.pdf", 4)
    hsbc_page = ensure_source_page("hsbc_p4.png", "report/VisionAlpha_Report_1773474758155.pdf", 4)
    jeff_page = ensure_source_page("jeff_va_p4.png", "report/VisionAlpha_Report_1773474668082.pdf", 4)
    jpm_crop = CROP_DIR / "jpm_crop_metrics.png"
    ms_crop = CROP_DIR / "ms_crop_risk_reward.png"
    ubs_crop = CROP_DIR / "ubs_crop_tpu.png"
    hsbc_crop = CROP_DIR / "hsbc_crop_revisions.png"
    jeff_crop = CROP_DIR / "jeff_crop_charts.png"
    if not jpm_crop.exists():
        img = Image.open(jpm_page)
        img.crop((120, 380, 1480, 2600)).save(jpm_crop)
    if not ms_crop.exists():
        img = Image.open(ms_page)
        img.crop((120, 620, 1940, 2850)).save(ms_crop)
    if not ubs_crop.exists():
        img = Image.open(ubs_page)
        img.crop((90, 120, 2340, 2380)).save(ubs_crop)
    if not hsbc_crop.exists():
        img = Image.open(hsbc_page)
        img.crop((90, 220, 2330, 3180)).save(hsbc_crop)
    if not jeff_crop.exists():
        img = Image.open(jeff_page)
        img.crop((80, 230, 2290, 2540)).save(jeff_crop)


def generate_charts() -> dict[str, Path]:
    style_matplotlib()
    charts = {
        "targets": CHART_DIR / "price_targets.png",
        "rev_ai": CHART_DIR / "revenue_ai_progression.png",
        "q1_mix": CHART_DIR / "q1_revenue_mix.png",
        "ai_mix": CHART_DIR / "ai_mix_q1_q2.png",
        "ai_frames": CHART_DIR / "fy27_ai_frameworks.png",
        "cash": CHART_DIR / "cash_returns.png",
        "annual_mix": CHART_DIR / "annual_mix.png",
        "scenario": CHART_DIR / "valuation_scenarios.png",
        "broker_eps": CHART_DIR / "broker_eps_fy27.png",
    }
    save_barh_chart(
        charts["targets"],
        "Price targets sampled from the broker pack",
        "Broadcom remains broadly rated bullish after the March quarter",
        [x[0] for x in BROKER_PTS],
        [x[1] for x in BROKER_PTS],
        [GOLD, BLUE, NAVY, TEAL, LIGHT, "#B8CBD9"],
    )
    save_grouped_bar_chart(
        charts["rev_ai"],
        "Revenue and AI revenue are stepping up together",
        "Q1 FY26 actuals and Q2 FY26 guidance versus prior-year quarter",
        ["Q1 FY25A", "Q1 FY26A", "Q2 FY26G"],
        [14.9, 19.3, 22.0],
        [4.1, 8.4, 10.7],
        "Total revenue",
        "AI revenue",
    )
    save_stacked_bar_chart(
        charts["q1_mix"],
        "Q1 FY26 revenue mix",
        "AI is already the largest piece of the Broadcom story",
        ["Q1 FY26"],
        [5.63 + 2.77, ][:1],
        [4.12 + 6.80, ][:1],
        "AI revenue",
        "Non-AI semi + software",
    )
    save_stacked_bar_chart(
        charts["ai_mix"],
        "AI mix is shifting toward networking",
        "One-third of AI revenue in Q1; management says 40% in Q2",
        ["Q1 FY26A", "Q2 FY26G"],
        [5.63, 6.42],
        [2.77, 4.28],
        "AI compute / ASIC",
        "AI networking",
    )
    save_barh_chart(
        charts["ai_frames"],
        "FY27 AI revenue frameworks in the source pack",
        "Management floor sits materially below the most bullish broker views",
        [x[0] for x in FY27_AI_FRAMES],
        [x[1] for x in FY27_AI_FRAMES],
        [BLUE, TEAL, GOLD, NAVY],
    )
    save_capital_chart(charts["cash"])
    save_stacked_bar_chart(
        charts["annual_mix"],
        "Our revenue mix view: AI semis drive growth, software anchors cash flow",
        "Amounts in $bn; semis include both AI and non-AI semiconductor solutions",
        list(ANNUAL_MODEL.keys()),
        [ANNUAL_MODEL[k]["semis"] for k in ANNUAL_MODEL],
        [ANNUAL_MODEL[k]["software"] for k in ANNUAL_MODEL],
        "Semiconductor solutions",
        "Infrastructure software",
    )
    save_scenario_chart(
        charts["scenario"],
        "12-month valuation scenarios",
        "Base case uses 25x FY27E adjusted EPS of $20.00",
        ["Bear", "Base", "Bull"],
        [360, 500, 575],
        CURRENT_PRICE,
    )
    save_barh_chart(
        charts["broker_eps"],
        "FY27 EPS framing across key brokers",
        "Wide range reflects disagreement on how fast AI revenue converts into earnings",
        [x[0] for x in BROKER_EPS_27],
        [x[1] for x in BROKER_EPS_27],
        [GOLD, NAVY, BLUE, TEAL, LIGHT],
    )
    return charts


def header_footer(canvas, doc) -> None:
    canvas.saveState()
    w, h = letter
    canvas.setStrokeColor(colors.HexColor(LIGHT))
    canvas.line(doc.leftMargin, h - 0.42 * inch, w - doc.rightMargin, h - 0.42 * inch)
    canvas.setFont("Helvetica-Bold", 8.5)
    canvas.setFillColor(colors.HexColor(NAVY))
    canvas.drawString(doc.leftMargin, h - 0.34 * inch, "Broadcom Inc. - Equity Research Update")
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor(GRAY))
    canvas.drawRightString(w - doc.rightMargin, h - 0.34 * inch, REPORT_DATE)
    canvas.line(doc.leftMargin, 0.44 * inch, w - doc.rightMargin, 0.44 * inch)
    canvas.drawString(doc.leftMargin, 0.30 * inch, "Prepared for private-banker use from the provided report pack and prior extracted company data.")
    canvas.drawRightString(w - doc.rightMargin, 0.30 * inch, f"Page {canvas.getPageNumber()}")
    canvas.restoreState()


def make_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="Body",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=7.3,
            leading=8.4,
            textColor=colors.HexColor(DARK),
            spaceAfter=2.5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Tiny",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=6.6,
            leading=7.4,
            textColor=colors.HexColor(DARK),
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ERBullet",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=7.2,
            leading=8.2,
            leftIndent=12,
            firstLineIndent=-8,
            textColor=colors.HexColor(DARK),
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TitleLarge",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=18,
            textColor=colors.HexColor(NAVY),
            alignment=TA_LEFT,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Subtitle",
            parent=styles["Heading2"],
            fontName="Helvetica",
            fontSize=8.8,
            leading=10,
            textColor=colors.HexColor(GRAY),
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Section",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11.5,
            leading=12.5,
            textColor=colors.HexColor(NAVY),
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CenterNote",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=6.2,
            leading=7.0,
            textColor=colors.HexColor(GRAY),
            alignment=TA_CENTER,
            spaceAfter=1,
        )
    )
    return styles


def info_table(rows: list[list[str]], col_widths: list[float], highlight_col: int | None = None) -> Table:
    table = Table(rows, colWidths=col_widths, hAlign="LEFT")
    style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(NAVY)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 7.2),
        ("TOPPADDING", (0, 0), (-1, 0), 4),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 4),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(LIGHTER)]),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 1), (-1, -1), 7.0),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C7D4DE")),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 1), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 2),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]
    if highlight_col is not None:
        style.extend(
            [
                ("BACKGROUND", (highlight_col, 1), (highlight_col, -1), colors.HexColor(LIGHT)),
                ("FONTNAME", (highlight_col, 1), (highlight_col, -1), "Helvetica-Bold"),
            ]
        )
    table.setStyle(TableStyle(style))
    return table


def image_flow(path: Path, width: float, height: float, caption: str | None = None, styles=None):
    img = RLImage(str(path), width=width, height=height)
    if caption:
        return [img, Paragraph(caption, styles["CenterNote"])]
    return [img]


def side_by_side(left_flowables: list, right_flowables: list, widths: tuple[float, float] = (3.45 * inch, 3.45 * inch)) -> Table:
    tbl = Table([[left_flowables, right_flowables]], colWidths=list(widths), hAlign="LEFT")
    tbl.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    return tbl


def make_pdf(charts: dict[str, Path]) -> None:
    styles = make_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.42 * inch,
        rightMargin=0.42 * inch,
        topMargin=0.56 * inch,
        bottomMargin=0.50 * inch,
    )
    story = []

    def add_title(title: str, subtitle: str) -> None:
        story.append(Paragraph(title, styles["TitleLarge"]))
        story.append(Paragraph(subtitle, styles["Subtitle"]))

    add_title(TITLE, f"{SUBTITLE} | Rating: Outperform | Target Price: ${TARGET_PRICE:.0f} | Price as of {PRICE_DATE}: ${CURRENT_PRICE:.2f}")
    rating_table = info_table(
        [
            ["Rating box", "Value"],
            ["Recommendation", "Outperform"],
            ["Current price", f"${CURRENT_PRICE:.2f}"],
            ["Target price", f"${TARGET_PRICE:.0f}"],
            ["Implied upside", f"{UPSIDE:.1%}"],
            ["Market cap", f"${MARKET_CAP_B:,.0f}bn"],
            ["Enterprise value", f"${ENTERPRISE_VALUE_B:,.0f}bn"],
            ["Quarterly dividend", f"${DIVIDEND_Q:.2f}"],
        ],
        [1.6 * inch, 1.45 * inch],
        highlight_col=1,
    )
    company_snapshot = [
        Paragraph("<b>What Broadcom is:</b> Broadcom helps build the hidden infrastructure that allows large AI systems to train and run efficiently. It helps hyperscale customers such as Google, Meta, and Anthropic design custom accelerator chips — known in the industry as XPUs or ASICs — and it also sells the networking hardware and silicon that move data between those chips at very high speed. Additionally, it owns a large enterprise-software business, mainly through its VMware-related infrastructure platform, which provides steady, recurring cash flow that anchors the rest of the story. In simple terms, if you think of a modern AI data center as a factory, Broadcom builds both the custom engines and the highway system that connects them.", styles["Body"]),
        Paragraph("<b>Why the stock matters now:</b> The January-quarter results and April-quarter guidance changed the conversation around Broadcom from one of quarterly optionality to multi-year visibility. Management disclosed line of sight to more than $100bn of AI semiconductor revenue in FY27 across roughly 10 GW of deployed compute capacity for six named or identifiable XPU customers, while every major broker in our source pack raised estimates and price targets in the days following the print. Crucially, margin fears around rack-level shipments — one of the biggest pre-earnings overhangs — were substantially eased, and the company announced it had secured all critical components through FY28. At the same time, the $10bn share-repurchase authorization and $8.3bn of Q1 operating cash flow make Broadcom an AI story that pays its way rather than requiring investors to suspend disbelief on free-cash-flow timing.", styles["Body"]),
        Paragraph("<b>Who this note is for:</b> This report is designed for a private banker who needs to explain Broadcom to clients clearly: what the company actually does, what changed in the latest quarter, where the earnings-power trajectory is heading, what the main risks are, and why the valuation still offers attractive upside despite the stock's strong performance. It is written to be accessible without sacrificing analytical depth.", styles["Body"]),
    ]
    story.append(side_by_side([rating_table], company_snapshot))
    story.append(Spacer(1, 0.06 * inch))
    summary_table = info_table(
        [
            ["Q1 FY26 summary", "Actual", "Street", "Variance"],
            ["Revenue", "$19.3bn", "$19.2bn", "+0.5%"],
            ["AI revenue", "$8.4bn", "$8.2bn", "+2.4%"],
            ["Gross margin", "77.0%", "76.8%", "+20bps"],
            ["Operating margin", "66.4%", "65.5%", "+90bps"],
            ["Adj. EPS", "$2.05", "$2.03", "+1.0%"],
            ["Q2 revenue guide", "$22.0bn", "$20.5bn", "+7.3%"],
            ["Q2 AI guide", "$10.7bn", "$9.3bn", "+15.1%"],
        ],
        [1.8 * inch, 1.05 * inch, 1.05 * inch, 1.1 * inch],
        highlight_col=1,
    )
    summary_bullets = [
        Paragraph("1. <b>The quarter itself was solid; the guidance was transformational.</b> Q1 FY26 revenue of $19.3bn and adjusted EPS of $2.05 both edged past consensus, but the real story was the April-quarter guide: $22.0bn of revenue and $10.7bn of AI revenue, which were 7% and 15% above consensus respectively. AI revenue is now running at more than twice the level of a year ago, and management's commentary — supported by J.P. Morgan, UBS, and Goldman Sachs analyses — suggests the acceleration is broadening across customers rather than concentrating further in Google.", styles["ERBullet"]),
        Paragraph("2. <b>Broadcom is the scaled alternative to merchant-GPU economics.</b> Hyperscalers increasingly want custom accelerators that are optimized for their specific workloads and that do not carry Nvidia's pricing premium. Broadcom is one of the only companies that can deliver both the custom chip design and the high-speed networking silicon needed to connect thousands of those chips inside a single cluster. That dual capability — compute plus interconnect — is the heart of the investment case and is why CEO Hock Tan has repeatedly noted that for every dollar spent on chips, an additional $0.15-0.20 is spent on networking.", styles["ERBullet"]),
        Paragraph("3. <b>The next leg depends on duration, not discovery.</b> The market already knows Broadcom benefits from AI. The open question is whether AI revenue can progress from a FY26 run-rate story into visible, multi-year FY27/FY28 customer deployments without gross-margin damage or customer concentration risk. That is why the four variables to monitor are networking mix (heading toward 40% of AI revenue), customer breadth (six programs and growing), capital-return discipline ($10bn buyback, 0.7% dividend yield), and software stability (VMware renewals approaching a cyclical test).", styles["ERBullet"]),
        Paragraph("4. <b>Robotics and simulation workloads quietly strengthen the infrastructure thesis.</b> Training large robotics foundation models — world models, multimodal perception stacks, simulation-heavy reinforcement learning — requires the same dense accelerator clusters used for frontier LLM training, but with even heavier data movement due to video and sensor volumes. As these workloads scale, they amplify demand for the very networking and compute bottlenecks Broadcom monetizes, adding an incremental demand layer that the market has not yet fully priced.", styles["ERBullet"]),
    ]
    story.append(side_by_side([summary_table], image_flow(charts["targets"], 3.35 * inch, 2.4 * inch, "Figure 1. Price targets sampled from the provided broker pack.", styles)))
    story.append(Spacer(1, 0.04 * inch))
    story.extend(summary_bullets)
    story.append(PageBreak())

    add_title("Latest Earnings Update", "Q1 FY26 was not just a beat — it changed the shape of forward expectations and eased the market's biggest margin concerns")
    story.append(
        Paragraph(
            "Broadcom's Q1 FY26 results, reported on 4 March 2026, delivered a clean beat on virtually every line item while simultaneously raising the credibility of the company's medium-term AI revenue framework. Total revenue of $19.3bn exceeded the Street's $19.2bn, driven by semiconductor solutions at $12.5bn versus consensus of $12.3bn. Infrastructure software was stable at $6.8bn, slightly below the $6.9bn expectation but well within normal variance. Adjusted EPS of $2.05 came in above the $2.03 consensus, and the quality of the beat was arguably more important than its size: operating margin hit 66.4% versus the Street's 65.5%, indicating strong cost flow-through as AI volumes scaled.",
            styles["Body"],
        )
    )
    story.append(Paragraph("The real inflection, however, came from the guidance. Management pointed to Q2 FY26 revenue of $22.0bn — 7.3% above the $20.5bn consensus — and AI revenue of $10.7bn, a full 15% above the $9.3bn consensus. That guidance gap is meaningful because it is not the kind of small variance that can be explained by timing or accounting; it signals that customer order books are filling faster than the Street had modeled, and that Broadcom's forward visibility has improved materially. As J.P. Morgan noted in its 5 March note, the combination of a wider customer base and stronger networking attach puts Broadcom on track to deliver over $65bn of AI revenue in FY26, which would represent roughly 105% growth year-over-year.", styles["Body"]))
    beat_table = info_table(
        [
            ["Beat / miss dashboard", "Actual", "Street", "Comment"],
            ["Total revenue", "$19.3bn", "$19.2bn", "Beat driven by semiconductor solutions; AI revenue more than offset soft legacy markets"],
            ["AI revenue", "$8.4bn", "$8.2bn", "Custom ASIC plus networking upside; AI compute grew ~28% QoQ, networking ~34% QoQ per UBS estimates"],
            ["Semi solutions", "$12.5bn", "$12.3bn", "Non-AI semis roughly stable at $4.1bn; enterprise networking, broadband and storage up YoY, offset by seasonal wireless decline"],
            ["Infrastructure software", "$6.8bn", "$6.9bn", "Stable and highly profitable (93% gross margin, 78% operating margin); slightly light but within normal range"],
            ["Gross margin", "77.0%", "76.8%", "In line to slightly better; no rack-related dilution visible, easing the biggest pre-earnings overhang"],
            ["Operating margin", "66.4%", "65.5%", "Meaningfully better cost flow-through; opex at $2.0bn came in below the $2.2bn consensus, reflecting operating discipline"],
            ["EBITDA margin", "68.0%", "66.3%", "170bps better than expected; underscores that the AI ramp is financially attractive, not just a revenue story"],
            ["Adjusted EPS", "$2.05", "$2.03", "Small headline beat, but materially better quality than the size suggests once margin expansion is accounted for"],
        ],
        [1.55 * inch, 1.0 * inch, 1.0 * inch, 3.35 * inch],
    )
    story.append(beat_table)
    story.append(Spacer(1, 0.05 * inch))
    story.append(
        side_by_side(
            image_flow(charts["rev_ai"], 3.35 * inch, 2.35 * inch, "Figure 2. Revenue and AI revenue progression.", styles),
            image_flow(charts["q1_mix"], 3.35 * inch, 2.35 * inch, "Figure 3. Q1 revenue mix shows AI is already central.", styles),
        )
    )
    story.append(
        Paragraph(
            "For a banker explaining the quarter to a client, the simplest takeaway is this: Broadcom is no longer merely an AI beneficiary; it is already operating at AI scale. AI revenue was 43% of total company revenue in Q1, up from roughly 27% a year ago, while non-AI semiconductors held roughly flat. That means the stock's next moves will be driven primarily by whether AI demand keeps compounding and whether those revenues remain high-quality. The fact that operating margins expanded rather than contracted during this scaling phase is especially important, because it undercuts the bear argument that Broadcom's incremental AI revenue would carry lower profitability than its legacy mix.",
            styles["Body"],
        )
    )
    story.append(Paragraph("One underappreciated signal was the inventory build. Broadcom's chip inventories rose approximately 30% quarter-over-quarter to $2.96bn, versus a consensus expectation of $2.64bn. In isolation, rising inventories can be a warning sign. But at the start of a major capacity ramp — with the company simultaneously guiding sharply higher revenue — this inventory build is more naturally interpreted as confidence in forward demand. J.P. Morgan's semiconductor team made exactly this point, noting that the inventory position signals Broadcom expects strong pull-through over the coming quarters rather than a slowdown. This reading is further supported by the company's statement that it has already secured all critical components — leading-edge wafers, HBM memory, substrates, and advanced packaging capacity — needed to support its revenue forecast through FY28. That is an unusually long supply-chain visibility window and suggests that management's comfort with the inventory position is grounded in contractual commitments, not hope.", styles["Body"]))
    story.append(Paragraph("The quarter also matters for the broader AI infrastructure debate. Broadcom's results confirm that the data-center spending cycle is not merely a GPU story. Custom silicon and high-speed networking are scaling alongside merchant GPUs, and Broadcom is capturing a growing share of that spend. The fact that AI networking revenue grew roughly 34% sequentially in Q1, according to UBS estimates, while AI compute grew roughly 28% sequentially, suggests that the networking attach rate is increasing — which is precisely what the company has been arguing would happen as cluster sizes expand and data-movement requirements intensify.", styles["Body"]))
    story.append(PageBreak())

    add_title("AI Engine And Customer Roadmap", "The AI debate has shifted from quarterly surprise to duration — and the customer base is wider than many investors realize")
    story.append(
        Paragraph(
            "To understand why Broadcom's AI business matters, it helps to start with what the company actually does inside an AI data center. Broadcom designs custom accelerator chips — sometimes called XPUs or ASICs — for hyperscale customers who want silicon tailored to their specific AI workloads. Google's TPUs, for example, are designed by Broadcom. These chips do not carry Nvidia's margins or Nvidia's general-purpose overhead; instead, they are optimized for particular training or inference tasks, which lowers the cost per computation for the customer. Broadcom also designs and sells the high-speed networking silicon — switches, serializer/deserializers (SerDes), and optical interconnect components — that moves data between accelerator chips and between racks inside a cluster. In the industry's language, SerDes IP allows chips to talk to each other at extremely high speeds, and Broadcom's proprietary SerDes technology is among the fastest in the world. The combination of custom compute and custom networking makes Broadcom a system-level infrastructure partner rather than a component vendor.",
            styles["Body"],
        )
    )
    story.append(Paragraph("Management disclosed on the Q1 FY26 call that it now has line of sight to more than $100bn of AI semiconductor revenue in FY27, generated across roughly 10 GW of deployed compute capacity for six XPU customers. That framing was not a casual target; it was built bottom-up from customer-by-customer deployment plans. J.P. Morgan's 5 March analysis pushed that framework to $120bn or more, estimating that Broadcom can generate $12-15bn of revenue per deployed gigawatt. UBS went further, arguing that its supply-chain work now suggests Broadcom's TPU unit shipments could reach approximately 6 million units next year, with revenue per GW in a range similar to AMD's roughly $15bn, yielding a plausible FY27 AI revenue figure above $130bn. Jefferies offered the most aggressive framing, arguing that if revenue per GW rises from approximately $13bn in CY25 to $20-25bn in CY27 — driven by networking attach and die-size growth — AI revenue could theoretically approach $200bn, although that requires assumptions at the top end of every variable.", styles["Body"]))
    story.append(
        side_by_side(
            image_flow(charts["ai_mix"], 3.35 * inch, 2.35 * inch, "Figure 4. AI networking is becoming a larger part of the mix.", styles),
            image_flow(charts["ai_frames"], 3.35 * inch, 2.35 * inch, "Figure 5. FY27 AI revenue frameworks vary widely.", styles),
        )
    )
    for line in [
        "1. <b>Google remains the foundational anchor.</b> Google's TPU programs — including the current Ironwood generation and the upcoming TPU v8 (Sunfish) — are the backbone of Broadcom's AI business. The relationship has been running for over a decade, and the source pack consistently identifies Google as providing the largest share of Broadcom's near-term AI revenue. Critically, management stated that it does not expect material impact from customer-owned tooling (COT) — the risk that Google or others internalize more chip-design work — for at least several more years. This matters because the COT debate has been one of the most persistent bear arguments against Broadcom's durability.",
        "2. <b>Anthropic is scaling fast.</b> Anthropic, which uses Google Cloud and Amazon Web Services for its infrastructure, is ramping its Broadcom-designed custom silicon aggressively. J.P. Morgan and Goldman Sachs both note that Anthropic is already deploying its initial gigawatt of capacity and expects to reach 3 GW in FY27, a threefold year-over-year increase. This was one of the biggest margin debates heading into the print, because the Anthropic relationship was initially structured around rack-level shipments, which investors feared would dilute Broadcom's gross margins. Management defused this concern, indicating that the mix of what it ships to Anthropic may now include more high-margin networking and that ODM partners may handle some of the rack assembly.",
        "3. <b>OpenAI is the newest and potentially the most significant long-term customer.</b> OpenAI is now a qualified Broadcom customer, expected to deploy more than 1 GW of compute capacity in FY27. While revenue contribution will be modest initially, the strategic significance is large: OpenAI's reported $10bn commitment to custom chip development with Broadcom represents a clear signal that the world's most prominent AI company is diversifying away from exclusive Nvidia reliance. Goldman Sachs emphasized this in its 4 March note, arguing that the customer disclosure should settle a key investor debate about Broadcom's addressable market.",
        "4. <b>Meta, SoftBank/Arm, and ByteDance round out the roster.</b> Meta's internal ASIC programs (Athena and Iris) remain on track with multiple gigawatts expected in FY27. HSBC's work estimates 900,000 Meta ASIC unit shipments this year, rising from 750,000 previously. Customers 5 and 6, widely believed to be SoftBank/Arm and ByteDance, are expected to more than double shipments into FY27. This widening customer base directly addresses the concentration concern that has historically weighed on the stock.",
        "5. <b>Networking may be the underappreciated upside driver.</b> AI networking accounted for one-third of Broadcom's AI revenue in Q1 and management guided it toward 40% in Q2. UBS, Goldman, Jefferies and HSBC all emphasize that this is not just a rounding change. Broadcom's Tomahawk 6 switch silicon, its merchant switching platform, its DSP (digital signal processor) leadership in 1.6T solutions, and its copper- and optical-based scale-up networking products are all benefiting from the same physics: as clusters get larger, the amount of data that must move between chips grows faster than the compute itself. CEO Hock Tan has estimated that for every $1 spent on AI chips, approximately $0.15-0.20 is spent on networking, and Broadcom captures a meaningful share of that spend. Jefferies projects networking to remain at 35-40% of AI revenue through CY27.",
    ]:
        story.append(Paragraph(line, styles["ERBullet"]))
    story.append(info_table(ROBOTICS_LINK, [1.35 * inch, 3.45 * inch, 2.1 * inch]))
    story.append(
        Paragraph(
            "The robotics angle deserves careful framing, because it is not the reason to own Broadcom today, but it is an important reason to believe the demand runway extends beyond the current LLM training cycle. Training large robotics foundation models — world models that simulate physical environments, multimodal perception stacks that fuse video, lidar, and tactile data, and reinforcement-learning agents that run billions of simulated episodes — requires the same kind of dense accelerator clusters used for frontier language-model training. In many cases, the data volumes are even larger because simulation-based robot training (think Nvidia Isaac Sim-like environments running at scale) generates enormous video and sensor-data streams that must be moved across the cluster at extremely high bandwidth.",
            styles["Body"],
        )
    )
    story.append(
        Paragraph(
            "This matters for Broadcom specifically because the networking and data-movement bottleneck becomes even more severe in robotics training than in text-centric LLM training. When the data flowing between accelerators is dense video rather than tokenized text, the bandwidth requirements per chip increase substantially, which favors Broadcom's networking silicon. In the broader AI/robotics value chain — from compute infrastructure through data, foundation models, simulation, analytics and fleet management, to application software — the compute and networking layers sit closest to near-term monetization because data-center spending must front-run the actual training work. Data itself is arguably the biggest moat in AI and robotics, but it is extremely hard to invest in directly in public markets: most pure-play data companies are private, and the data generated by vertical integrators like Tesla or Amazon is unlikely to be sold. Memory is important but may already be more fully priced in, given recent supply contracts and the difficulty of adding new memory assembly capacity beyond what the Street has modeled. That leaves networking and custom silicon as the portions of the infrastructure stack that are likely still underpriced relative to the demand they will face as robotics and simulation workloads join the existing LLM training wave.",
            styles["Body"],
        )
    )
    story.append(PageBreak())

    add_title("Guidance, Margins, And Capital Returns", "The quarter eased the two biggest practical concerns — margin dilution and supply-chain risk — while delivering aggressive shareholder returns")
    guide_table = info_table(
        [
            ["Q2 FY26 guide", "Management", "Street", "Read-through"],
            ["Revenue", "$22.0bn", "$20.5bn", "A meaningfully stronger top-line guide that was 7.3% above consensus, indicating broad demand acceleration"],
            ["AI revenue", "$10.7bn", "$9.3bn", "The largest guidance beat of any metric; networking and XPU programs are both driving the overshoot"],
            ["Non-AI semi revenue", "$4.1bn", "n.a.", "Roughly stable, no further deterioration; enterprise networking, storage and broadband all holding"],
            ["Software revenue", "$7.2bn", "n.a.", "Steady cash engine; sequential improvement from $6.8bn; VMware renewals appear healthy for now"],
            ["Adj. EBITDA margin", "68.0%", "67.1%", "Flat quarter-over-quarter at an elevated level, explicitly easing margin-compression fears"],
        ],
        [1.6 * inch, 1.15 * inch, 1.15 * inch, 3.2 * inch],
    )
    margin_commentary = [
        Paragraph("Broadcom did not merely guide revenue higher; it showed that the AI ramp can remain financially attractive as it scales. That distinction is central because, heading into the print, the single biggest investor concern was whether Broadcom's incremental AI business — particularly the Anthropic rack-level relationship and the broader shift toward full-system delivery — would dilute gross margins and drag earnings quality even as revenue accelerated. The margin fears were grounded in reasonable logic: when a chip company starts shipping more complete rack systems that include commodity memory and other pass-through components, gross margins typically compress because the revenue denominator grows faster than the margin contribution.", styles["Body"]),
        Paragraph("<b>Why margins matter — and why the fears have eased:</b> UBS's 5 March note provided perhaps the most useful framework for understanding the margin resolution. The firm noted that the mix of what Broadcom plans to ship to Anthropic is still evolving, but now appears likely to include more high-margin networking content and potentially a pivot to using ODM (original design manufacturer) partners for rack assembly. In other words, Broadcom may avoid absorbing the lowest-margin portions of the rack while still capturing the highest-value silicon inside it. Morgan Stanley independently reached a similar conclusion, noting that management was explicit that gross margins will not be diluted by racks, and that the prior overhang around how to model rack margins has been substantially removed. Goldman Sachs added that the company's constructive supply-chain commentary — having secured all components needed through FY28 — should help Broadcom maintain its current margin structure on AI products over the medium term. For a private banker, the practical implication is that Broadcom's margin profile looks resilient rather than fragile, which makes the stock significantly easier to recommend to clients who care about earnings quality as well as revenue growth.", styles["Body"]),
        Paragraph("<b>Supply chain secured through FY28 — an unusually strong visibility window:</b> One of the most important disclosures in the quarter was that Broadcom has fully secured capacity for leading-edge wafers, HBM (high-bandwidth memory), substrates, and advanced packaging (including 2.5D/3D stacking and CoWoS, the wafer-on-substrate technology used for advanced chip packaging) to support customer build-outs through FY28. This matters for two reasons. First, it removes the risk that component shortages or supplier margin demands erode Broadcom's profitability over the next two years. Second, it signals that management has enough forward order visibility to commit to multi-year supply agreements, which is a strong implicit endorsement of the demand outlook. Prior to this disclosure, investors feared that much of Broadcom's incremental margin would flow to suppliers of scarce components, particularly in HBM and advanced packaging where capacity is tight across the industry.", styles["Body"]),
        Paragraph("<b>Cash flow and capital returns make this story investable, not speculative:</b> Q1 FY26 operating cash flow was $8.3bn, against just $250mn of capital expenditure — an extraordinarily light capex intensity for a company at the center of the AI infrastructure build-out. Broadcom spent $7.9bn on share repurchases and $3.0bn on dividends in the quarter, then approved a fresh $10bn buyback authorization. The dividend yield is modest at roughly 0.7%, but the buyback program is aggressive and reflects management's confidence that the stock remains undervalued relative to forward earnings power. For private clients, this cash-generation profile is a genuinely differentiating feature: many AI-infrastructure names require investors to underwrite distant free cash flow, but Broadcom is already converting current earnings into immediate capital returns. The company is not asking shareholders to wait for a payoff; it is delivering one now.", styles["Body"]),
        Paragraph("<b>Software continues to do its job, but the cyclical test is ahead:</b> The infrastructure-software segment was only slightly light in Q1 at $6.8bn, but it remained stable and highly profitable with a 93% gross margin and 78% operating margin. The software franchise, which is primarily composed of VMware-related enterprise infrastructure products, functions as a cash-flow anchor for the broader Broadcom story. It provides recurring revenue and high-margin cash flow that can fund shareholder returns, service debt, and smooth the inherent cyclicality of semiconductor demand. The UBS software deep-dive, published on 23 February 2026, identified the main emerging risks: VMware customer churn when three-year deals begin to renew in 2026/2027, potential growth headwinds as the VCF (VMware Cloud Foundation) upsell is now being lapped, and the risk that AI coding tools accelerate cloud migration at the expense of on-prem modernization. We take these risks seriously in our downside scenario, but we do not think they break the investment case today. Software is a backstop, not a growth engine, and it only needs to stay stable — not accelerate — to do its job.", styles["Body"]),
    ]
    story.append(side_by_side([guide_table], image_flow(charts["cash"], 3.35 * inch, 2.35 * inch, "Figure 6. Capital return stayed aggressive in Q1 FY26.", styles)))
    story.append(Spacer(1, 0.03 * inch))
    for block in margin_commentary:
        story.append(block)
    story.append(PageBreak())

    add_title("Core Investment Thesis", "Three structural pillars — and one amplifier — explain why Broadcom still works for long-term client portfolios")
    thesis_blocks = [
        "<b>Thesis pillar 1 — Broadcom is evolving from component supplier to system-level AI infrastructure platform.</b> The most important conceptual shift in the Broadcom story is that the company is no longer selling individual chips into a supply chain it does not control. Instead, it is becoming a design partner for the largest AI infrastructure projects in the world. For each hyperscale customer, Broadcom provides the custom accelerator design (the XPU or ASIC itself), manages fabrication orchestration with TSMC across advanced process nodes (3nm, 5nm), handles advanced packaging (2.5D and 3D stacking, CoWoS), integrates its proprietary SerDes IP for ultra-high-speed chip-to-chip communication, and layers on its networking silicon portfolio — Tomahawk switches, DSPs, and optical interconnect — to connect thousands of those chips inside a cluster. This vertically integrated design-through-networking capability is extremely difficult to replicate and gives Broadcom multiple revenue streams per customer deployment. For a banker, the client message is that Broadcom is not betting on a single product; it is embedding itself in the architecture of AI data centers, which makes the revenue stream more durable and harder to displace than a merchant chip relationship would be.",

        "<b>Thesis pillar 2 — Earnings power is still rising faster than the stock narrative implies.</b> The broker pack spans a wide range of FY27 EPS views — from HSBC's $16.72 to UBS's $22.76 — but almost every note moved estimates up materially after the quarter. Our base-case FY27 adjusted EPS of $20.00 sits in the middle of the range and below the most bullish broker numbers, but it is comfortably above the market's old framing. What matters analytically is the trajectory: if Broadcom delivers $65bn+ of AI revenue in FY26 and anything close to the $100bn+ management framework in FY27 while margins hold near current levels, adjusted EPS should at least double from FY25 to FY27. That kind of earnings compounding can support further stock upside without requiring a speculative expansion in the valuation multiple. The risk, obviously, is that AI demand decelerates or that margins compress as the product mix shifts. But the Q1 FY26 results directly addressed both of those concerns, making the forward earnings trajectory more credible, not less. J.P. Morgan, in raising its price target to $500, specifically cited the combination of higher AI revenue visibility and better margin confidence as the drivers of its estimate revisions — a theme echoed by Goldman Sachs, Morgan Stanley, and Jefferies in their respective follow-ups.",

        "<b>Thesis pillar 3 — Capital returns and software cash flows create a genuine investability advantage.</b> Many AI-themed equities require investors to underwrite distant and uncertain free cash flow. Broadcom is different. The software franchise, even if it grows only in the low-to-mid single digits, provides roughly $27-34bn of annual recurring revenue with industry-leading margins, generating the cash flow needed to support a dividend, fund buybacks, and keep leverage under control. In Q1 FY26 alone, Broadcom generated $8.3bn of operating cash flow versus $250mn of capex, leaving more than $8bn available for shareholder returns and debt service. The $10bn buyback authorization, combined with the existing dividend, gives the stock a total capital-return profile that most AI infrastructure peers simply cannot match. For a private-banking client, this matters enormously: Broadcom can be held inside a diversified portfolio framework as a high-quality industrial grower with AI upside, rather than only as a speculative thematic bet. UBS's SOTP work supports this interpretation, estimating that even after adjusting for software risks and dis-synergies, Broadcom's semiconductor business alone is trading at only about 20x P/E, 23x EV/FCF, and 17x EV/EBITDA — multiples that are roughly one turn above Nvidia and the broader semiconductor peer set. That is not expensive for a business with this quality of growth and visibility.",

        "<b>Amplifier — Robotics and simulation workloads extend the demand runway beyond the current LLM cycle.</b> This is not a reason to buy the stock in isolation, but it is a reason to believe the infrastructure-spending cycle lasts longer than many investors currently model. Robotics foundation models, world models, and simulation-heavy training workloads are emerging as the next wave of compute-intensive AI applications, and they share a critical feature with LLM training: they require dense clusters of accelerators connected by extremely high-speed networking. The difference is that robotics training generates even larger data volumes — video, sensor fusion, physics simulation — which increases the networking bandwidth requirement per chip. As these workloads scale, they directly amplify demand for Broadcom's networking silicon and custom accelerators, creating an incremental demand layer beyond the hyperscaler LLM programs that currently dominate the company's order book. For every dollar spent on chips in this context, the $0.15-0.20 networking attach that CEO Hock Tan has referenced may actually increase rather than stay constant, because robotics data movement is more networking-intensive per unit of compute.",
    ]
    for block in thesis_blocks:
        story.append(Paragraph(block, styles["Body"]))
    story.append(Spacer(1, 0.05 * inch))
    story.append(info_table(THESIS_MATRIX, [1.45 * inch, 3.0 * inch, 2.45 * inch]))
    story.append(Paragraph("The coherence of the thesis matters as much as any individual data point. Customer breadth makes AI revenue more credible. Networking mix supports margins and revenue quality. Software cash flows stabilize earnings and fund capital returns. And capital returns make the stock easier to hold through the inevitable volatility that comes with AI-infrastructure exposure. These elements reinforce one another rather than pulling the story in different directions, which is why Broadcom can work as a serious portfolio position rather than just a trading vehicle. Even if the company lands below the highest broker estimates for FY27 AI revenue, the stock can still perform well if deployments prove broad, margins remain intact, and the cash-return profile continues to function as advertised.", styles["Body"]))
    story.append(PageBreak())

    add_title("Business Overview Beyond The AI Headline", "Private clients often know Broadcom as 'an AI stock' — they should understand the full business to make an informed allocation decision")
    story.append(
        side_by_side(
            [
                Paragraph("<b>Semiconductor solutions — the growth engine.</b> Broadcom's semiconductor business is where the AI story lives, but it is broader than AI alone. The segment includes: (1) custom AI accelerators (XPUs/ASICs) designed for hyperscale customers, which are the headline growth driver; (2) networking silicon including Tomahawk switches, Memory switches and Memory switch silicon, SerDes IP, DSPs, and optical interconnect components that connect chips and racks; (3) broadband and connectivity chips for cable modems, Wi-Fi access points, and set-top boxes; (4) server storage connectivity including SAS, RAID, and fibre-channel controllers; and (5) wireless chips sold primarily to Apple. In Q1 FY26, semiconductor solutions generated $12.5bn of revenue, of which $8.4bn was AI-related and $4.1bn was non-AI. The non-AI business is stable rather than exciting, but it matters for diversification: a broad cyclical recovery in enterprise networking, broadband or storage would be pure upside, not something required to justify the current stock price.", styles["Body"]),
                Paragraph("<b>Infrastructure software — the cash engine.</b> This segment is dominated by VMware-related enterprise infrastructure products, acquired through Broadcom's $69bn acquisition of VMware which closed in late 2023. It also includes CA Technologies, Symantec enterprise security, and Brocade software assets from prior acquisitions. The software franchise generated $6.8bn of revenue in Q1 FY26 with a 93% gross margin and 78% operating margin — numbers that would be the envy of virtually any standalone software company. Broadcom re-engineered VMware's go-to-market from perpetual licensing to subscription-based VMware Cloud Foundation (VCF) bundles, implemented substantial price increases (offering 'more for more' versus the old disaggregated product set), and cut costs dramatically. The result is industry-leading profitability, but also a business that now sits on a plateau: the easy upsell has been captured, three-year contracts will begin renewing in 2026/2027, and there is a legitimate question about whether AI coding tools may accelerate cloud migration at the expense of on-prem modernization, which would reduce VMware's addressable base over time.", styles["Body"]),
                Paragraph("<b>Why this dual structure matters for portfolio construction.</b> A client buying Broadcom is buying a company with one foot in structural AI growth and another in durable enterprise-infrastructure cash flow. That combination makes Broadcom more resilient in a drawdown than a pure-play AI semiconductor name, because the software business should still generate meaningful cash even if AI spending slows. Conversely, the AI semiconductor business provides the kind of earnings growth that a mature software platform cannot deliver on its own. UBS's SOTP analysis quantified this by noting that, even after adjusting for potential software dis-synergies in a hypothetical separation, Broadcom's semiconductor business trades at only about 20x P/E — roughly one turn above Nvidia. That is not an expensive multiple for a business with this growth trajectory and customer visibility.", styles["Body"]),
            ],
            image_flow(charts["annual_mix"], 3.35 * inch, 2.35 * inch, "Figure 7. Our annual mix view keeps software meaningful even as semis surge.", styles),
        )
    )
    story.append(
        side_by_side(
            image_flow(CROP_DIR / "jpm_crop_metrics.png", 3.35 * inch, 3.2 * inch, "Figure 8. Cropped JPMorgan chart / table view from the provided report pack.", styles),
            [
                Paragraph("<b>The software risk in detail — why it matters for downside support.</b> The UBS software deep dive published on 23 February 2026 is the most useful framework for understanding where the software backstop could weaken. The key risks are: (1) VMware customer churn when the first wave of three-year subscription contracts come up for renewal — some customers locked in at high prices during the forced transition and may not renew at the same level; (2) lapping of the VCF upsell, which means organic growth will need to come from new customers or expanded usage rather than migration-driven uplift; and (3) the possibility that AI-powered coding tools compress code and application modernization timelines, encouraging enterprises to move on-prem workloads to the public cloud faster than expected, which would shrink VMware's on-prem addressable market. UBS estimates that the software business can grow at a roughly 4-6% CAGR through 2029 even after accounting for these headwinds, which we view as a reasonable base case. The critical point is that software does not need to accelerate for the Broadcom story to work — it only needs to remain stable and profitable enough to fund capital returns and provide valuation support in a downturn.", styles["Body"]),
                Paragraph("<b>Non-AI semiconductors — stable, not dead.</b> The non-AI semiconductor business generated $4.1bn in Q1 FY26, roughly flat sequentially. Enterprise networking, broadband, and storage revenues were all up year-over-year, offset by a seasonal decline in wireless (which is driven primarily by Apple cycle timing). This part of the business is not a growth driver today, but it is also not a drag. Any broad recovery in enterprise IT spending — servers, storage, networking refresh — would provide incremental revenue and margin upside that is not currently embedded in most models. J.P. Morgan's full-take note specifically flagged this as a potential positive catalyst for FY27, noting that non-AI semiconductor growth should provide additional revenue and gross-margin upside as end-market recovery gains momentum.", styles["Body"]),
            ],
        )
    )
    story.append(PageBreak())

    add_title("Valuation And Price Target", "Our $500 target is built on a deliberately moderate framework — but the upside skew is real if execution continues")
    story.append(Paragraph("Valuing Broadcom today requires navigating a tension between two realities. On one hand, the stock has already rerated substantially from its 2024 levels and is trading at a premium to most semiconductor peers. On the other hand, the earnings trajectory is steeper than at any point in the company's history, and the Q1 FY26 results suggest that the forward estimates — which most investors have not fully updated — are more likely to move up than down. The framework below attempts to capture both of these dynamics by using conservative multiples applied to increasingly visible earnings power.", styles["Body"]))
    valuation_table = info_table(
        [
            ["Valuation scenario", "FY27E EPS", "P/E multiple", "Implied value", "Key assumptions"],
            ["Bear case", "$15.00", "24x", "$360", "AI demand digestion cycle begins in H2 FY27; software renewals weaken; 1-2 customer programs slip; multiple compresses toward historical average"],
            ["Base case", "$20.00", "25x", "$500", "Current AI ramp executes at roughly $110bn FY27 AI revenue; margins hold near 68% EBITDA; software grows low single digits; no new customer additions beyond current six"],
            ["Bull case", "$23.00", "25x", "$575", "FY27 AI revenue exceeds $130bn per UBS/HSBC frameworks; networking attach and revenue per GW both come in above base; OpenAI ramps faster than expected; FY28 growth visibility firms"],
        ],
        [1.1 * inch, 0.9 * inch, 0.9 * inch, 0.9 * inch, 3.2 * inch],
        highlight_col=3,
    )
    story.append(valuation_table)
    story.append(Spacer(1, 0.05 * inch))
    story.append(
        side_by_side(
            image_flow(charts["scenario"], 3.35 * inch, 2.35 * inch, "Figure 9. House valuation scenarios.", styles),
            image_flow(charts["broker_eps"], 3.35 * inch, 2.35 * inch, "Figure 10. FY27 EPS expectations across brokers.", styles),
        )
    )
    story.append(Paragraph("<b>Why we use 25x FY27E EPS as our base-case multiple.</b> This multiple sits below Broadcom's peak trading range but above where a mature hardware company would typically trade. The justification is that Broadcom is not a typical hardware company: it has a recurring-revenue software franchise worth $27-34bn annually, it has multi-year AI design wins with contractual visibility, and it generates free cash flow that supports aggressive buybacks. At the same time, 25x is not heroic — it does not require the market to award Broadcom a software-like premium. We think this balance is appropriate because the software segment's growth challenges (VMware renewal risk, VCF lapping) argue against stretching the blended multiple further, while the AI semiconductor segment's earnings trajectory argues against compressing it to commodity-hardware levels. For context, Nvidia currently trades at a modestly lower forward P/E despite carrying more customer concentration and more competition risk on the custom-silicon front.", styles["Body"]))
    story.append(Paragraph("<b>How the broker pack frames value.</b> The broker range is constructive but not uniform. HSBC is the most aggressive at $535, using a 32x multiple on its FY27 EPS estimate of $16.72 (which was set before the Q1 FY26 print and is likely to be revised higher). J.P. Morgan targets $500, broadly aligned with our framework. UBS maintains $475 despite raising estimates significantly, because it has chosen to apply a lower multiple — roughly 21x CY27 EPS — to discount concerns about the sustainability of AI spending levels. Morgan Stanley sits at $470, below our target, but its risk/reward analysis still shows meaningful upside from the current share price. Goldman Sachs maintains a Buy with conviction-list placement but has not published a specific target revision post-quarter. Jefferies, while not publishing a formal target in the same format, has noted that the math supports $25-30 of EPS in CY27 on the higher end of AI revenue scenarios, which at 25x would imply $625-750 per share — well above our base case.", styles["Body"]))
    story.append(Paragraph("<b>The SOTP perspective adds confidence.</b> UBS's sum-of-the-parts work, published on 23 February 2026, provides a useful cross-check. Stripping out the software business at a conservative standalone multiple (which UBS estimates with room for dis-synergies) and valuing the semiconductor business separately, UBS arrives at an implied semiconductor P/E of roughly 20x — which is only one turn above Nvidia. That finding is important because it suggests that even if the software business disappoints, the semiconductor franchise is not being valued at a bubble premium. In fact, the semiconductor business may be undervalued relative to its growth profile if the AI revenue trajectory executes anywhere near management's framework.", styles["Body"]))
    story.append(Paragraph("<b>Why we are not more aggressive.</b> Our target deliberately leaves room for execution to create upside rather than baking in the most optimistic scenarios. The stock is already carrying meaningful AI expectations, and the software segment should probably receive a more sober multiple than it earned during peak VMware enthusiasm in 2024. We would rather preserve credibility and let performance create the rerating than stretch for a target that requires everything to go right. If networking attach exceeds 40% of AI revenue on a sustained basis, or if FY27 AI revenue tracks toward the upper end of the $120-130bn range that UBS and HSBC have modeled, we would revisit the target upward — but only once the data supports it.", styles["Body"]))
    story.append(RLImage(str(CROP_DIR / "jeff_crop_charts.png"), width=6.8 * inch, height=2.15 * inch))
    story.append(Paragraph("Figure 11. Jefferies chart page from the attached report pack, showing historical surprise, forward P/E, margins and short interest (4 Mar 2026).", styles["CenterNote"]))
    story.append(PageBreak())

    add_title("Key Risks And Near-Term Catalysts", "The stock can still work well, but the debate is now about durability and execution — not discovery")
    risks_table = info_table(
        [
            ["Risk", "Why it matters", "What to monitor"],
            ["Hyperscaler capex normalization", "Broadcom is now judged on multi-year AI deployment duration. If hyperscaler capex budgets flatten or decline in FY28, the stock will de-rate regardless of Q1/Q2 strength.", "Cloud capex budgets, AI cluster build timing, customer commentary on FY28 plans."],
            ["Customer insourcing / COT", "Customer-owned-tooling risk means Google, Meta or others could internalize more chip design, shrinking Broadcom's content per program over time.", "Management commentary on COT, design-win cadence, share of next-gen programs."],
            ["Gross-margin dilution", "Rack-level or pass-through content can lower earnings conversion if the product mix shifts unfavorably toward lower-margin system shipments.", "Networking share of AI revenue, incremental gross margin trend, Anthropic shipment mix."],
            ["Software renewal churn", "VMware three-year contracts renew starting 2026/2027; weaker renewals would reduce the valuation backstop in a downturn.", "VMware renewal rates, VCF adoption metrics, enterprise spending sentiment."],
            ["Supply chain / advanced packaging", "HBM, CoWoS packaging and substrate availability can cap upside or shift revenue timing even with secured commitments.", "TSMC CoWoS capacity, HBM supplier commentary, lead times, booking patterns."],
            ["Valuation risk", "At 25x FY27E EPS, the stock is not cheap in absolute terms. Any negative revision to the earnings trajectory could trigger a sharp de-rate.", "Forward P/E relative to semiconductor peers, earnings revision momentum, sector rotation."],
        ],
        [1.45 * inch, 3.35 * inch, 2.05 * inch],
    )
    story.append(risks_table)
    story.append(Spacer(1, 0.04 * inch))
    story.append(Paragraph("<b>Risk deep-dive: hyperscaler capex normalization.</b> The most important risk for Broadcom is also the most difficult to model: the possibility that hyperscaler capital-expenditure growth slows before Broadcom's AI revenue framework fully materializes. The current AI spending cycle has been driven by competitive urgency — every large cloud provider feels compelled to invest in AI infrastructure to avoid falling behind peers. That dynamic has pushed capex budgets to record levels. But capex cycles are cyclical by nature, and if the major hyperscalers begin to see diminishing returns on their AI investments, or if macroeconomic conditions force budget discipline, spending could flatten or decline before Broadcom reaches its FY27 revenue targets. Broadcom's stock is now pricing duration — multi-year sustained demand — rather than the initial surprise that drove the first leg of the rally. That makes it more vulnerable to any signal that duration is shorter than expected.", styles["Body"]))
    story.append(Paragraph("<b>Risk deep-dive: customer insourcing and COT.</b> Customer-owned-tooling is the semiconductor industry's term for when a customer brings chip-design capability in-house rather than relying on a third-party design partner like Broadcom. Google has its own TPU design team, Meta has its MTIA team, and Amazon has Annapurna Labs. The question is whether these internal teams will progressively take over more of the design work that Broadcom currently performs, reducing Broadcom's content per customer over time. Management addressed this directly on the Q1 FY26 call, stating that it does not expect any major competitive dynamics from COT for at least several more years. J.P. Morgan supported this view, noting that Broadcom's full-stack integration — from advanced-node fabrication orchestration through packaging and SerDes IP — creates high switching costs. But the risk is real on a 3-5 year horizon and should be monitored.", styles["Body"]))
    catalyst_table = info_table(
        [
            ["Catalyst", "Timeline", "Why clients should care"],
            ["Next quarterly earnings (Q2 FY26)", "June 2026", "Confirms whether AI revenue reaches the $10.7bn guide and whether margins hold at 68% EBITDA; the single most important near-term data point."],
            ["Customer deployment disclosures", "Ongoing", "Concrete evidence that OpenAI, Meta and Anthropic are ramping on schedule reduces concentration risk and supports the $100bn+ FY27 framework."],
            ["Networking mix trajectory", "Q2 FY26 onward", "If networking sustains at or above 40% of AI revenue, it signals that Broadcom's attach rate is structural rather than a one-quarter anomaly."],
            ["Buyback execution", "Ongoing", "Consistent repurchases at current levels soften drawdowns and improve per-share earnings power for client portfolios."],
            ["Non-AI recovery signals", "H2 FY26", "Any cyclical uptick in enterprise networking, storage or broadband would be pure upside not embedded in current estimates."],
            ["FY27 framework updates", "Next 2-3 quarters", "Management commentary on GW deployment pace, revenue per GW, and FY28 growth expectations will shape the medium-term multiple."],
        ],
        [2.15 * inch, 0.85 * inch, 3.9 * inch],
    )
    story.append(catalyst_table)
    story.append(Spacer(1, 0.04 * inch))
    story.append(side_by_side(
        [Paragraph("<b>What would make us more bullish:</b> Evidence that networking mix and customer breadth continue to expand without margin leakage; OpenAI ramp accelerates beyond 1 GW; Anthropic margins prove better than feared; software renewals stabilize and churn concerns prove manageable; non-AI semiconductor recovery begins; FY28 framework becomes visible. Any combination of these would justify a move toward the bull-case valuation.", styles["Body"]),
         Paragraph("<b>What would make us less bullish:</b> Hyperscaler capex budgets begin to flatten; AI revenue guide disappoints or does not accelerate sequentially; margins compress due to rack mix or supply-chain cost absorption; software renewals deteriorate meaningfully; buybacks slow without a clear reinvestment rationale; a major customer program slips or is delayed. Any of these would warrant re-examining the position and potentially moving toward the bear-case framing.", styles["Body"]),
        ],
        image_flow(CROP_DIR / "ms_crop_risk_reward.png", 3.35 * inch, 3.0 * inch, "Figure 12. Cropped Morgan Stanley risk / reward exhibit from the source pack.", styles)))
    story.append(PageBreak())

    add_title("Street Read-Through And Broker Debate", "The broker pack is broadly constructive — but the real value is in understanding where the analysts disagree")
    story.append(info_table(BROKER_GRID, [1.35 * inch, 0.8 * inch, 4.85 * inch]))
    story.append(
        Paragraph(
            "The broad message across the broker pack is bullish, but the notes are not identical, and the differences reveal where the real analytical debates sit. Understanding those debates is more useful for client conversations than simply knowing that every broker has a Buy rating.",
            styles["Body"],
        )
    )
    story.append(Paragraph("<b>J.P. Morgan ($500 target)</b> is the most comprehensive in its customer-by-customer framework. The 5 March full-take note anchors on the $100bn+ FY27 AI revenue disclosure and pushes it to $120bn+ by applying $12-15bn of revenue per deployed gigawatt. JPM's Harlan Sur was among the first to identify OpenAI as a qualified customer and to quantify Anthropic's 3 GW deployment plan. The note also emphasized the inventory build — $2.96bn versus consensus of $2.64bn — as a bullish signal of forward demand confidence rather than a bearish warning of slowing sell-through. JPM's estimate revisions were significant: FY26E EPS moved from $10.92 to $12.29, a 12.5% increase, and the price target rose from $475 to $500. For our purposes, JPM provides the most detailed factual base for the quarter, the AI revenue path, customer ramps, and the FY26-FY27 EPS framing.", styles["Body"]))
    story.append(Paragraph("<b>UBS ($475 target)</b> took a different angle by focusing on margin resilience rather than the headline AI revenue number. UBS analyst Timothy Arcuri noted that the $100bn FY27 AI revenue figure was actually roughly in line with his pre-call model of $106bn and with most pre-call investor expectations — the more incremental revelation was on margins. UBS's supply-chain work now suggests TPU unit shipments may reach approximately 6 million next year, and management's GW-based commentary implies revenue per GW similar to AMD's ~$15bn. That math drives UBS's FY27 AI revenue forecast above $130bn. Despite these much higher estimates, UBS left its $475 target unchanged because the implied multiples — roughly 21x CY27 P/E and 17x CY28 — already discount concerns about AI spending sustainability. UBS is also the most explicit on software risks, having published a companion note with its software team that flags VMware renewal churn, VCF lapping, and AI-driven cloud migration as headwinds.", styles["Body"]))
    story.append(Paragraph("<b>Morgan Stanley ($470 target)</b> framed the quarter as strengthening the thesis with modest upside impact. Joseph Moore's team emphasized two comfort factors: (1) rack-margin fears have substantially receded because management was explicit that gross margins will not be diluted, and (2) customer-owned-tooling commentary was constructive, with Broadcom continuing to expect majority share in the TPU ecosystem. MS is currently modeling approximately $120bn of AI revenue in FY27 and sees potential for upward revisions. The note uses a per-GW revenue framework of roughly $20bn as a bull case, which yields very large numbers but which MS views as closer to an upside scenario than a base case. MS's contribution to the debate is its risk/reward analysis, which shows attractive upside from the current price even using a below-consensus multiple.", styles["Body"]))
    story.append(info_table(DEBATE_MATRIX, [1.75 * inch, 2.55 * inch, 2.7 * inch]))
    story.append(Paragraph("<b>Goldman Sachs (Buy, Conviction List)</b> published both a first-take reaction and a follow-up note. Goldman's James Schneider described the quarter as providing guidance that should address several key investor debates and drive the stock meaningfully higher. Goldman was particularly focused on the supply-chain assurance — components secured through FY28 — and the margin clarity, arguing that these disclosures are significant positives that address key investor concerns. Goldman's revised AI semiconductor revenue estimates now stand at $60bn/$130bn/$170bn for FY26/27/28, representing one of the more aggressive trajectories in the pack. The note also emphasized that Broadcom's leadership in AI networking and custom silicon enables the lowest inference cost for its hyperscaler customers, a key competitive differentiator.", styles["Body"]))
    story.append(Paragraph("<b>Jefferies (Franchise Pick / Top Pick)</b> offered the most aggressive upside framing. The note by Mark Lipacis argues that the $100bn FY27 figure likely proves conservative, based on a revenue-per-GW analysis that shows $13bn/GW in CY25 rising to $18bn in CY26, with potential for $20-25bn/GW in CY27. At the high end, that math implies AI revenue could reach $200bn — far above any consensus estimate. Jefferies also highlighted that networking is outgrowing ASICs within the AI mix, with dominance in 1.6T DSPs and Tomahawk 6 (with Tomahawk 7 at 200T coming next year) driving the networking attach rate higher. The note framed the debate as one about capex sustainability into CY28 rather than Broadcom's competitive positioning, concluding that it is difficult to see how these compute names will remain this cheap.", styles["Body"]))
    story.append(Paragraph("<b>HSBC ($535 target — highest in the pack)</b> was set before the Q1 FY26 print based on its 24 November 2025 initiation note, and is likely to be revised higher. HSBC argued that the Street was still underestimating the ASIC opportunity even after substantial upward revisions, and that Broadcom had increased its FY26 CoWoS capacity allocation from 150,000 to 250,000 wafers. HSBC's FY26/FY27 ASIC revenue estimates of $44.7bn/$72.0bn were 44%/59% above consensus at the time of publication. The note also emphasized networking TAM expansion as a further upside driver, with the overall AI addressable market expanding beyond traditional hyperscalers into emerging hyperscalers, enterprise AI, and sovereign AI deployments.", styles["Body"]))
    story.append(
        side_by_side(
            image_flow(CROP_DIR / "ubs_crop_tpu.png", 3.35 * inch, 2.05 * inch, "Figure 13. UBS TPU unit estimates table (5 Mar 2026).", styles),
            image_flow(CROP_DIR / "hsbc_crop_revisions.png", 3.35 * inch, 2.05 * inch, "Figure 14. HSBC consensus revision exhibits (24 Nov 2025 note).", styles),
        )
    )
    story.append(PageBreak())

    add_title("Detailed Estimates And Portfolio Framing", "Appendix tables for bankers who want the numbers, the client framing, and the practical monitoring framework")
    story.append(info_table(ESTIMATE_MATRIX, [2.0 * inch, 1.25 * inch, 1.25 * inch, 1.25 * inch], highlight_col=2))
    story.append(Spacer(1, 0.04 * inch))
    story.append(
        Paragraph(
            "<b>House view summary:</b> Broadcom remains one of the best large-cap ways to express AI infrastructure demand without sacrificing balance-sheet strength and cash-flow visibility. We rate the shares Outperform with a $500 twelve-month price target, implying approximately 57% upside from the 4 March closing price of $317.53. The stock is appropriate for growth-oriented clients who can tolerate periodic AI-capex-driven volatility and who want their AI exposure backed by real free cash flow, a meaningful dividend, and aggressive buybacks. Our FY26 revenue estimate of $92.5bn reflects approximately 45% year-over-year growth driven by AI semiconductor acceleration, while our FY27 estimate of $121.0bn assumes AI revenue of approximately $110bn — conservatively within the range established by management's $100bn+ floor and the broker pack's upper bounds of $130bn+.",
            styles["Body"],
        )
    )
    story.append(info_table(CLIENT_FIT, [1.8 * inch, 5.2 * inch]))
    story.append(Spacer(1, 0.03 * inch))
    story.append(info_table(MONITORING_FRAME, [1.55 * inch, 3.1 * inch, 2.35 * inch]))
    story.append(Spacer(1, 0.03 * inch))
    story.append(Paragraph("<b>How to use this stock in client portfolios.</b> Broadcom is best positioned as a core AI/technology holding within a diversified equity allocation, rather than as a satellite or thematic trade. The combination of high-single-digit free-cash-flow yield (on a forward basis), aggressive capital returns, and structural earnings growth gives it a quality-growth profile that can sit alongside other large-cap industrial or technology holdings without creating undue concentration risk. For clients who already own Nvidia, Broadcom offers complementary exposure to the custom-silicon and networking segments of the AI stack, diversifying the AI bet away from a single merchant GPU vendor. For clients who are underweight semiconductors, Broadcom is a cleaner entry point than many alternatives because the software cash-flow base reduces downside risk relative to pure-play chip companies.", styles["Body"]))
    story.append(Paragraph("<b>Position sizing considerations.</b> Given that Broadcom is a mega-cap stock ($1.6 trillion market cap) with high daily liquidity (average daily trading value of approximately $8bn), it is suitable for most portfolio sizes. However, the stock carries AI-cycle sensitivity and will remain volatile around earnings, hyperscaler capex announcements, and any shifts in AI sentiment. We would suggest a position size consistent with a high-conviction large-cap holding — typically 2-5% of a growth-oriented equity portfolio — with room to add on any pullbacks driven by macro or sentiment rather than fundamental deterioration.", styles["Body"]))
    story.append(Paragraph("<b>What would trigger a rating change.</b> We would consider downgrading to Market Perform if: (1) AI revenue guidance fails to meet or exceed the $10.7bn Q2 guide, suggesting demand is plateauing rather than accelerating; (2) gross margins compress more than 200bps on a sustained basis, indicating that the rack-mix and pass-through concerns have returned; (3) software revenue declines sequentially for two consecutive quarters, signaling that the VMware renewal cycle is worse than expected; or (4) hyperscaler capex budgets show clear signs of flattening across multiple customers. Conversely, we would consider raising our target above $500 if: (1) FY27 AI revenue tracking exceeds $120bn; (2) networking consistently maintains 40%+ share of AI revenue; (3) OpenAI or additional customers ramp faster than current expectations; or (4) the non-AI semiconductor cycle turns positive, providing incremental upside not currently in estimates.", styles["Body"]))
    story.append(
        Paragraph(
            "If this report is being used in a client conversation, the cleanest summary is that Broadcom offers a combination many AI beneficiaries do not: it participates in a very large infrastructure buildout, it monetizes both compute and networking bottlenecks, it has a software cash-flow base that most chip companies lack, and it is already returning capital to shareholders at a rate that makes the stock investable for mainstream portfolios rather than only for high-conviction thematic accounts. The emerging robotics and simulation demand wave — while not the primary investment rationale — adds a further layer of confidence that the infrastructure-spending cycle has more duration than the market currently models.",
            styles["Body"],
        )
    )

    doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(8.5)


def add_docx_table(document: Document, rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = Inches(width)
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            set_cell_text(table.cell(r, c), value, bold=(r == 0))
            if r == 0:
                shade_cell(table.cell(r, c), "10324A")
        if r == 0:
            for cell in table.rows[r].cells:
                for run in cell.paragraphs[0].runs:
                    run.font.color.theme_color = None
    document.add_paragraph()


def set_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)


def add_header(section, title: str) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Broadcom Inc. - Equity Research Update | {title}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)


def add_docx_page(document: Document, header_title: str) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    set_page(section)
    add_header(section, header_title)


def make_docx(charts: dict[str, Path]) -> None:
    document = Document()
    set_page(document.sections[0])
    add_header(document.sections[0], "Summary")
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(8.7)

    def title(text: str, subtitle: str) -> None:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(19)
        p2 = document.add_paragraph()
        r2 = p2.add_run(subtitle)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(10)

    title(TITLE, f"{SUBTITLE} | {REPORT_DATE}")
    document.add_paragraph("Broadcom helps build the hidden infrastructure that allows large AI systems to train and run efficiently. It helps hyperscale customers design custom accelerator chips (XPUs/ASICs), sells the networking hardware and silicon that move data between those chips at very high speed, and owns a large enterprise-software business through VMware that provides steady, recurring cash flow. The January-quarter results and April-quarter guidance changed the conversation from quarterly optionality to multi-year visibility, with management disclosing line of sight to more than $100bn of AI semiconductor revenue in FY27 across roughly 10 GW of deployed compute for six customers.")
    add_docx_table(
        document,
        [
            ["Key metric", "Value", "Commentary"],
            ["Recommendation", "Outperform", "AI custom silicon plus networking remain under-modeled; software cash flow anchors downside"],
            ["Current price", f"${CURRENT_PRICE:.2f}", f"As of {PRICE_DATE}"],
            ["Target price", f"${TARGET_PRICE:.0f}", f"{UPSIDE:.1%} implied upside from base-case 25x FY27E EPS of $20.00"],
            ["Market cap", f"${MARKET_CAP_B:,.0f}bn", "Mega-cap AI infrastructure platform with mature cash generation"],
            ["Enterprise value", f"${ENTERPRISE_VALUE_B:,.0f}bn", "Backed by $8.3bn quarterly operating cash flow and $10bn fresh buyback authorization"],
        ],
        [1.7, 1.2, 4.2],
    )
    document.add_picture(str(charts["targets"]), width=Inches(6.9))
    document.add_paragraph("Key takeaways: (1) The quarter itself was solid but the guidance was transformational — Q2 AI revenue guided to $10.7bn, 15% above consensus; (2) Broadcom is the scaled alternative to merchant GPU economics, combining custom chip design with networking silicon; (3) Margin fears around rack-level shipments were substantially eased, with management explicit that gross margins will not be diluted; (4) The $10bn buyback, $8.3bn quarterly cash flow, and secured components through FY28 make this an AI story that pays its way.")

    add_docx_page(document, "Latest earnings")
    document.add_heading("Latest Earnings Update", level=1)
    document.add_paragraph("Q1 FY26 revenue of $19.3bn and adjusted EPS of $2.05 both edged past consensus, but the real inflection came from guidance. Management pointed to Q2 revenue of $22.0bn and AI revenue of $10.7bn — 7.3% and 15.1% above consensus respectively. AI revenue was 43% of total company revenue, up from 27% a year ago. Operating margin hit 66.4% versus the Street's 65.5%, and EBITDA margin came in 170bps better at 68.0%. The combination of stronger AI revenue, a materially better Q2 guide, and steady margin guidance tells investors that Broadcom is scaling into demand rather than chasing it.")
    document.add_paragraph("One underappreciated signal was the inventory build: $2.96bn versus consensus of $2.64bn. At the start of a major capacity ramp with sharply higher guided revenue, this is most naturally interpreted as confidence in forward demand. The company also disclosed that it has secured all critical components — leading-edge wafers, HBM, substrates, and advanced packaging — through FY28.")
    add_docx_table(
        document,
        [
            ["Metric", "Actual", "Street", "Comment"],
            ["Revenue", "$19.3bn", "$19.2bn", "Beat driven by semiconductor solutions"],
            ["AI revenue", "$8.4bn", "$8.2bn", "Custom ASIC plus networking upside; AI compute +28% QoQ, networking +34% QoQ"],
            ["Gross margin", "77.0%", "76.8%", "No rack-related dilution visible"],
            ["Operating margin", "66.4%", "65.5%", "+90bps; opex at $2.0bn below $2.2bn consensus"],
            ["Adj. EPS", "$2.05", "$2.03", "Better quality than size"],
            ["Q2 guide revenue", "$22.0bn", "$20.5bn", "+7.3% above consensus"],
            ["Q2 AI guide", "$10.7bn", "$9.3bn", "+15.1% above consensus"],
        ],
        [1.7, 1.1, 1.1, 3.1],
    )
    document.add_picture(str(charts["rev_ai"]), width=Inches(6.8))
    document.add_picture(str(charts["q1_mix"]), width=Inches(6.8))
    document.add_paragraph("AI networking revenue grew roughly 34% sequentially while AI compute grew 28%, suggesting the networking attach rate is increasing as cluster sizes expand. This confirms CEO Hock Tan's framework that for every $1 spent on AI chips, approximately $0.15-0.20 is spent on networking.")

    add_docx_page(document, "AI engine")
    document.add_heading("AI Engine And Customer Roadmap", level=1)
    for para in [
        "Broadcom designs custom accelerator chips (XPUs/ASICs) for hyperscale customers who want silicon tailored to their specific AI workloads — chips that do not carry Nvidia's margins or general-purpose overhead. It also designs the high-speed networking silicon that moves data between those chips. The combination makes Broadcom a system-level infrastructure partner.",
        "Management disclosed line of sight to more than $100bn of FY27 AI semiconductor revenue across roughly 10 GW of deployed compute for six XPU customers. J.P. Morgan pushes that framework to $120bn+, UBS to above $130bn based on supply-chain TPU unit work suggesting ~6 million units next year, and Jefferies argues the upper bound could be far higher if revenue per GW rises from $13bn in CY25 to $20-25bn in CY27.",
        "Google remains the anchor TPU customer (Ironwood + TPU v8 Sunfish). Anthropic is scaling to 3 GW in FY27 (3x YoY). OpenAI is newly qualified with 1+ GW expected. Meta's ASIC programs (Athena/Iris) remain on track. Customers 5-6 (likely SoftBank/Arm and ByteDance) are expected to more than double shipments. This customer breadth directly addresses the concentration risk that historically weighed on the stock.",
        "Networking may be the underappreciated upside driver. AI networking rose from one-third of AI revenue in Q1 to a guided 40% in Q2, driven by Tomahawk 6, 1.6T DSPs, and higher attach rates. Jefferies projects networking to remain at 35-40% through CY27. CEO Hock Tan has estimated that for every $1 spent on AI chips, $0.15-0.20 is spent on networking.",
    ]:
        document.add_paragraph(para)
    document.add_picture(str(charts["ai_mix"]), width=Inches(6.8))
    document.add_picture(str(charts["ai_frames"]), width=Inches(6.8))
    add_docx_table(document, ROBOTICS_LINK, [1.5, 3.2, 1.8])
    document.add_paragraph("Robotics foundation models — world models, sensor fusion, simulation-heavy RL — require the same dense accelerator clusters as frontier LLM training, but with even heavier data movement due to video and sensor volumes. As these workloads scale, they amplify demand for networking and compute bottlenecks that Broadcom monetizes. Data is the biggest moat in AI/robotics but is hard to invest in directly. Memory is important but may be more fully priced in. Networking and custom silicon are likely still underpriced relative to the demand they will face.")
    document.add_paragraph("This is one reason the Broadcom story is broader than a single-quarter chip narrative. It sits close to the near-term infrastructure bottlenecks where spending must front-run actual monetization, meaning value accrues at the chip and networking layer earlier than at application-software or data layers that remain largely private.")

    add_docx_page(document, "Margins and cash")
    document.add_heading("Guidance, Margins, And Capital Returns", level=1)
    add_docx_table(
        document,
        [
            ["Q2 guide", "Management", "Street", "Implication"],
            ["Revenue", "$22.0bn", "$20.5bn", "+7.3% above consensus; demand accelerating across customers"],
            ["AI revenue", "$10.7bn", "$9.3bn", "+15.1% above consensus; networking and XPU both driving overshoot"],
            ["Adj. EBITDA margin", "68.0%", "67.1%", "Flat QoQ at elevated level; margin compression fears easing"],
            ["Software revenue", "$7.2bn", "n.a.", "Cash engine remains stable; sequential improvement from $6.8bn"],
        ],
        [1.8, 1.1, 1.1, 2.8],
    )
    document.add_picture(str(charts["cash"]), width=Inches(6.8))
    document.add_paragraph("Broadcom generated $8.3bn of operating cash flow in Q1 FY26, spent $7.9bn on buybacks and $3.0bn on dividends, and approved a new $10bn repurchase authorization. Capex was only $250mn. The company stated it has secured all critical components — leading-edge wafers, HBM, substrates, advanced packaging — through FY28, removing the risk that supplier margin demands erode Broadcom's profitability.")
    document.add_paragraph("UBS provided the most useful margin framework: the mix shipped to Anthropic may now include more high-margin networking and potentially pivot to ODM partners for rack assembly, meaning Broadcom avoids the lowest-margin portions while capturing highest-value silicon. Morgan Stanley independently confirmed management was explicit that gross margins will not be diluted by racks.")
    document.add_paragraph("For bankers, this is one of the features that makes Broadcom easier to recommend than more speculative AI names. The company is already converting strength into capital returns — it is not asking shareholders to wait.")

    add_docx_page(document, "Investment thesis")
    document.add_heading("Core Investment Thesis", level=1)
    for para in [
        "Pillar 1 — Broadcom is evolving from component supplier to system-level AI infrastructure platform. It provides custom accelerator design, fabrication orchestration with TSMC, advanced packaging, proprietary SerDes IP, and networking silicon. This vertically integrated capability gives it multiple revenue streams per deployment and creates high switching costs. UBS's SOTP work shows the semiconductor business alone trades at roughly 20x P/E — only one turn above Nvidia.",
        "Pillar 2 — Earnings power is still rising faster than the stock narrative implies. The broker pack spans FY27 EPS from HSBC's $16.72 to UBS's $22.76. Our base case of $20.00 implies earnings roughly doubling from FY25 to FY27. That compounding can support further stock upside without requiring a speculative multiple expansion.",
        "Pillar 3 — Software cash flows and capital returns make the stock investable, not speculative. The software franchise provides $27-34bn of annual recurring revenue with industry-leading margins. The $10bn buyback and quarterly $8.3bn cash flow make this an AI story that pays its way. This fits clients who want AI exposure with mature free cash flow and shareholder returns.",
        "Amplifier — Robotics and simulation workloads extend the demand runway. Training world models and simulation-heavy RL generates even larger data volumes than LLM training, amplifying networking demand. This is not the reason to buy today, but it adds confidence that the infrastructure-spending cycle has more duration than currently modeled.",
    ]:
        document.add_paragraph(para)
    add_docx_table(document, THESIS_MATRIX, [1.6, 2.8, 2.3])
    document.add_paragraph("The coherence of the thesis matters more than any individual estimate. Customer breadth, networking mix, software resilience and capital return discipline all reinforce one another rather than pulling the story in different directions.")

    add_docx_page(document, "Business overview")
    document.add_heading("Business Overview Beyond The AI Headline", level=1)
    document.add_paragraph("Broadcom's semiconductor business includes: (1) custom AI accelerators (XPUs/ASICs) for hyperscale customers; (2) networking silicon including Tomahawk switches, SerDes IP, DSPs, and optical interconnect; (3) broadband and connectivity chips; (4) server storage connectivity; and (5) wireless chips sold primarily to Apple. In Q1 FY26, semiconductor solutions generated $12.5bn, of which $8.4bn was AI-related and $4.1bn non-AI.")
    document.add_paragraph("Infrastructure software is dominated by VMware-related products with 93% gross margin and 78% operating margin. Broadcom re-engineered VMware from perpetual licensing to VCF subscription bundles with substantial price increases. The result is industry-leading profitability, but risks ahead include VMware renewal churn, VCF upsell lapping, and AI-driven cloud migration. UBS estimates software can grow at 4-6% CAGR through 2029 even after accounting for these headwinds.")
    document.add_paragraph("This dual structure matters for portfolio construction: clients get structural AI growth from semiconductors plus durable cash flow from software, making Broadcom more resilient than a pure-play AI semiconductor name.")
    document.add_picture(str(charts["annual_mix"]), width=Inches(6.8))
    document.add_picture(str(CROP_DIR / "jpm_crop_metrics.png"), width=Inches(6.8))
    document.add_paragraph("Software is not the headline growth engine, but it remains a meaningful valuation backstop. The main risk is whether VMware renewals and VCF upsell stay healthy enough to preserve that role.")

    add_docx_page(document, "Valuation")
    document.add_heading("Valuation And Price Target", level=1)
    document.add_paragraph("Our $500 target is built on 25x FY27 adjusted EPS of $20.00 — deliberately moderate to preserve credibility and let execution create the rerating. The 25x multiple sits below Broadcom's peak trading range but above mature-hardware levels, reflecting the recurring software franchise, multi-year AI design wins, and aggressive buybacks.")
    add_docx_table(document, [["Scenario", "FY27E EPS", "P/E", "Value", "Key assumptions"], ["Bear", "$15.00", "24x", "$360", "AI demand digestion; software renewals weaken; 1-2 customer programs slip"], ["Base", "$20.00", "25x", "$500", "Current AI ramp executes at ~$110bn FY27; margins hold; software grows low single digits"], ["Bull", "$23.00", "25x", "$575", "FY27 AI revenue exceeds $130bn; networking attach above base; OpenAI ramps faster"]], [1.2, 1.0, 0.8, 0.8, 3.1])
    document.add_picture(str(charts["scenario"]), width=Inches(6.8))
    document.add_picture(str(charts["broker_eps"]), width=Inches(6.8))
    document.add_picture(str(CROP_DIR / "jeff_crop_charts.png"), width=Inches(6.8))
    document.add_paragraph("The broker range spans $470 (Morgan Stanley) to $535 (HSBC). UBS left its $475 target unchanged despite much higher estimates, applying a ~21x CY27 P/E to discount AI sustainability concerns. HSBC is most aggressive, using 32x on its FY27 EPS. Our target sits in the middle, aligned with JPM's $500.")
    document.add_paragraph("UBS's SOTP analysis provides a useful cross-check: stripping out software at a conservative standalone multiple, the semiconductor business trades at only ~20x P/E — one turn above Nvidia. That suggests the semiconductor franchise is not at a bubble premium.")

    add_docx_page(document, "Risks and catalysts")
    document.add_heading("Key Risks And Catalysts", level=1)
    add_docx_table(
        document,
        [
            ["Risk", "Why it matters"],
            ["Hyperscaler capex normalization", "Stock is pricing multi-year AI deployment duration; if capex flattens in FY28, the stock de-rates regardless of near-term strength."],
            ["Customer insourcing / COT", "Google, Meta, Amazon all have internal chip teams. Management expects no major impact for several years, but this is a real 3-5 year risk."],
            ["Gross-margin dilution", "Rack mix or pass-through content could weaken EPS conversion; eased by Q1 commentary but needs monitoring."],
            ["Software renewal churn", "VMware 3-year contracts renew in 2026/2027; weaker renewals reduce valuation backstop in a downturn."],
            ["Supply chain / packaging", "HBM, CoWoS, and substrate availability can cap upside or shift timing even with secured commitments."],
            ["Valuation risk", "At 25x FY27E EPS, any negative earnings revision triggers sharp de-rate. The stock is pricing duration, not surprise."],
        ],
        [2.2, 4.6],
    )
    document.add_picture(str(CROP_DIR / "ms_crop_risk_reward.png"), width=Inches(6.8))
    document.add_paragraph("Near-term catalysts to watch: (1) Q2 FY26 earnings confirming $10.7bn AI revenue guide; (2) customer deployment disclosures reducing concentration risk; (3) networking mix sustaining at 40%+; (4) buyback execution at current aggressive pace; (5) non-AI semiconductor recovery signals; (6) FY27 framework updates from management on GW pace and revenue per GW.")

    add_docx_page(document, "Street read-through")
    document.add_heading("Street Read-Through And Broker Debate", level=1)
    add_docx_table(document, BROKER_GRID, [1.4, 0.8, 4.6])
    document.add_paragraph("J.P. Morgan ($500) anchors on customer-by-customer GW framework and $12-15bn revenue per GW, pushing FY27 AI revenue to $120bn+. UBS ($475) focused on margin resilience; its supply-chain work suggests ~6M TPU units next year with AI revenue above $130bn, but applied ~21x CY27 P/E to discount sustainability. Morgan Stanley ($470) emphasized rack-margin fears receding and constructive COT commentary. Goldman Sachs (Buy, CL) highlighted supply-chain assurance through FY28 and revised AI estimates to $60bn/$130bn/$170bn for FY26/27/28. Jefferies argues the $100bn FY27 floor likely proves conservative with revenue per GW potentially reaching $20-25bn. HSBC ($535) remains most aggressive, arguing Street still underestimates ASIC plus networking upside.")
    add_docx_table(document, DEBATE_MATRIX, [1.8, 2.4, 2.3])
    document.add_picture(str(CROP_DIR / "ubs_crop_tpu.png"), width=Inches(6.8))
    document.add_picture(str(CROP_DIR / "hsbc_crop_revisions.png"), width=Inches(6.8))
    document.add_paragraph("The disagreement across brokers is useful: it shows the real questions are not whether Broadcom has AI exposure, but how far the current framework stretches into FY27-FY28, whether networking monetization sustains, and whether software resilience deserves a premium multiple. That is a healthier debate for a long position than binary quarter-to-quarter surprise.")

    add_docx_page(document, "Appendix")
    document.add_heading("Detailed Estimates And Portfolio Framing", level=1)
    add_docx_table(document, ESTIMATE_MATRIX, [2.2, 1.1, 1.1, 1.1])
    add_docx_table(document, CLIENT_FIT, [1.9, 4.5])
    add_docx_table(document, MONITORING_FRAME, [1.7, 3.0, 2.0])
    document.add_paragraph("House view: Broadcom remains an attractive AI infrastructure holding for clients who want real earnings and free cash flow behind the theme. Outperform with $500 target. Best positioned as a core AI/technology holding within a diversified equity allocation — the combination of forward FCF yield, aggressive capital returns, and structural earnings growth gives it a quality-growth profile. For clients who already own Nvidia, Broadcom offers complementary custom-silicon and networking exposure.")
    document.add_paragraph("We would downgrade to Market Perform if: AI revenue guidance disappoints, margins compress 200bps+, software declines sequentially for two quarters, or hyperscaler capex budgets flatten. We would raise target above $500 if: FY27 AI revenue tracks above $120bn, networking holds 40%+ of AI, or new customers ramp faster than expected.")
    document.add_paragraph("The cleanest client message: Broadcom participates in a very large infrastructure buildout, monetizes both compute and networking bottlenecks, has a software cash-flow base most chip companies lack, and is already returning capital at a rate that makes it investable for mainstream portfolios. The emerging robotics demand wave adds further confidence in duration.")

    document.save(str(DOCX_PATH))


def main() -> None:
    ensure_dirs()
    crop_source_charts()
    charts = generate_charts()
    make_pdf(charts)
    make_docx(charts)
    print(PDF_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
