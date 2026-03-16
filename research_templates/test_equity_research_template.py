from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer


def normalize_text(value: str) -> str:
    return (
        value.replace("\u2014", "-")
        .replace("\u2013", "-")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
    )


def workspace_root() -> Path:
    return Path(__file__).resolve().parents[2]


def load_run_config(root: Path) -> dict:
    return json.loads((root / "context" / "run-config.json").read_text(encoding="utf-8"))


def load_text_snippets(root: Path) -> list[dict]:
    snippets: list[dict] = []
    for path in sorted((root / "extracted" / "text").glob("*.txt")):
        text = normalize_text(path.read_text(encoding="utf-8", errors="ignore")).strip()
        if not text:
            continue
        snippets.append(
            {
                "name": path.stem,
                "text": text,
                "excerpt": text[:900].strip(),
            }
        )
    return snippets


def load_visuals(root: Path) -> list[Path]:
    visuals = list((root / "generated" / "charts").glob("*.png"))
    visuals.extend((root / "generated" / "diagrams").glob("*.png"))
    return sorted(visuals)


def build_sections(root: Path, config: dict) -> list[tuple[str, list[str]]]:
    snippets = load_text_snippets(root)
    objective = normalize_text((root / "context" / "objective.md").read_text(encoding="utf-8").strip())
    source_points = [
        f"{snippet['name']}: {snippet['excerpt']}" for snippet in snippets[:4]
    ] or ["No extracted source text was available when the template ran."]

    return [
        (
            "Executive Summary",
            [
                f"Objective: {objective or 'Research request staged from the app.'}",
                "This template is a working baseline that Codex can edit in-place for a deeper final report.",
                "The draft is intentionally source-backed and keeps the tone plain, direct, and finance-friendly.",
            ],
        ),
        (
            "Source-Backed Notes",
            source_points,
        ),
        (
            "How To Extend This Template",
            [
                "Add deeper thesis sections, valuation commentary, and custom charts before generating the final client-ready draft.",
                "Prefer full-sentence bullets and short analytical paragraphs over compressed fragments.",
                "Keep all finished outputs inside result/ so the desktop app can open them automatically.",
            ],
        ),
    ]


def build_docx(root: Path, config: dict) -> Path:
    output_path = root / "result" / "report.docx"
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    document.add_heading("Research Report Draft", level=0)
    document.add_paragraph(
        "Prepared from staged source material. This file is designed to be edited by Codex before the final handoff."
    )

    for title, bullets in build_sections(root, config):
        document.add_heading(title, level=1)
        for bullet in bullets:
            document.add_paragraph(normalize_text(bullet), style="List Bullet")

    for visual in load_visuals(root):
        document.add_heading(visual.stem.replace("_", " ").title(), level=2)
        document.add_picture(str(visual), width=Inches(6.0))
        document.add_paragraph(f"Source visual: {visual.name}")

    document.save(output_path)
    return output_path


def build_pdf(root: Path, config: dict) -> Path:
    output_path = root / "result" / "report.pdf"
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=LETTER,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.85 * inch,
        bottomMargin=0.85 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = styles["Heading1"]
    body_style = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor("#1f2933"),
        spaceAfter=8,
    )

    story = [
        Paragraph("Research Report Draft", title_style),
        Paragraph(
            "Prepared from staged source material. This PDF is a working baseline that Codex can deepen and refine.",
            body_style,
        ),
        Spacer(1, 0.18 * inch),
    ]

    for title, bullets in build_sections(root, config):
        story.append(Paragraph(title, title_style))
        for bullet in bullets:
            story.append(Paragraph(f"- {normalize_text(bullet)}", body_style))
        story.append(Spacer(1, 0.12 * inch))

    for visual in load_visuals(root):
        story.append(Paragraph(visual.stem.replace("_", " ").title(), title_style))
        story.append(Image(str(visual), width=6.0 * inch, height=3.2 * inch))
        story.append(Paragraph(f"Source visual: {visual.name}", body_style))
        story.append(Spacer(1, 0.12 * inch))

    doc.build(story)
    return output_path


def build_workbook(root: Path) -> Path:
    output_path = root / "result" / "valuation.xlsx"
    workbook = Workbook()
    dcf = workbook.active
    dcf.title = "DCF"
    dcf["A1"] = "Assumption"
    dcf["B1"] = "Value"
    dcf["A2"] = "WACC"
    dcf["B2"] = 0.1
    dcf["A3"] = "Terminal growth"
    dcf["B3"] = 0.03
    dcf["A5"] = "This workbook is a structured starting point."

    comps = workbook.create_sheet("Comps")
    comps.append(["Company", "EV/EBITDA", "P/E", "Note"])
    comps.append(["Peer A", 12.0, 18.0, "Replace with real comps."])
    comps.append(["Peer B", 14.0, 21.0, "Replace with real comps."])

    scenarios = workbook.create_sheet("Scenarios")
    scenarios.append(["Scenario", "Value / Share", "Comment"])
    scenarios.append(["Bear", 0, "Fill with downside case."])
    scenarios.append(["Base", 0, "Fill with base case."])
    scenarios.append(["Bull", 0, "Fill with upside case."])

    header_fill = PatternFill(fill_type="solid", fgColor="1E2D6F")
    header_font = Font(color="FFFFFF", bold=True)
    for sheet in workbook.worksheets:
        for cell in sheet[1]:
            cell.fill = header_fill
            cell.font = header_font
        sheet.column_dimensions["A"].width = 22
        sheet.column_dimensions["B"].width = 18
        sheet.column_dimensions["C"].width = 18
        sheet.column_dimensions["D"].width = 36

    workbook.save(output_path)
    return output_path


def main() -> None:
    root = workspace_root()
    config = load_run_config(root)
    output_format = config.get("output_format", "docx").lower()

    if output_format == "pdf":
        build_pdf(root, config)
    else:
        build_docx(root, config)

    if config.get("valuation_required"):
        build_workbook(root)


if __name__ == "__main__":
    main()
